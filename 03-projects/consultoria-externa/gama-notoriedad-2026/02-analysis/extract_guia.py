import sys
sys.stdout.reconfigure(encoding='utf-8')
import os
import json
from docx import Document

src = r"G:\Mi unidad\RAUL\colaboradores\Genteca\Cora-Urrea\01_De_Cora_Para_Raoul\Notoriedad V2.0\GUIA DE PREGUNTAS NOTORIEDAD 2026.docx"
out_md = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\02-analysis\guia_preguntas_2026.md"

doc = Document(src)
lines = []
for p in doc.paragraphs:
    text = p.text.strip()
    if text:
        runs_info = []
        for r in p.runs:
            color = None
            try:
                if r.font.color and r.font.color.rgb:
                    color = str(r.font.color.rgb)
            except Exception:
                pass
            runs_info.append((r.text, color))
        # Marcar lineas con rojo (color con FF en el primer hex)
        has_red = any(c and c.upper().startswith('FF') and ('00' in c[2:4]) for _, c in runs_info if c)
        marker = "[ROJO] " if has_red else ""
        lines.append(f"{marker}{text}")

# Tambien procesar tablas si las hay
for ti, table in enumerate(doc.tables):
    lines.append(f"\n--- TABLA {ti+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        lines.append(" | ".join(cells))

content = "\n".join(lines)
with open(out_md, 'w', encoding='utf-8') as f:
    f.write(content)

print(f"Guardado: {out_md}")
print(f"Lineas: {len(lines)}")
print(f"Caracteres: {len(content):,}")
