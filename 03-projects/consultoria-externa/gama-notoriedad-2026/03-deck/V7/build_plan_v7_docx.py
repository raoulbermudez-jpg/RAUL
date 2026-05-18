import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V7\2026-05-18_Plan-V7_Notoriedad-Gama.docx"
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)

RED_GAMA  = RGBColor(0x7A, 0x12, 0x12)
RED_DARK  = RGBColor(0x4A, 0x0A, 0x0A)
GRAY_DARK = RGBColor(0x33, 0x33, 0x33)
GRAY_MID  = RGBColor(0x66, 0x66, 0x66)
BLUE_INFO = RGBColor(0x1A, 0x56, 0x8A)
ORANGE_W  = RGBColor(0xD9, 0x73, 0x06)
GREEN_OK  = RGBColor(0x1A, 0x70, 0x2A)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color='808080', size='4'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_para(doc, text, size=11, bold=False, italic=False, color=GRAY_DARK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RED_GAMA


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RED_DARK


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = GRAY_DARK


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("•  ")
    run.font.size = Pt(size)
    run.bold = True
    run.font.color.rgb = RED_GAMA
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GRAY_DARK


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, '7A1212')
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            set_cell_borders(cell)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'F5F5F5')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()


def add_callout(doc, text, color=RED_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("★ ")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = color
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = color


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # PORTADA
    add_para(doc, "PLAN V7 — Notoriedad Gama 2026",
             size=24, bold=True, color=RED_GAMA, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)
    add_para(doc, "Documento de Plan tras reunión Cora + Raoul 18-05-2026",
             size=14, italic=True, color=GRAY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, "Para revisión y aprobación de Cora antes de ejecutar CU-9 + Word V7",
             size=11, italic=True, color=GRAY_MID, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "Fecha: 2026-05-18  ·  Confidencial NDA  ·  Cora Urrea + Raoul Bermudez",
             size=10, color=GRAY_MID, align=WD_ALIGN_PARAGRAPH.CENTER)

    # CONTEXTO
    add_h1(doc, "1. Contexto")
    add_para(doc, "Tras la reunión del 18/05 a las 9:38, Cora y Raoul acordaron rehacer la propuesta V7 incorporando: "
                  "(a) bifurcación sistemática Total + C+/C ('clase media') + Pref-Gama, (b) análisis profundo de "
                  "importancia atributos + posicionamiento + mundo de marca, (c) recuperación de análisis V3 que "
                  "se habían perdido en V5/V6, (d) NO incluir cualitativo del 2026 como conclusión (era de otro "
                  "estudio: App Gama + Gama Club), (e) HTML interactivo con data bruta, (f) documento metodológico separado, "
                  "(g) workflow Word V7 primero → revisión → PPTX después.")

    # APROBACIONES
    add_h1(doc, "2. Aprobaciones explícitas de Cora")
    add_bullet(doc, "Plan V7 aprobado en bloque")
    add_bullet(doc, "HTML interactivo: alcance simple (tabla cruzada con 5 variables de cruce)")
    add_bullet(doc, "Documento metodológico: formato .docx")
    add_bullet(doc, "Recuperar análisis V3 e incluirlos en V7")
    add_bullet(doc, "Validar consistencia entre análisis V3 recuperados y datos actuales. Si NO son consistentes, explicar por qué y avisar antes de proceder.")

    # ESTRUCTURA WORD V7
    add_h1(doc, "3. Estructura del Word V7")
    add_h3(doc, "Capítulos del cuerpo principal")
    headers = ["#", "Capítulo", "Bifurcación"]
    rows = [
        ["0", "Marco metodológico (resumen — detalle en doc separado)", "—"],
        ["1", "Embudo de marcas + drivers de preferencia espontáneos (P21.1)",
         "Total + C+/C, comparativo 2025-2026"],
        ["2", "Posicionamiento + Mundo de marca + DNA competitivo",
         "Total + C+/C + Pref-Gama (3ª capa)"],
        ["3", "Precios y categorías + misiones de compra",
         "Total + C+/C + Pref-Gama donde aplique"],
        ["4", "Efectividad comunicacional PTL/DTLS",
         "Re-segmentación por recall + cruces con posicionamiento"],
        ["5", "Síntesis estratégica y recomendaciones priorizadas",
         "Bifurcadas por segmento"],
    ]
    add_table(doc, headers, rows, col_widths=[1, 7, 6])

    add_h3(doc, "Anexos")
    add_bullet(doc, "Anexo A — Geografía / municipios (secundario, no central — Cora pidió explícitamente no comenzar por aquí)")
    add_bullet(doc, "Anexo B — Hipótesis Express formato (compatible con datos pero variable formato no en BBDD)")
    add_bullet(doc, "Anexo C — P43-P45 El Recreo apertura sucursal (táctico, no modifica análisis principal)")

    add_h3(doc, "Entregables adicionales")
    add_bullet(doc, "Word V7 principal (estructura arriba)")
    add_bullet(doc, "Documento metodológico separado (.docx) — qué análisis se usó, por qué, dónde")
    add_bullet(doc, "HTML interactivo con data bruta + variables de cruce (edad, clase social, género, marca preferida, municipio — C+/C ya unido en BBDD)")

    # ANÁLISIS V3 RECUPERADOS
    add_h1(doc, "4. Análisis V3 a recuperar")
    add_para(doc, "Cuatro análisis específicos del V3 que se habían perdido en V5/V6 y que vamos a reincorporar:")

    add_h3(doc, "4.1  DNA de Gama por z-scores (V3 Slide 13)")
    add_para(doc, "Gama sobreindexa significativamente en 4 atributos experienciales y subindexa en 3 atributos precio/valor "
                  "(z-scores vs media del mercado):")
    headers = ["Atributo", "Z-score Gama", "Tipo"]
    rows = [
        ["Tienda atractiva", "+1.09", "Sobreíndice (DNA)"],
        ["Calidad", "+0.97", "Sobreíndice (DNA)"],
        ["Seguro", "+0.76", "Sobreíndice (DNA)"],
        ["Limpieza", "+0.72", "Sobreíndice (DNA)"],
        ["Menor precio", "-0.72", "Subíndice"],
        ["Hacer valer dinero", "-0.67", "Subíndice"],
        ["Promociones", "-0.67", "Subíndice"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 4, 5])
    add_para(doc, "Para V7: recalcular sobre BBDD 2026 con bifurcación Total + C+/C + Pref-Gama. Comparar contra valores 2025.",
             italic=True, color=GRAY_MID, size=10)

    add_h3(doc, "4.2  Modelo mental: Atención-dominante vs Precio-dominante (V3 Slide 29)")
    add_para(doc, "Mercado dividido en dos modelos mentales según razón espontánea de preferencia (P21.1):")
    headers = ["Modelo", "Cadenas", "% razón dominante"]
    rows = [
        ["Atención-dominante", "Gama, Central Madeirense, Rio", "Gama 53%, CM 53%, Rio 51% (atención)"],
        ["Precio-dominante", "Páramo, Plan Suárez, La Granja, Forum, Plazas",
         "Páramo 81%, Plan Suárez 71%, La Granja 82%, Forum 52%, Plazas 60% (precio)"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 5, 7])
    add_para(doc, "Para V7: recalcular sobre BBDD 2026 con bifurcación Total + C+/C. Validar si Rio sigue en el modelo atención-dominante o migró a precio-dominante (CU-7 WoW muestra Rio creciendo agresivamente en TOM).",
             italic=True, color=GRAY_MID, size=10)

    add_h3(doc, "4.3  Tres segmentos K-means y perfil del Núcleo Leal (V3 Slides 24-27)")
    add_para(doc, "K-means k=3 identifica tres segmentos con perfiles diferenciados:")
    headers = ["Segmento", "Tamaño", "Perfil clave"]
    rows = [
        ["Seg 1 — Mayoría Exigente", "59% (n≈237)", "Alta exigencia, no exclusivos Gama, misión múltiple"],
        ["Seg 2 — Pragmáticos Convertibles", "33% (n≈133)",
         "Resistencia precio 3.44/5 (vs 3.66 Seg1), 0% pref-Gama actual"],
        ["Seg 3 — Núcleo Leal", "8% (n=32, ref.)",
         "100% pref-Gama, NSE 2.16, precio 2.94/5, rapidez 4.81/5, atención 4.69/5"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 3, 8])
    add_para(doc, "Para V7: vigentes sin cambios. Incorporar como capítulo principal con estrategia diferenciada por segmento.",
             italic=True, color=GRAY_MID, size=10)

    add_h3(doc, "4.4  Cinco temas cualitativos V3 (V3 Slide 28)")
    add_bullet(doc, "Tema 1 — Atención como diferenciador simbólico (53% razón espontánea + triple convergencia)")
    add_bullet(doc, "Tema 2 — Precio como zona de tensión no resuelta (paradoja 54% caro vs OR=1.03 NS)")
    add_bullet(doc, "Tema 3 — Cercanía como primer escalón de la relación (40.6% razón #2)")
    add_bullet(doc, "Tema 4 — Campaña como oportunidad no aprovechada (recall 4.2%)")
    add_bullet(doc, "Tema 5 — Modelo de dos velocidades (Seg 3 ya resolvió, Seg 1+2 no)")
    add_para(doc, "Para V7: estos 5 temas se mantienen, pero solo como validación interpretativa del cuantitativo. "
                  "NO se incluye el cualitativo del 2026 (App Gama + Gama Club) como conclusión propia, como pidió Cora.",
             italic=True, color=GRAY_MID, size=10)

    # VALIDACIÓN CONSISTENCIA
    add_h1(doc, "5. Validación de consistencia V3 vs análisis actual — preliminar")
    add_callout(doc, "Conclusión preliminar: NO se detectan tensiones críticas. V3 es consistente direccionalmente "
                     "con los análisis V4-V6. CU-9 hará la validación formal con cifras exactas.")
    headers = ["Análisis V3", "Hallazgo V3", "Estado en análisis actual (V4-V6)", "¿Consistente?"]
    rows = [
        ["Modelo Atención-dominante",
         "Gama 53% razón atención",
         "CU-3: OR=5.73 atención con convergencia 4 métodos (logit + RF + SHAP + razón espontánea)",
         "✅ CONSISTENTE (refuerza)"],
        ["DNA z-scores experienciales",
         "Gama sobreíndice +0.7 a +1.1 en 4 atributos",
         "CU-7 WoW: tendencia +6.3pp 'atractiva', +5.8pp 'seguro' (no sig pero direccional)",
         "✅ CONSISTENTE (direccional)"],
        ["DNA z-scores precio negativo",
         "Gama subíndice -0.7 en menor precio",
         "CU-8 P32: Gama no lidera precio en ninguna categoría de alta rotación, gap -33pp en proteínas vs Páramo",
         "✅ CONSISTENTE (más fuerte)"],
        ["3 segmentos k-means",
         "Seg 3 Núcleo Leal n=32, NSE 2.16, precio 2.94/5",
         "CU-4 v2: idéntico (mismo dataset, mismo algoritmo)",
         "✅ IDÉNTICO"],
        ["5 temas cualitativos V3",
         "Atención, Precio, Cercanía, Campaña, 2 velocidades",
         "IN-7 v2: confirma + agrega 'sifrinaje' (barrera identitaria) y 'arquetipo femenino'",
         "✅ EXPANDIDO (no contradice)"],
        ["Modelo mental Rio = Atención-dominante",
         "Rio 51% atención en V3",
         "CU-7 WoW: Rio crece +17pp TOM (sig 99%). Necesita re-verificar si su razón espontánea sigue siendo atención o migró a precio.",
         "⚠ VERIFICAR en CU-9"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 4, 6, 3])

    add_para(doc, "Punto de atención: solo el modelo mental de Rio merece verificación adicional. Rio creció agresivamente "
                  "en TOM (+17pp sig 99%) — vale verificar si su razón espontánea de preferencia (P21.1) sigue siendo "
                  "atención (51% en V3) o migró hacia otro atributo. Esto lo confirma CU-9.",
             italic=True, color=BLUE_INFO)

    # ANÁLISIS CU-9
    add_h1(doc, "6. Análisis nuevos requeridos (CU-9)")
    add_para(doc, "Cuanti procesará 10 análisis nuevos sobre la BBDD 2026 (n=402):", size=11)
    headers = ["#", "Análisis", "Bifurcación"]
    rows = [
        ["1", "P22 importancia atributos", "Total + C+/C + Pref-Gama"],
        ["2", "P23 asociación marca×atributo (matriz completa)", "Total + C+/C + Pref-Gama"],
        ["3", "Correlación nominal P21.1 (razones espontáneas) × P23 (asociación)", "Total + C+/C"],
        ["4", "Mapa de posicionamiento perceptual multi-segmento", "Total + C+/C + Pref-Gama"],
        ["5", "Mundo de marca / DNA: atributos únicos vs compartidos por cadena (recálculo V3 Slide 13)",
         "Total + C+/C + Pref-Gama"],
        ["6", "Modelo Atención-dominante vs Precio-dominante (recálculo V3 Slide 29, verifica Rio)", "Total + C+/C"],
        ["7", "Perfil del recordador PTL y DTLS (P35-P42 × demográficos + preferencia)", "Total"],
        ["8", "Re-segmentación de TODAS las preguntas según recall publicitario", "Recall sí / no"],
        ["9", "Cruce P30 (hábito por categoría) × P32 (mejor precio por categoría) × P26 (misión)",
         "Total + C+/C"],
        ["10", "Categorías 'cuidar precio' vs 'ofrecer valor' — síntesis estratégica",
         "Total + C+/C + Pref-Gama"],
    ]
    add_table(doc, headers, rows, col_widths=[1, 9, 5])

    add_para(doc, "Adicionalmente CU-9 validará formalmente la consistencia V3 con datos actuales y documentará "
                  "cualquier tensión detectada para revisión Owner antes de proceder con el Word V7.",
             italic=True, color=BLUE_INFO)

    # LO QUE NO VA AL ESTUDIO
    add_h1(doc, "7. Lo que NO va al estudio (decisión explícita Cora)")
    add_bullet(doc, "Cualitativo 2026 como conclusión (era de otro estudio: App Gama + Gama Club). Solo se usa para confirmar posicionamiento cuantitativo.")
    add_bullet(doc, "Recomendaciones digitales (fuera del scope original de notoriedad y brand health)")
    add_bullet(doc, "Mensajes alternativos para PTL/DTLS (Cora no quiere modificar mensajes — solo validar y entender perfil de quien recordó)")
    add_bullet(doc, "Análisis geográfico como capítulo central (queda como Anexo A)")
    add_bullet(doc, "Hipótesis Express como conclusión (queda como Anexo B)")
    add_bullet(doc, "Comparativo 2025-2026 en precios o comunicación (solo en embudo + posicionamiento)")

    # WORKFLOW
    add_h1(doc, "8. Workflow operativo aprobado")
    add_para(doc, "Acordado en reunión Cora + Raoul 18/05:")
    headers = ["#", "Paso", "Output"]
    rows = [
        ["1", "Owner + Cora aprueban este plan", "Plan V7 aprobado"],
        ["2", "Drop del plan a Drive Cora (este documento)", "Plan-V7.docx en Drive"],
        ["3", "Cuanti CU-9 procesa los 10 análisis + valida consistencia V3", "CU-9 reporte + JSON"],
        ["4", "Main Claude escribe Word V7", "Word V7 .docx"],
        ["5", "Drop Word V7 a Drive Cora", "Word V7 en Drive Cora"],
        ["6", "Cora + Owner revisan Word V7", "Feedback"],
        ["7", "Si OK → generar documento metodológico (.docx) + HTML interactivo", "Doc metodológico + HTML"],
        ["8", "Vivienne genera PPTX (3 invocaciones separadas, tono rojo corregido)", "PPTX deck + RE"],
        ["9", "Drop final completo + draft email", "Entrega final"],
    ]
    add_table(doc, headers, rows, col_widths=[1, 8, 6])

    # TIEMPOS
    add_h1(doc, "9. Tiempos estimados")
    headers = ["Etapa", "Tiempo estimado"]
    rows = [
        ["Cuanti CU-9 (10 análisis + recuperación V3 + validación consistencia)", "60-90 min"],
        ["Word V7 (estructura, contenido, tablas, charts inline)", "60-90 min"],
        ["Documento metodológico (.docx separado)", "30 min"],
        ["HTML interactivo simple (tabla cruzada con 5 variables)", "60-90 min"],
        ["PPTX V7 (Vivienne 3 invocaciones, anti-token-explosion)", "90 min"],
        ["Drop final + draft email a Gama", "20 min"],
        ["TOTAL ESTIMADO", "~5-6 horas"],
    ]
    add_table(doc, headers, rows, col_widths=[10, 5])

    # PUNTO DE ATENCIÓN
    add_h1(doc, "10. Único punto a verificar antes de arrancar")
    add_callout(doc, "Cora: ¿confirmas que el plan está OK tal como está, o querés ajustar algo antes de que Cuanti arranque CU-9?",
                color=BLUE_INFO)
    add_para(doc, "Una vez que confirmes, Cuanti CU-9 arranca en background (~60-90 min). Mientras tanto preparamos "
                  "los esqueletos del Word V7 + documento metodológico + script HTML interactivo.")

    add_para(doc, "")
    add_para(doc, "—  FIN DEL PLAN V7  —",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RED_GAMA, bold=True, size=12)
    add_para(doc, "Cora Urrea + Raoul Bermudez  ·  2026-05-18",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY_MID, size=10)

    doc.save(OUTPUT)
    print(f"Guardado: {OUTPUT}")


if __name__ == '__main__':
    main()
