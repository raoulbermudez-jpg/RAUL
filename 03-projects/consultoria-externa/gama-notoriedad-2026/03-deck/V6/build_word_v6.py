"""
Word V6 — Notoriedad Gama 2026
Documento formal de consulta + base para PPTX posterior.

Estructura: portada + ToC + 11 secciones principales + anexos.
Charts: 9 charts V5 reusados + 11 charts V6 nuevos = ~20 charts inline.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# ============================================================================
# CONFIGURACIÓN
# ============================================================================

OUTPUT = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V6\2026-05-18_Notoriedad-Gama-2026_V6.docx"
CHARTS_V5 = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V5\charts"
CHARTS_V6 = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V6\charts"
CU8_HEATMAP = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\02-analysis\cuanti\outputs\plots\cu8_v2_heatmap_desconexion.png"

RED_GAMA  = RGBColor(0x7A, 0x12, 0x12)
RED_DARK  = RGBColor(0x4A, 0x0A, 0x0A)
BLUE_INFO = RGBColor(0x1A, 0x56, 0x8A)
GRAY_MID  = RGBColor(0x66, 0x66, 0x66)
GRAY_DARK = RGBColor(0x33, 0x33, 0x33)
ORANGE_W  = RGBColor(0xD9, 0x73, 0x06)
GREEN_OK  = RGBColor(0x1A, 0x70, 0x2A)

# ============================================================================
# HELPERS
# ============================================================================

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color='808080', size='4'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_para(doc, text, size=11, bold=False, italic=False, color=GRAY_DARK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RED_GAMA
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = RED_DARK
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY_DARK
    return p


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    run.italic = True
    run.font.name = 'Calibri'
    run.font.color.rgb = BLUE_INFO
    return p


def add_bullet(doc, text, level=0, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7 + level*0.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run("•  ")
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = RED_GAMA
    run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = GRAY_DARK


def add_table(doc, headers, rows, col_widths=None, header_color='7A1212'):
    """Crea tabla formateada con encabezados rojos y bordes."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, header_color)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            set_cell_borders(cell)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'F5F5F5')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
    doc.add_paragraph()
    return table


def add_chart(doc, chart_path, width_in=6.5, caption=None):
    if not os.path.exists(chart_path):
        add_para(doc, f"[Chart faltante: {os.path.basename(chart_path)}]", italic=True, color=ORANGE_W)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_MID
        run.font.name = 'Calibri'
    doc.add_paragraph()


def add_metodologia_box(doc, text):
    """Nota metodológica anexa — caja gris al final de sección."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run("📋 Nota metodológica  ·  ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = BLUE_INFO
    run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = GRAY_MID
    run.font.name = 'Calibri'


def add_caveat_box(doc, text):
    """Caja de caveat (naranja)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run("⚠ Caveat  ·  ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ORANGE_W
    run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = ORANGE_W
    run.font.name = 'Calibri'


def add_callout_box(doc, text, color=RED_DARK):
    """Caja de hallazgo destacado."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run("★ ")
    run.font.size = Pt(11)
    run.font.color.rgb = color
    run.bold = True
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = color
    run.font.name = 'Calibri'


def add_verbatim(doc, quote, attribution):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'"{quote}"')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run(f"— {attribution}")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_MID
    run.italic = True


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("◆ ◆ ◆")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_MID


# ============================================================================
# SECCIONES
# ============================================================================

def build_portada(doc):
    # Título principal
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOTORIEDAD Y PREFERENCIA DE MARCA")
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RED_GAMA
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gama  ·  2026")
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = RED_GAMA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nVersión V6  ·  Documento de Consulta")
    run.font.size = Pt(18)
    run.font.color.rgb = GRAY_DARK
    run.italic = True

    for _ in range(3):
        doc.add_paragraph()

    # Metadatos
    metadata = [
        ("Fecha", "2026-05-18"),
        ("Equipo analítico", "Cora Urrea + Raoul Bermúdez"),
        ("Tipo de documento", "Reporte de Investigación de Mercado"),
        ("Base muestral", "n=402 (2026)  ·  n=785 (2025 referencia)"),
        ("Margen de error", "±4.89% al 95% de confianza (Total)"),
        ("Confidencialidad", "NDA  ·  Uso interno Gama y consultoría"),
    ]
    t = doc.add_table(rows=len(metadata), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(metadata):
        cell_k = t.rows[i].cells[0]
        cell_v = t.rows[i].cells[1]
        cell_k.width = Cm(5)
        cell_v.width = Cm(10)
        p = cell_k.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(k + ":")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RED_DARK
        p = cell_v.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(v)
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY_DARK

    for _ in range(4):
        doc.add_paragraph()

    add_para(doc, "Este documento es la base de consulta del estudio Notoriedad Gama 2026. "
                  "Contiene el análisis completo desde los datos de Notoriedad 2.0, incluye "
                  "todas las preguntas del cuestionario, integra los insights cualitativos como "
                  "hipótesis (no como conclusiones), y queda como referencia técnica para Gama y "
                  "para el diseño de la próxima ola.",
             size=10, italic=True, color=GRAY_MID, align=WD_ALIGN_PARAGRAPH.CENTER)


def build_toc(doc):
    add_h1(doc, "Índice de contenido")

    toc_items = [
        ("1.", "Resumen Ejecutivo", ""),
        ("2.", "Marco del documento V6 — qué es, cómo leerlo, iconografía", ""),
        ("3.", "Ficha técnica del estudio", ""),
        ("", "", ""),
        ("PARTE I.", "Análisis Cuantitativo 2026", ""),
        ("4.1", "NSE C+/C — el segmento natural de Gama", ""),
        ("4.2", "Salud de marca Gama — embudo 2026", ""),
        ("4.3", "Posición competitiva 2026", ""),
        ("4.4", "Drivers de preferencia", ""),
        ("4.5", "Análisis profundo de PRECIOS (capítulo nuevo)", ""),
        ("4.6", "Lugar de compra × precio × preferencia × misiones (NEW)", ""),
        ("4.7", "Balance óptimo de categorías a igualar precio", ""),
        ("4.8", "Subset de shoppers con percepción favorable", ""),
        ("4.9", "Comportamiento y misiones de compra", ""),
        ("4.10", "Comunicación: PTL vs DTLS (análisis completo)", ""),
        ("4.11", "Evolución comparativa 2025→2026", ""),
        ("", "", ""),
        ("PARTE II.", "Insights Cualitativos — hipótesis y señales", ""),
        ("5.1", "Regla de lectura del bloque cualitativo", ""),
        ("5.2", "Mecanismo del driver de atención: \"acompañamiento-guía\"", ""),
        ("5.3", "Barrera del precio: hipótesis del \"sifrinaje\"", ""),
        ("5.4", "Datos digitales como contexto (sin recomendaciones)", ""),
        ("5.5", "Arquetipo femenino emergente", ""),
        ("", "", ""),
        ("PARTE III.", "Hipótesis especiales", ""),
        ("6.1", "Hipótesis del formato Express — validación parcial", ""),
        ("", "", ""),
        ("PARTE IV.", "Síntesis y Recomendaciones", ""),
        ("7.1", "Tesis V6 — diagnóstico central", ""),
        ("7.2", "Los 4 argumentos MECE", ""),
        ("7.3", "Recomendaciones priorizadas", ""),
        ("7.4", "Dos lecturas válidas — decisión Owner CO-2", ""),
        ("", "", ""),
        ("PARTE V.", "Agenda Ola 2027", ""),
        ("8.1", "Hipótesis abiertas que requieren ola 2027", ""),
        ("8.2", "5 mejoras metodológicas obligatorias", ""),
        ("", "", ""),
        ("9.", "Análisis comparativo V1-V5 + lecciones aprendidas", ""),
        ("10.", "Anexo metodológico", ""),
        ("11.", "Anexo táctico — P43-P45 El Recreo (fuera del estudio principal)", ""),
    ]

    t = doc.add_table(rows=len(toc_items), cols=2)
    for i, (num, title, _) in enumerate(toc_items):
        cell1 = t.rows[i].cells[0]
        cell2 = t.rows[i].cells[1]
        cell1.width = Cm(3)
        cell2.width = Cm(14)
        p = cell1.paragraphs[0]
        run = p.add_run(num)
        if "PARTE" in num:
            run.bold = True
            run.font.color.rgb = RED_GAMA
            run.font.size = Pt(11)
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY_DARK
        p = cell2.paragraphs[0]
        run = p.add_run(title)
        if "PARTE" in num:
            run.bold = True
            run.font.color.rgb = RED_GAMA
            run.font.size = Pt(11)
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY_DARK


def build_executive_summary(doc):
    add_h1(doc, "1. Resumen Ejecutivo")

    add_callout_box(doc,
        "Gama tiene el activo relacional más diferenciador del mercado venezolano de "
        "supermercados (Atención: OR=5.73, convergente en 4 métodos) y su posición en 2026 "
        "es estadísticamente estable — pero comunica en el territorio equivocado, no activa al "
        "tercio de shoppers convertibles, y su segmento natural (NSE C+/C) está siendo "
        "penetrado por Río y Páramo a una velocidad que exige respuesta antes de la ola 2027.")

    add_h3(doc, "Los 5 hallazgos centrales del estudio 2026")

    add_para(doc, "1. Atención al cliente es el driver más fuerte y único de la preferencia por Gama.",
             bold=True, size=11)
    add_para(doc, "Quien asocia Gama con atención tiene 5.7 veces más probabilidades de "
                  "preferirla. La convergencia de 4 métodos independientes (logit, Random "
                  "Forest, SHAP, razón espontánea) es el mayor respaldo metodológico del "
                  "estudio. Rio y Páramo no tienen equivalente documentado en atención.", size=10)

    add_para(doc, "2. El precio no predice preferencia, pero la percepción de precio sí pesa.",
             bold=True, size=11)
    add_para(doc, "El atributo \"precio\" es el predictor #10 de preferencia (OR=1.03, no "
                  "significativo). Sin embargo, cuando se cruza la percepción global de precios "
                  "(P33) con la preferencia (P21), se observa un gradiente claro: pasar de "
                  "\"Gama es cara\" a \"Gama es igual o económica\" más que duplica la "
                  "probabilidad de preferencia (OR=2.4, IC95 [1.12, 5.12]). La paradoja se "
                  "resuelve así: el precio no es driver de elección entre quien ya conoce Gama, "
                  "pero la percepción de precio es barrera de entrada para quien no ha "
                  "experimentado el activo relacional.",
             size=10)

    add_para(doc, "3. Gama es la única gran cadena con posición estable en 2026.",
             bold=True, size=11)
    add_para(doc, "Cero de ocho indicadores del embudo de Gama muestra variación "
                  "estadísticamente significativa entre 2025 y 2026. En un mercado donde 9 de "
                  "10 cambios significativos son aumentos de competidores (Rio +17pp TOM, "
                  "Páramo +12pp TOM, ambos al 99% de confianza), la estabilidad de Gama es una "
                  "fortaleza defensiva resultado de sus activos diferenciadores.", size=10)

    add_para(doc, "4. El Segmento 2 (Pragmáticos Convertibles, 33% del mercado) es la mayor "
                  "oportunidad sin explotar.", bold=True, size=11)
    add_para(doc, "Este segmento tiene 0% de preferencia actual por Gama y la menor "
                  "resistencia al precio (3.44 vs 3.66 en escala 1-5). El cualitativo "
                  "sugiere que su barrera no es económica sino identitaria — la activación "
                  "correcta es una primera experiencia guiada, no un descuento.", size=10)

    add_para(doc, "5. Rio y Páramo están creciendo agresivamente en el segmento natural de Gama.",
             bold=True, size=11)
    add_para(doc, "En NSE C+/C, donde Gama tiene su posición más fuerte (TOM 60.6%), Rio creció "
                  "+25.8pp y Páramo +22.3pp en TOM entre 2025 y 2026 (exploratorio, requiere "
                  "confirmación en ola 2027). La ventana de respuesta ofensiva está abierta hoy.",
             size=10)

    add_h3(doc, "Las 3 decisiones del Owner que estructuran este reporte")
    add_para(doc, "Este documento aplica las siguientes decisiones, tomadas por el Owner el "
                  "2026-05-18, después de revisar las versiones V3, V4 y V5:", size=10)
    add_bullet(doc, "CO-1: La señal de alerta sobre Rio en proteínas se EXCLUYE del cuerpo "
                    "principal del documento (un solo verbatim cualitativo, sin sustento "
                    "cuantitativo de P32). Se traslada al memo interno.")
    add_bullet(doc, "CO-2: El tono presenta MEZCLA EQUILIBRADA — las dos lecturas válidas "
                    "(\"liderazgo defensivo sólido\" y \"señal de alerta competitiva\") se "
                    "muestran lado a lado con datos de soporte para cada una. La Junta de "
                    "Gama decide cuál enfatizar.")
    add_bullet(doc, "CO-3: Las recomendaciones sobre canal digital (Gama Club, inventario, "
                    "frescos online) NO se incluyen como recomendaciones del estudio. Los "
                    "datos digitales se comparten como contexto en una sección informativa, "
                    "fuera del scope del brief original.")

    add_metodologia_box(doc,
        "Este Resumen Ejecutivo sintetiza los hallazgos del análisis cuantitativo (n=402, "
        "margen de error ±4.89% al 95%) y los insights del análisis cualitativo (corpus de 12 "
        "documentos de focus groups, ~42 mil caracteres). Las afirmaciones \"hipótesis cuali\" "
        "no son proyectables estadísticamente. Detalle metodológico completo en Sección 10.")


def build_marco(doc):
    add_h1(doc, "2. Marco del documento V6")

    add_h3(doc, "Qué es V6")
    add_para(doc, "V6 es la versión consolidada y profundizada del estudio Notoriedad Gama 2026. "
                  "Es un documento independiente, autocontenido, diseñado tanto para presentación "
                  "ejecutiva como para consulta de referencia futura.")

    add_para(doc, "V6 incorpora todos los análisis previos (V1 a V5) corregidos y ampliados con:",
             bold=True)
    add_bullet(doc, "Las 9 preguntas del cuestionario que no habían sido procesadas en versiones "
                    "anteriores (PF8, PF9, PF10, P25, P26, P30, P32, P40-P42, P44, P45).")
    add_bullet(doc, "El capítulo profundo de precios (P31, P32, P33, P34) que reemplaza la cobertura "
                    "fragmentaria de V5.")
    add_bullet(doc, "El análisis completo de la frase \"De tu lado siempre\" (DTLS), que no se "
                    "había procesado en ninguna versión previa.")
    add_bullet(doc, "El capítulo nuevo de lugar de compra × precio × preferencia × misiones, "
                    "que responde directamente a las preguntas estratégicas de la consultora "
                    "del lado Gama.")
    add_bullet(doc, "La validación parcial de la hipótesis sobre el formato Express.")
    add_bullet(doc, "El análisis comparativo de las 5 versiones previas y las lecciones aprendidas.")

    add_h3(doc, "Cómo leer este documento")
    add_para(doc, "El documento sigue una arquitectura modular:")
    add_bullet(doc, "Para una lectura ejecutiva rápida: Sección 1 (Resumen Ejecutivo) + Sección 7 "
                    "(Síntesis y recomendaciones). 20 minutos.")
    add_bullet(doc, "Para una lectura analítica completa: Partes I a V. ~2 horas.")
    add_bullet(doc, "Para una consulta puntual: usar el índice (página 2) por número de sección.")
    add_bullet(doc, "Cada sección termina con una nota metodológica en azul que explica el análisis "
                    "técnico de los datos mostrados.")

    add_h3(doc, "Iconografía de certeza")
    add_para(doc, "Cada afirmación cuantitativa y cada hipótesis de este documento lleva un "
                  "ícono que indica su nivel de evidencia:")
    headers = ["Ícono", "Tipo", "Significado", "Cómo interpretar"]
    rows = [
        ["✅", "Conclusión cuanti", "Evidencia estadística significativa (≥95% confianza), n≥30",
         "Afirmación robusta — apta para decisiones estratégicas"],
        ["⚠", "Hipótesis apoyada", "Tendencia estadística (p 0.05-0.10) o soporte cualitativo robusto",
         "Dirección clara, requiere validación adicional"],
        ["💡", "Insight cuali", "Hallazgo de focus group — patrón en N grupos",
         "Útil para hipótesis y diseño; NO para extrapolación estadística"],
        ["📊", "Referencia evolutiva", "Dato 2025 como comparación (caveats aplican)",
         "Comparativo histórico — caveats de comparabilidad documentados"],
    ]
    add_table(doc, headers, rows, col_widths=[1.5, 3.5, 5.5, 5.5])

    add_h3(doc, "Lo que cambió respecto a V5")
    add_para(doc, "V6 mantiene todos los datos y cifras de V5 (que son a su vez idénticos a V4). "
                  "Lo que cambia es:")
    add_bullet(doc, "Profundidad analítica del tema precios (V5 tenía precio fragmentado en 4 lugares; "
                    "V6 dedica un capítulo completo).")
    add_bullet(doc, "Recuperación de la mención de Plan Suárez en el empate técnico de C+/C "
                    "(V3 lo tenía, V5 lo perdió).")
    add_bullet(doc, "Análisis de DTLS completo (P40-P42), que no aparece en ninguna versión previa.")
    add_bullet(doc, "Capítulo de lugar × precio × preferencia × misiones que responde literalmente "
                    "las preguntas estratégicas de Cora.")
    add_bullet(doc, "Eliminación de las recomendaciones digitales del cuerpo del documento (CO-3).")
    add_bullet(doc, "Eliminación de la señal de Rio en proteínas (CO-1).")
    add_bullet(doc, "Formato más limpio para presentación ejecutiva y para consulta futura.")

    add_metodologia_box(doc,
        "V6 reusa los datos cuantitativos de la BBDD Notoriedad 2026 (n=402) y la BBDD "
        "Notoriedad 2025 (n=785). Los análisis nuevos (CU-8 v1 y v2) procesan preguntas que "
        "no habían sido extraídas previamente. Los análisis previos (CU-1 a CU-7) se mantienen "
        "vigentes y se citan en las secciones correspondientes. El gate de Bruna BR-2 V4 sigue "
        "vigente sobre todos los claims que se mantienen de V5.")


def build_ficha_tecnica(doc):
    add_h1(doc, "3. Ficha técnica del estudio")

    headers = ["Dimensión", "Especificación"]
    rows = [
        ["Universo", "Shoppers regulares (≥1 compra/mes) en supermercados de cadena. Caracas + Altos Mirandinos."],
        ["Muestra principal 2026", "n=402 entrevistas face-to-face"],
        ["Margen de error Total 2026", "±4.89% al 95% de confianza (Wilson)"],
        ["NSE (proporción)", "C+/C combinado: n=104 (25.9%) · D: n=127 (31.6%) · E: n=171 (42.5%)"],
        ["Distribución geográfica", "Baruta (122) · Libertador (80) · Sucre (79) · Chacao (70) · El Hatillo (31) · Altos Mirandinos (20)"],
        ["Muestra referencia 2025", "n=785 — comparable en preguntas comunes (con caveats)"],
        ["Modelos estadísticos", "Logit (AUC=0.929, Pseudo R²=0.4371) · Random Forest + SHAP · K-means (k=3) · Z-test BH-FDR · Newcombe-Wilson IC95"],
        ["Subgrupo drivers (preferentes Gama)", "n=32 (referencial — IC95 amplios)"],
        ["Cualitativo", "12 documentos focus groups Gama (~42.094 caracteres) — 6 segmentos cubiertos"],
        ["Total preguntas del cuestionario", "PF1-PF10 perfil · P16-P42 cuerpo principal · P43-P45 anexo El Recreo"],
        ["Preguntas procesadas en V6", "100% de las preguntas con datos válidos en BBDD"],
        ["Fecha de campo", "[Pendiente confirmación con Cora]"],
        ["Versión del análisis", "V6 — 2026-05-18 (sucede a V5 del 2026-05-18)"],
    ]
    add_table(doc, headers, rows, col_widths=[4.5, 12])

    add_h3(doc, "Caveats metodológicos generales")
    add_para(doc, "Estos caveats aplican transversalmente a todo el documento. Caveats específicos "
                  "por sección se reportan al final de cada capítulo.")
    add_bullet(doc, "NSE C+/C combinado: en la base de datos no son separables C+ y C. n=104 implica "
                    "margen de error ±9.8% al 95% — análisis con mayor incertidumbre que el Total.")
    add_bullet(doc, "Sin ponderación 2025: el factor de ponderación muestral 2025 (variable "
                    "@PONDERAR_1) era todo cero en la base recibida. Los estimados 2025 pueden tener "
                    "sesgo de diseño si la muestra original era estratificada.")
    add_bullet(doc, "Composición geográfica difiere entre olas: Libertador está sobrerrepresentado "
                    "en 2025 vs 2026. Los cambios WoW interpretar con precaución.")
    add_bullet(doc, "Análisis WoW por NSE no pre-registrados en diseño 2025 — interpretar como "
                    "hipótesis a confirmar en ola 2027.")
    add_bullet(doc, "Subgrupo preferentes de Gama (n=32) es referencial — IC95 amplios. La "
                    "robustez de los hallazgos sobre drivers se basa en convergencia de 4 métodos, "
                    "no en el tamaño del subgrupo.")
    add_bullet(doc, "Análisis cualitativo realizado por analista único (sin Kappa inter-codificador). "
                    "Práctica estándar en investigación aplicada cuando recursos no permiten doble "
                    "codificación.")

    add_metodologia_box(doc,
        "La metodología completa del estudio está documentada en los artefactos técnicos internos: "
        "ME-1 (research design), ME-3 (methodology rationale), ME-4 (vigilance Q2 2026), y ME-5 "
        "(methodology plan V4 — vigente para V6). Disponibles bajo solicitud.")


# ============================================================================
# PARTE I — ANÁLISIS CUANTITATIVO
# ============================================================================

def build_parte1_header(doc):
    add_h1(doc, "PARTE I — Análisis Cuantitativo 2026")
    add_para(doc, "Esta parte presenta los hallazgos cuantitativos del estudio, organizados por "
                  "bloque temático. El eje de análisis es 2026; los datos 2025 se usan como "
                  "referencia comparativa en la Sección 4.11.", italic=True, size=11)
    add_para(doc, "Cada subsección incluye: (a) hallazgos clave, (b) tablas con datos completos, "
                  "(c) gráficos cuando aplique, (d) nota metodológica anexa.", italic=True, size=10, color=GRAY_MID)


def build_sec_4_1_cpc(doc):
    add_h2(doc, "4.1  NSE C+/C — el segmento natural de Gama")

    add_callout_box(doc,
        "El segmento NSE C+/C (n=104, 25.9% de la muestra) define el campo de batalla "
        "principal del estudio. Es donde Gama tiene su posición más fuerte (TOM 60.6%) y donde "
        "Rio y Páramo están creciendo agresivamente. La defensa de C+/C es la prioridad "
        "estratégica de Gama.")

    add_h3(doc, "4.1.1  Posición de Gama en C+/C 2026")

    headers = ["Métrica", "C+/C 2026", "Total 2026", "Diferencial pp", "Certeza"]
    rows = [
        ["TOM (Top of Mind)", "60.6%", "44.3%", "+16.3 pp", "✅ Alta"],
        ["Preferida", "13.5%", "8.0%", "+5.5 pp", "✅ Alta"],
        ["Asistida (consideración)", "~50%", "31.8%", "—", "Referencial"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 3, 3, 3, 2.5])

    add_chart(doc, os.path.join(CHARTS_V6, "C10_p33_percepcion_precio.png"), width_in=6.5,
              caption="Percepción de precios de Gama por NSE (P33). C+/C es el segmento más exigente con el precio.")

    add_para(doc, "Lectura clave:", bold=True)
    add_para(doc, "Gama tiene una sobrerrepresentación significativa en C+/C: 60.6% de TOM "
                  "vs 44.3% en el Total — un diferencial de 16.3 puntos porcentuales. En "
                  "preferencia el patrón se mantiene: 13.5% en C+/C vs 8.0% en el Total. "
                  "C+/C es donde el DNA de marca de Gama (experiencia premium, atención, "
                  "calidad) tiene mayor resonancia.")

    add_h3(doc, "4.1.2  Plan Suárez en C+/C — empate técnico recuperado")

    add_para(doc, "Un hallazgo importante de V3 que se había perdido en V5 y se recupera en V6: "
                  "en preferencia C+/C, Gama (13.5%) NO se separa estadísticamente de Páramo "
                  "(~16%) ni de Plan Suárez (~12%). Los intervalos de confianza se solapan.")

    add_chart(doc, os.path.join(CHARTS_V6, "C16_plan_suarez_cpc.png"), width_in=6.0,
              caption="Preferencia en C+/C: tres cadenas en empate técnico (IC95 solapados).")

    add_para(doc, "Esto significa que la batalla por el segmento natural de Gama está abierta. "
                  "Páramo, Gama y Plan Suárez compiten por el mismo shopper de C+/C, y ninguna "
                  "tiene una ventaja perceptual decisiva sobre las otras dos. Cualquier "
                  "movimiento competitivo de las tres puede alterar el equilibrio.")

    add_h3(doc, "4.1.3  La amenaza: Rio y Páramo crecen explosivamente en C+/C")

    headers = ["Cadena", "TOM C+/C 2025", "TOM C+/C 2026", "Delta", "Significancia"]
    rows = [
        ["Rio", "~30.9%", "56.7%", "+25.8 pp", "sig 99% (exploratorio)"],
        ["Páramo", "~18.0%", "40.4%", "+22.3 pp", "sig 99% (exploratorio)"],
        ["Gama", "~58.4%", "60.6%", "+2.2 pp", "no significativo"],
    ]
    add_table(doc, headers, rows, col_widths=[3, 3.5, 3.5, 3, 4])

    add_chart(doc, os.path.join(CHARTS_V5, "D2_tom_cpc_wow.png"), width_in=6.5,
              caption="TOM en C+/C 2025→2026 — la amenaza de Rio y Páramo en el segmento natural de Gama.")

    add_para(doc, "Si las tendencias son reales, el segmento natural de Gama está siendo "
                  "penetrado agresivamente por sus dos competidores principales a una velocidad "
                  "de +22-26 puntos porcentuales por ola. La posición actual de Gama es sólida; "
                  "la velocidad de Rio y Páramo es la señal de alerta más importante del "
                  "estudio — pendiente confirmación en ola 2027.")

    add_caveat_box(doc,
        "Los análisis WoW por NSE no fueron pre-registrados en el diseño 2025. Deben "
        "interpretarse como hipótesis a confirmar en ola 2027. La composición geográfica 2025 "
        "puede sesgar los deltas. Causalidad no inferible. (CV-WOW-005)")

    add_h3(doc, "4.1.4  El Núcleo Leal de Gama vive en C+/C")

    add_para(doc, "El Segmento 3 (Núcleo Leal, 8% de la muestra total) está sobrerrepresentado "
                  "en C+/C. El perfil dominante: frecuentes de 18-30 y 50+ con uso del horario "
                  "24h. El servicio de 24 horas es un activo especialmente diferenciador en "
                  "este segmento — Rio y Páramo no tienen un equivalente funcional documentado.")

    add_verbatim(doc,
        "El de 24 horas es un tiro al piso, de verdad que sí.",
        "Frecuentes 18-30 años, focus group cualitativo")
    add_verbatim(doc,
        "Eso es impelable.",
        "Ocasionales 50+ años, focus group cualitativo")

    add_metodologia_box(doc,
        "Datos cuantitativos por NSE basados en filtrado de la BBDD 2026 por variable NSE "
        "(C+/C combinado). Tests de proporciones con corrección Newcombe-Wilson para IC95. "
        "Comparaciones WoW (2025→2026) calculadas con corrección Benjamini-Hochberg FDR para "
        "controlar tasa de falsos descubrimientos en múltiples tests. Verbatims cualitativos "
        "del corpus de focus groups Gama (12 documentos, ~42.094 caracteres).")


def build_sec_4_2_salud_gama(doc):
    add_h2(doc, "4.2  Salud de marca Gama — embudo 2026")

    add_callout_box(doc,
        "El embudo de Gama es estadísticamente estable en 2026: ninguno de los 8 indicadores "
        "muestra variación significativa entre 2025 y 2026. En un mercado donde 9 de 10 cambios "
        "significativos son aumentos de competidores, la estabilidad de Gama es fortaleza "
        "defensiva — sus activos diferenciadores (24h, atención relacional) generan una base "
        "estable.")

    add_h3(doc, "4.2.1  Embudo 2026 vs referencia 2025")

    headers = ["Etapa del embudo", "Gama 2026", "Gama 2025 (ref.)", "Delta WoW", "Significancia"]
    rows = [
        ["TOM (Top of Mind)", "44.3%", "42.0%", "+2.2 pp", "no sig"],
        ["Asistida", "60.2%", "~58.0%", "+2.2 pp", "no sig"],
        ["Consideración", "31.8%", "27.5%", "+4.3 pp", "no sig"],
        ["Compra últimos 3m", "17.7%", "17.8%", "-0.2 pp", "no sig"],
        ["Preferida", "8.0%", "9.7%", "-1.7 pp", "no sig"],
        ["Última compra", "—", "—", "-1.2 pp", "no sig"],
        ["Habitual", "20.2%", "19.4%", "+0.8 pp", "no sig"],
        ["Misiones", "—", "—", "+4.9 pp", "no sig"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 3, 3, 2.5, 3])

    add_chart(doc, os.path.join(CHARTS_V5, "D3_funnel_gama.png"), width_in=6.5,
              caption="Embudo Gama 2026 vs referencia 2025 — 0/8 indicadores con variación significativa.")

    add_para(doc, "Hallazgos del embudo:", bold=True)
    add_bullet(doc, "TOM alto: 44.3%. Conocimiento espontáneo robusto. Gama no tiene problema de awareness.")
    add_bullet(doc, "Conversión TOM→Preferida muy baja: 44.3% → 8.0%. De cada 100 personas que conocen "
                    "Gama espontáneamente, solo 8 la prefieren. Aquí está el problema central.")
    add_bullet(doc, "Compra 3m (17.7%) es más del doble que Preferida (8.0%). Hay shoppers que compran en "
                    "Gama sin preferirla — uso transaccional o de conveniencia.")
    add_bullet(doc, "El embudo es estable año contra año. No hay erosión ni crecimiento — Gama mantiene "
                    "su posición en un mercado dinámico.")

    add_metodologia_box(doc,
        "Embudo construido sobre n=402 (2026) y n=785 (2025). Tests WoW con Benjamini-Hochberg FDR "
        "para corrección por comparaciones múltiples. Hipótesis nula \"sin cambio\" no se rechaza en "
        "ninguno de los 8 indicadores (p_adj > 0.10 en todos los casos). Caveats CV-WOW-001 (sin "
        "ponderación 2025) y CV-WOW-002 (composición geográfica difiere) aplican.")


def build_sec_4_3_posicion(doc):
    add_h2(doc, "4.3  Posición competitiva 2026")

    add_callout_box(doc,
        "Rio superó a Gama en TOM por primera vez en este tracker: 45.0% vs 44.3%. Páramo "
        "creció +12pp en TOM. Central Madeirense pierde compra (-7.7pp, sig 95%). Gama "
        "permanece estable — fortaleza relativa en un mercado muy dinámico.")

    headers = ["Cadena", "TOM 2026", "Preferida 2026", "Movimiento WoW", "Diagnóstico"]
    rows = [
        ["Rio", "45.0%", "10.2%", "+17.0 pp TOM sig 99%", "Ganador absoluto del mercado"],
        ["Gama", "44.3%", "8.0%", "0 pp (estable)", "Defensa sólida en mercado turbulento"],
        ["Páramo", "~39.1%", "~6-8%", "+12.1 pp TOM sig 99%", "Consolidación en cadena del embudo"],
        ["Central Madeirense", "—", "11.2%", "-7.7 pp Compra sig 95%", "Pérdida significativa"],
        ["Plan Suárez", "—", "—", "—", "Competidor relevante en C+/C"],
    ]
    add_table(doc, headers, rows, col_widths=[3.5, 2.5, 3, 4, 4.5])

    add_chart(doc, os.path.join(CHARTS_V5, "D4_posicionamiento.png"), width_in=6.5,
              caption="Mapa de posicionamiento competitivo TOM vs Preferida 2026 con cambio WoW.")

    add_para(doc, "Lectura competitiva:", bold=True)
    add_bullet(doc, "Gama mantiene la mayor Preferida del top 3 (8.0% vs 10.2% Rio) — su core fidelizado "
                    "es más leal que el de Rio.")
    add_bullet(doc, "Central Madeirense tiene la Preferida más alta (11.2%) pero pierde Compra activa — "
                    "shoppers que la prefieren pero compran en otra cadena.")
    add_bullet(doc, "La oportunidad de capturar el espacio que CM está perdiendo está abierta para Gama "
                    "y Rio.")

    add_metodologia_box(doc,
        "Análisis WoW con tests de proporciones (z-test Newcombe-Wilson) + corrección BH-FDR. "
        "Solo 10 de 57 ítems WoW pasan el filtro de significancia BH-FDR (q<0.05). 9 son aumentos "
        "de competidores (Rio y Páramo), 1 es pérdida de CM, 0 son cambios de Gama. Caveats WoW "
        "aplicables.")


def build_sec_4_4_drivers(doc):
    add_h2(doc, "4.4  Drivers de preferencia")

    add_callout_box(doc,
        "Atención al cliente es el driver #1 de la preferencia por Gama (OR=5.73, p<0.01, con "
        "convergencia de 4 métodos independientes). Quien asocia Gama con atención tiene 5.7 "
        "veces más probabilidades de preferirla. Precio es el predictor #10 (OR=1.03, no "
        "significativo).")

    add_h3(doc, "4.4.1  El modelo logístico de preferencia")

    headers = ["Driver", "OR logit", "IC95", "Significancia", "SHAP rank", "Convergencia"]
    rows = [
        ["Atención al cliente", "5.73", "[1.6, 20.4]", "*** p<0.01", "#1",
         "Logit + RF + SHAP + razón espontánea"],
        ["Limpieza", "3.99", "[0.94, 16.91]", "* p<0.10", "#2", "Logit borderline + SHAP + Gini"],
        ["Promociones", "3.64", "[1.1, 11.8]", "** p<0.05", "#3", "Logit + SHAP + razón espontánea"],
        ["Precio (menor)", "1.03", "—", "NS (p=0.966)", "#10", "4 métodos: NO es driver"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 2, 3, 2.5, 1.8, 5])

    add_chart(doc, os.path.join(CHARTS_V5, "D5_forest_drivers.png"), width_in=6.5,
              caption="Forest plot de drivers de preferencia con IC95 — Atención es el único driver con OR>5.")

    add_para(doc, "Tres conclusiones de alta certeza:", bold=True)
    add_bullet(doc, "Conclusión 1: Atención es el driver más fuerte. Quien asocia Gama con atención "
                    "tiene 5.7 veces más odds de preferirla. La convergencia de 4 métodos "
                    "independientes (logit, Random Forest, SHAP, razón espontánea) es el mayor "
                    "respaldo metodológico del estudio.")
    add_bullet(doc, "Conclusión 2: El precio NO predice preferencia. Comunicar precio en Gama es "
                    "invertir en el atributo menos relevante para la conversión. (Esta es la "
                    "conclusión que se desarrolla profundamente en la Sección 4.5.)")
    add_bullet(doc, "Conclusión 3: Las promociones son palanca táctica (OR=3.64**). Su función es "
                    "atraer a la primera visita, no comunicar precio bajo.")

    add_h3(doc, "4.4.2  Segmentación K-means: tres tipos de shoppers")

    headers = ["Segmento", "Tamaño", "Perfil cuantitativo", "Oportunidad"]
    rows = [
        ["Seg 1 — Mayoría Exigente", "59% (n≈237)", "Alta exigencia, no exclusivos Gama, misión múltiple",
         "Defensa + comunicación de experiencia"],
        ["Seg 2 — Pragmáticos Convertibles", "33% (n≈133)",
         "Resistencia precio 3.44/5 (vs 3.66 Seg1), 0% pref-Gama actual",
         "Mayor retorno C/P — target de conversión prioritario"],
        ["Seg 3 — Núcleo Leal", "8% (n≈32)",
         "Alta frecuencia, uso 24h, lealtad alta, vínculo afectivo",
         "Retención + advocacy"],
    ]
    add_table(doc, headers, rows, col_widths=[4.5, 2.5, 5, 5])

    add_chart(doc, os.path.join(CHARTS_V5, "D6_kmeans_scatter.png"), width_in=6.0,
              caption="K-means k=3 con centroides — los 3 perfiles de shopper en el mercado.")

    add_para(doc, "El Segmento 2 es el de mayor potencial de conversión a corto plazo: 0% de "
                  "preferencia actual + menor resistencia al precio + n=133 (33% del mercado). "
                  "El cualitativo sugiere que su barrera no es económica sino identitaria — "
                  "se trata en profundidad en la Sección 5.3.")

    add_caveat_box(doc,
        "Silhouette score ~0.20 (moderado). Los segmentos son tendencias interpretativas, no "
        "categorías discretas. Los perfiles se solapan en los márgenes. El tamaño del Segmento 3 "
        "(n=32) es referencial.")

    add_metodologia_box(doc,
        "Regresión logística (statsmodels) sobre Preferida=Gama (1/0) con 10 atributos de P22 "
        "como predictores. Pseudo R²=0.4371, AUC=0.929. Random Forest (sklearn) con SHAP para "
        "validar contribución no-lineal de cada feature. K-means k=3 (sklearn), silhouette "
        "validado contra k=2..6 (máximo en k=3) y BayesianGaussianMixture (BIC mínimo en k=3). "
        "n pref-Gama=32 (referencial — IC95 amplios en OR pero patrón robusto por convergencia "
        "de 4 métodos).")


# === SECCIÓN CLAVE: PRECIOS ===
def build_sec_4_5_precios(doc):
    add_h2(doc, "4.5  Análisis profundo de PRECIOS — capítulo nuevo de V6")

    add_callout_box(doc,
        "Este capítulo es nuevo en V6. Reemplaza la cobertura fragmentaria del tema precios en "
        "V5 (que lo dispersaba en 4 lugares distintos del documento). Integra las 4 preguntas "
        "del cuestionario sobre precio (P31, P32, P33, P34) y agrega el cruce nuevo P33×P21 "
        "que responde una pregunta estratégica clave: ¿cómo se relaciona la percepción de "
        "precio de Gama con la preferencia?")

    add_h3(doc, "4.5.1  Percepción de precios de Gama vs competencia (P33)")

    add_para(doc, "P33 mide la percepción agregada de precio de Gama vs otras cadenas. "
                  "Resultado central: 54% de los shoppers percibe a Gama como más cara. La "
                  "brecha entre NETO caro (54%) y NETO económico (15.4%) es de 38.6 puntos "
                  "porcentuales.")

    headers = ["Categoría", "n", "%", "IC95"]
    rows = [
        ["Mucho más económicos", "6", "1.5%", "[0.7, 3.2]"],
        ["Poco más económicos", "56", "13.9%", "[10.9, 17.7]"],
        ["Igual que otros", "123", "30.6%", "[26.3, 35.3]"],
        ["Poco más caros", "152", "37.8%", "[33.2, 42.7]"],
        ["Mucho más caros", "65", "16.2%", "[12.9, 20.1]"],
        ["NETO caro (Poco + Mucho)", "217", "54.0%", "[49.1, 58.8]"],
        ["NETO económico (Poco + Mucho)", "62", "15.4%", "[12.1, 19.4]"],
    ]
    add_table(doc, headers, rows, col_widths=[7, 2, 2.5, 4])

    add_chart(doc, os.path.join(CHARTS_V6, "C10_p33_percepcion_precio.png"), width_in=6.5,
              caption="Percepción de precios de Gama por NSE — el problema es transversal a todos los segmentos.")

    add_para(doc, "Por NSE — un problema transversal:", bold=True)
    headers = ["NSE", "n", "NETO caro %", "IC95"]
    rows = [
        ["C+/C", "104", "60.6%", "[51.0, 69.4]"],
        ["D", "127", "49.6%", "[41.1, 58.2]"],
        ["E", "171", "53.2%", "[45.7, 60.5]"],
    ]
    add_table(doc, headers, rows, col_widths=[3, 2, 3, 4])

    add_para(doc, "Hallazgo clave de P33:", bold=True)
    add_para(doc, "La percepción de precio caro es transversal a todos los NSE. No es específica "
                  "del segmento premium ni del popular. C+/C la percibe marginalmente más cara "
                  "(60.6%) pero la diferencia con D y E no es estadísticamente significativa al 95%. "
                  "Implicación: el problema de imagen de precio no se resuelve solo en el segmento "
                  "premium — es un problema de marca general.")

    add_h3(doc, "4.5.2  Entre los propios preferentes de Gama — el activo latente")

    headers = ["Segmento", "n", "NETO caro %", "Lectura"]
    rows = [
        ["Preferentes de Gama", "32 (referencial)", "34.4%", "Más bajo que el promedio del mercado"],
        ["No preferentes de Gama", "~370", "~57% (estimado)", "Más alto que el promedio"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 5])

    add_para(doc, "Los propios preferentes de Gama reconocen que no es la opción más económica "
                  "(34.4% NETO caro) — pero la diferencia con el resto del mercado (~57%) "
                  "muestra que conocen la marca diferente. Su lealtad se construye sobre otros "
                  "atributos (atención, limpieza, 24h), no sobre precio. Esto es una fortaleza "
                  "(la lealtad no depende del precio) pero también una vulnerabilidad (si el "
                  "precio sube sin compensación en servicio, el switching es probable).")

    add_h3(doc, "4.5.3  Evolución percibida de precios en 6 meses (P34)")

    headers = ["Categoría", "n", "%"]
    rows = [
        ["Mucho más económicos", "10", "2.5%"],
        ["Poco más económicos", "32", "8.0%"],
        ["Igual que hace 6 meses", "179", "44.5%"],
        ["Poco más caros", "138", "34.3%"],
        ["Mucho más caros", "43", "10.7%"],
        ["NETO caros en 6m", "181", "45.0%"],
        ["NETO económicos en 6m", "42", "10.5%"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 2, 2])

    add_para(doc, "El 45% percibe que Gama subió precios en el último semestre. Solo 10.5% "
                  "percibe bajada. La combinación de P33 y P34 es relevante: Gama no solo es "
                  "percibida como cara hoy (54%), sino que la mayoría percibe que lo es cada "
                  "vez más (45% ve subida en 6 meses). Esto crea un ciclo de percepción de "
                  "precio que puede acelerar la erosión en shoppers sensibles al precio.")

    add_h3(doc, "4.5.4  Ranking de precio entre cadenas (P31)")

    add_chart(doc, os.path.join(CHARTS_V6, "C11_p31_ranking_precio.png"), width_in=6.8,
              caption="Ranking de precio percibido — Gama es la 6ta cadena más económica del mercado.")

    headers = ["Cadena", "Mean rank", "Mediana", "% Top-3 precio bajo"]
    rows = [
        ["Páramo", "4.24", "4", "47.8%"],
        ["Central Madeirense", "4.06", "4", "45.0%"],
        ["Forum", "4.53", "4", "40.0%"],
        ["Rio", "5.01", "5", "32.6%"],
        ["Plan Suárez", "5.25", "5", "29.1%"],
        ["Gama", "5.77", "6", "21.9%"],
        ["Luz", "6.30", "7", "21.1%"],
        ["La Muralla", "8.49", "9", "4.7%"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 2.5, 4])

    add_para(doc, "Gama ocupa la posición #6 en ranking forzado de precio. Solo 21.9% la ubica "
                  "entre las 3 cadenas más económicas. Páramo es el líder claro (4.24 mean rank, "
                  "47.8% Top-3).")

    add_h3(doc, "4.5.5  P32 — Mejor precio por categoría (15 categorías)")

    add_para(doc, "P32 es la pregunta clave para responder la consulta estratégica de Cora: "
                  "¿en qué categorías Gama es percibida como la más económica? La respuesta "
                  "es contundente: en ninguna categoría de alta rotación.")

    add_chart(doc, os.path.join(CHARTS_V6, "C12_p32_categorias.png"), width_in=7.0,
              caption="Mejor precio percibido por categoría: Gama vs líder de cada categoría. Páramo domina 11/15.")

    headers = ["Categoría", "Líder precio", "Líder %", "Gama %", "Pos.", "Gap pp"]
    rows = [
        ["Farmacia", "Plan Suárez", "3.7%", "3.0%", "#2", "-0.7"],
        ["Licores", "Forum", "4.7%", "3.5%", "#4", "-1.2"],
        ["Galletas y confitería", "Páramo", "11.4%", "5.7%", "#8", "-5.7"],
        ["Congelados", "Páramo", "11.4%", "4.7%", "#6", "-6.7"],
        ["Salsas y Enlatados", "Páramo", "20.9%", "4.5%", "#9", "-16.4"],
        ["Productos básicos", "Páramo", "18.4%", "4.5%", "#8", "-13.9"],
        ["Cuidado y limpieza", "Forum", "9.7%", "3.7%", "#8", "-6.0"],
        ["Gaseosas, jugos", "Páramo", "10.2%", "3.2%", "#8", "-7.0"],
        ["Cuidado Personal", "Páramo", "10.0%", "3.0%", "#8", "-7.0"],
        ["Carne de res", "Páramo", "36.3%", "3.0%", "#7", "-33.3"],
        ["Pollo", "Páramo", "35.3%", "2.5%", "#9", "-32.8"],
        ["Charcutería", "Páramo", "34.8%", "1.7%", "#9", "-33.1"],
        ["Frutas, legumbres", "Central M.", "13.2%", "2.2%", "#9", "-11.0"],
        ["Alimento mascotas", "Páramo", "4.0%", "1.2%", "#9", "-2.8"],
        ["Pescados y mariscos", "Páramo", "6.0%", "1.0%", "#9", "-5.0"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 2, 2, 1.5, 1.8])

    add_para(doc, "Patrones de P32:", bold=True)
    add_bullet(doc, "Páramo es el líder de precio percibido en 11 de 15 categorías. Su ventaja "
                    "es sistemática, no sectorial.")
    add_bullet(doc, "Gama no lidera precio en ninguna categoría de alta rotación. Sus mejores "
                    "posiciones son Farmacia (#2) y Licores (#4) — ambas con bases de percepción "
                    "muy pequeñas (85% y 66% \"Ninguno\").")
    add_bullet(doc, "Las proteínas (carne, pollo, charcutería) tienen los gaps más grandes — "
                    "hasta 33 puntos porcentuales vs Páramo. Estas son las categorías de mayor "
                    "volumen y mayor penalización a la imagen de precio total.")
    add_bullet(doc, "En NSE C+/C, los gaps en proteínas son de 29-30 puntos — el segmento natural "
                    "de Gama percibe claramente que Páramo es mejor en precio en proteínas.")

    add_h3(doc, "4.5.6  Cruce nuevo P33 × P21 — la relación entre percepción de precio y preferencia")

    add_para(doc, "Este es uno de los análisis nuevos de V6 que responde directamente la pregunta "
                  "de Cora: \"¿cómo se relaciona la percepción de precio de Gama con la preferencia?\". "
                  "El resultado: hay una relación dosis-respuesta clara y estadísticamente significativa.")

    add_chart(doc, os.path.join(CHARTS_V6, "C13_p33xp21_gradiente.png"), width_in=6.5,
              caption="Gradiente monótono entre percepción de precio de Gama y preferencia.")

    headers = ["Percepción P33", "n en categoría", "n prefiere Gama", "% prefiere Gama", "IC95"]
    rows = [
        ["Mucho más caros", "65", "0", "0.0%", "[0.0, 5.6]"],
        ["Poco más caros", "152", "11", "7.2%", "[4.1, 12.5]"],
        ["Igual precio", "123", "11", "8.9%", "[5.1, 15.3]"],
        ["Poco más económico", "56", "7", "12.5%", "[6.2, 23.6]"],
        ["Mucho más económico (n=6, EXCLUIR)", "6", "3", "50.0%", "—"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 2.5, 2.5, 3, 3])

    add_para(doc, "Análisis estadístico:", bold=True)
    headers = ["Estadístico", "Valor", "Interpretación"]
    rows = [
        ["Chi-cuadrado (4 gl)", "21.94", "Significativo p<0.001"],
        ["Cramér's V", "0.234", "Asociación pequeña-moderada"],
        ["OR (favorable vs cara)", "2.40", "IC95 [1.12, 5.12]"],
        ["OR atención (comparativo CU-5)", "5.73", "Driver más potente — Atención > Precio"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 3, 6])

    add_callout_box(doc,
        "Pasar de \"Gama es más cara\" a \"Gama es igual o más económica\" más que duplica la "
        "probabilidad de preferir Gama (OR=2.4, significativo). Pero el OR de atención (5.73) "
        "sigue siendo 2.4 veces más potente. La batalla de precio importa, pero la de atención "
        "importa más.", color=BLUE_INFO)

    add_para(doc, "Esto resuelve la paradoja aparente: el atributo precio no predice preferencia "
                  "(OR=1.03 NS) porque entre quienes ya conocen y experimentaron Gama, el precio "
                  "no es el criterio de elección. Pero la percepción agregada de precio (P33) sí "
                  "predice preferencia porque actúa como barrera de entrada — quien percibe a "
                  "Gama como cara nunca llega a experimentar el activo de atención.")

    add_metodologia_box(doc,
        "P31 ranking forzado 1-10, mean rank y % Top-3 calculados sobre n=402. "
        "P32 multi-mention por categoría, % sobre n=402 que asigna cadena (excluyendo \"Ninguno\"). "
        "P33 5-puntos Likert, NETO caro = Poco + Mucho más caros. "
        "P34 5-puntos Likert evolución 6 meses. "
        "Cruce P33×P21: tabla cruzada con chi-cuadrado de independencia y Cramér's V. "
        "OR favorable calculado agrupando Igual+Poco+Mucho económico (n=185) vs Poco+Mucho caro (n=217). "
        "Bases bajas en NEW-1 son referenciales (CV-11, CV-12, CV-14 del CU-8 v2).")


def build_sec_4_6_lugar_precio_preferencia(doc):
    add_h2(doc, "4.6  Lugar de compra × precio percibido × preferencia × misiones")

    add_callout_box(doc,
        "Capítulo nuevo de V6 que responde literalmente a las preguntas estratégicas de Cora "
        "sobre cómo se relacionan el lugar de compra, la percepción de precio por categoría, "
        "la preferencia y las misiones de compra. El hallazgo clave: existe un \"territorio "
        "protegido\" de Gama en 3 categorías donde el shopper compra habitualmente sin "
        "percibirla como la más económica — evidencia directa de que atención/cercanía/24h "
        "generan retención incluso sin liderazgo de precio.")

    add_h3(doc, "4.6.1  Matriz P30 × P32 — congruencia entre hábito y precio")

    add_para(doc, "En 11 de 15 categorías el líder de hábito (P30) coincide con el líder de "
                  "precio percibido (P32). Esto significa que en la mayor parte del mercado, el "
                  "precio sí es el criterio principal de elección de cadena. Las 4 excepciones "
                  "son ricas y se analizan a continuación.")

    headers = ["Categoría", "Líder hábito P30", "Líder precio P32", "Congruencia"]
    rows = [
        ["Galletas y confitería", "Central Madeirense", "Páramo", "NO — CM tiene mejor surtido percibido"],
        ["Licores", "Luz", "Forum", "NO — Luz por cercanía/surtido, no por precio"],
        ["Farmacia", "Gama (líder)", "Plan Suárez", "NO — Gama por conveniencia, no por precio"],
        ["Cuidado y limpieza", "Páramo", "Forum", "NO — Páramo retiene hábito pese a Forum mejor en precio"],
        ["Las otras 11 categorías", "—", "—", "SÍ — coincidencia hábito-precio"],
    ]
    add_table(doc, headers, rows, col_widths=[4.5, 3.5, 3.5, 5])

    add_h3(doc, "4.6.2  El territorio protegido de Gama — hábito sin precio")

    add_para(doc, "Gama tiene shoppers habituales en varias categorías sin figurar entre los top "
                  "de mejor precio percibido. En esas categorías, el shopper elige Gama por otro "
                  "atributo: conveniencia, cercanía, atención, horario 24h.")

    headers = ["Categoría", "Gama hábito %", "Gama precio %", "Brecha (hábito - precio)", "Líder precio"]
    rows = [
        ["Gaseosas, jugos y aguas", "7.2%", "3.2%", "+4.0 pp", "Páramo"],
        ["Congelados", "8.0%", "4.7%", "+3.3 pp", "Páramo"],
        ["Salsas y Enlatados", "7.5%", "4.5%", "+3.0 pp", "Páramo"],
        ["Galletas y confitería", "7.0%", "5.7%", "+1.3 pp", "Páramo"],
        ["Cuidado y limpieza", "5.0%", "3.7%", "+1.3 pp", "Forum"],
        ["Productos básicos", "5.5%", "4.5%", "+1.0 pp", "Páramo"],
        ["Cuidado Personal", "4.0%", "3.0%", "+1.0 pp", "Páramo"],
        ["Farmacia (Gama líder hábito)", "4.2%", "3.0%", "+1.2 pp", "Plan Suárez"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 2.5, 2.5, 3.5, 3])

    add_para(doc, "Las 3 categorías de mayor desconexión Gama (Gaseosas +4.0pp, Congelados +3.3pp, "
                  "Salsas y Enlatados +3.0pp) son evidencia empírica de que el activo de marca de "
                  "Gama (atención, conveniencia, 24h) retiene cliente incluso sin liderazgo de "
                  "precio. Es el territorio protegido.")

    add_h3(doc, "4.6.3  Desconexión × Preferencia × Misión")

    headers = ["Categoría", "Gama hábito %", "% pref. Gama compra acá", "Misión principal", "Misión %"]
    rows = [
        ["Gaseosas, jugos", "7.2%", "59.4%", "Reabastecimiento parcial", "82.8%"],
        ["Salsas y Enlatados", "7.5%", "56.2%", "Reabastecimiento parcial", "73.3%"],
        ["Congelados", "8.0%", "37.5%", "Reabastecimiento parcial", "78.1%"],
        ["Frutas, legumbres", "4.0%", "34.4%", "Reabastecimiento parcial", "62.5%"],
    ]
    add_table(doc, headers, rows, col_widths=[4.5, 2.5, 3, 4, 2.5])

    add_callout_box(doc,
        "Patrón triple identificado: las categorías de mayor desconexión hábito-precio (el "
        "shopper va a Gama sin percibir mejor precio) son las mismas que (a) concentran mayor "
        "proporción de los preferentes de Gama, y (b) tienen la misión de Reabastecimiento "
        "Parcial como la principal. Es decir: el shopper leal de Gama va a completar el mercado, "
        "compra Gaseosas/Enlatados/Congelados habitualmente ahí, y no necesita que Gama sea la "
        "más barata para hacerlo.", color=BLUE_INFO)

    add_para(doc, "Implicación estratégica:", bold=True)
    add_para(doc, "Este patrón define el \"territorio seguro\" de Gama — categorías donde la "
                  "propuesta de valor es la experiencia (atención, conveniencia, 24h), no el "
                  "precio. La comunicación de precio en estas categorías podría convertir el "
                  "hábito de conveniencia en convicción de precio — sin necesidad de ser el "
                  "más barato, solo dejar de ser \"el percibido como caro\".")

    add_metodologia_box(doc,
        "Matriz P30 (supermercado habitual por categoría) × P32 (mejor precio percibido por "
        "categoría) calculada sobre n=402. Líderes por categoría identificados sobre la "
        "intersección de respondientes que asignan cadena en cada pregunta. Brecha calculada "
        "como (Gama hábito % - Gama precio %). Todos los valores de Gama por categoría son "
        "referenciales (n por categoría < 30) — los patrones son direccionales, no concluyentes "
        "individualmente, pero el patrón conjunto sí es robusto (CV-14 del CU-8 v2).")


def build_sec_4_7_balance(doc):
    add_h2(doc, "4.7  Balance óptimo de categorías a igualar precio")

    add_callout_box(doc,
        "Este análisis responde la pregunta estratégica de Cora: \"¿cuál es el balance óptimo "
        "de categorías a precios iguales a los competidores para capitalizar el segmento que se "
        "mueve por precio, sin perder su segmento natural de calidad, limpieza y los otros "
        "diferenciadores?\". La respuesta es una clasificación de las 15 categorías en 3 grupos "
        "estratégicos.")

    add_chart(doc, os.path.join(CHARTS_V6, "C18_balance_categorias.png"), width_in=6.5,
              caption="Balance óptimo de precios: 3 grupos estratégicos según el gap vs líder de cada categoría.")

    add_h3(doc, "4.7.1  Grupo 1 — CRÍTICO: igualar precio (gap >10pp)")

    add_para(doc, "Categorías donde Páramo lidera precio con brechas grandes. Igualar precio "
                  "aquí es la acción de mayor impacto visible para capturar al segmento "
                  "precio-sensible, especialmente en NSE C+/C donde las brechas son aún mayores.")
    headers = ["Categoría", "Líder", "Líder %", "Gama %", "Gap pp", "Recomendación"]
    rows = [
        ["Carne de res", "Páramo", "36.3%", "3.0%", "-33.3", "IGUALAR PRECIO"],
        ["Charcutería", "Páramo", "34.8%", "1.7%", "-33.1", "IGUALAR PRECIO"],
        ["Pollo", "Páramo", "35.3%", "2.5%", "-32.8", "IGUALAR PRECIO"],
        ["Salsas y Enlatados", "Páramo", "20.9%", "4.5%", "-16.4", "IGUALAR PRECIO"],
        ["Productos básicos", "Páramo", "18.4%", "4.5%", "-13.9", "IGUALAR PRECIO"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 2.5, 2, 2, 2, 3.5])

    add_para(doc, "Las proteínas (carne/pollo/charcutería) tienen gaps de 33pp — abismales. "
                  "Igualar precio aquí es la acción de mayor visibilidad para el segmento "
                  "precio-sensible. El costo de margen sería alto pero el impacto en imagen "
                  "de precio sería el mayor posible. Salsas y Enlatados + Productos básicos "
                  "tienen gaps de 14-16pp — también relevantes y probablemente con mejor ROI "
                  "de margen.")

    add_h3(doc, "4.7.2  Grupo 2 — EASY WIN: ajuste marginal (gap <2pp, vacante de percepción)")

    headers = ["Categoría", "Líder", "Líder %", "Gama %", "Gap pp", "% sin percepción", "Recomendación"]
    rows = [
        ["Farmacia", "Plan Suárez", "3.7%", "3.0%", "-0.7", "84.6%", "EASY WIN — comunicar precio"],
        ["Licores", "Forum", "4.7%", "3.5%", "-1.2", "66.4%", "EASY WIN — comunicar precio"],
    ]
    add_table(doc, headers, rows, col_widths=[3, 2.5, 2, 2, 1.5, 2.5, 3.5])

    add_para(doc, "El gap es mínimo (-0.7 y -1.2pp) pero el mercado de percepción es muy pequeño "
                  "— 84% y 66% de respondientes no asigna a nadie como el más económico en estas "
                  "categorías. Esto significa que no hay una batalla de precio activa que ganar "
                  "— hay una vacante de percepción que llenar. Comunicar precio en Farmacia "
                  "puede hacer que Gama sea el primer nombre que viene a la mente, sin necesidad "
                  "de recortar precio realmente.")

    add_h3(doc, "4.7.3  Grupo 3 — DIFERENCIAL DE EXPERIENCIA: mantener precio + comunicar atributo")

    headers = ["Categoría", "Gama hábito %", "Gama precio %", "Gap precio pp", "Líder", "Recomendación"]
    rows = [
        ["Congelados", "8.0%", "4.7%", "-6.7", "Páramo", "MANTENER + COMUNICAR DIFERENCIAL"],
        ["Salsas y Enlatados*", "7.5%", "4.5%", "-16.4", "Páramo", "MANTENER + COMUNICAR (ver nota)"],
        ["Gaseosas, jugos", "7.2%", "3.2%", "-7.0", "Páramo", "MANTENER + COMUNICAR DIFERENCIAL"],
        ["Galletas, confitería", "7.0%", "5.7%", "-5.7", "Páramo", "MANTENER + COMUNICAR DIFERENCIAL"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 2.5, 2.5, 2.5, 2.5, 3.5])

    add_para(doc, "(*) Salsas y Enlatados aparece en dos grupos. Tiene gap grande (Grupo 1 dice "
                  "\"igualar\") pero también tiene hábito sin precio (Grupo 3 dice \"mantener\"). "
                  "La tensión es real: es la categoría donde el shopper leal sigue comprando sin "
                  "importar precio — pero el gap es tan grande que igualar precio podría expandir "
                  "la base. Decisión de negocio Gama.", italic=True, size=10)

    add_h3(doc, "4.7.4  Recomendación de balance estratégico")

    headers = ["Prioridad", "Acción", "Categorías", "Justificación"]
    rows = [
        ["1", "Igualar precio", "Proteínas (carne/pollo/charcutería)",
         "Mayor impacto visible, captura segmento precio. Alto costo margen."],
        ["2", "Comunicar precio (easy win)", "Farmacia, Licores",
         "Gap mínimo, vacante de percepción disponible. Bajo costo."],
        ["3", "Mantener precio + comunicar diferencial", "Congelados, Gaseosas, Galletas",
         "Shopper leal ya está — no sacrificar margen, comunicar por qué vale Gama."],
        ["4", "Monitorear", "Productos básicos, Frutas",
         "Brecha media, base habitual moderada. Evaluar ROI según recursos."],
    ]
    add_table(doc, headers, rows, col_widths=[2, 3.5, 4.5, 5])

    add_caveat_box(doc,
        "Este análisis es 100% perceptual (P30/P32 de encuesta). Las cifras de Gama en precio "
        "percibido son REFERENCIALES (n<30 en todos los casos). Las decisiones de ajuste de "
        "precio deben validarse con datos reales de pricing interno antes de ejecutar. Un "
        "shopper que percibe precio alto puede estar equivocado — o puede estar en lo correcto. "
        "La encuesta no distingue. (CV-15 del CU-8 v2)")

    add_metodologia_box(doc,
        "Clasificación basada en gap entre líder de precio percibido (P32) y posición de Gama. "
        "Umbrales: >10pp = crítico, <2pp = easy win con vacante, intermedios = mantener si hay "
        "hábito existente. La columna \"% sin percepción\" refleja respondientes que eligen "
        "\"Ninguno en particular\" en P32. Recomendaciones son evaluación analítica — requieren "
        "validación con pricing operativo interno de Gama antes de ejecutar.")


def build_sec_4_8_subset(doc):
    add_h2(doc, "4.8  Subset de shoppers con percepción favorable")

    add_callout_box(doc,
        "Capítulo nuevo de V6 que responde a la pregunta de Cora: \"¿entre los que consideran "
        "que Gama tiene el precio igual o menor que otros, cómo es el mix de categorías que "
        "compran en Gama y la percepción de marca más económica?\". El subset es n=185 (46% "
        "del mercado) — un grupo grande con percepción favorable que merece atención específica.")

    add_h3(doc, "4.8.1  Composición del subset")

    headers = ["Percepción P33", "n", "%"]
    rows = [
        ["Mucho más caros", "65", "16.2%"],
        ["Poco más caros", "152", "37.8%"],
        ["Igual que en otros lugares", "123", "30.6%"],
        ["Poco más económicos", "56", "13.9%"],
        ["Mucho más económicos", "6", "1.5%"],
        ["Subset \"Igual o menos caro\" (IGUAL + ECONÓMICO)", "185", "46.0%"],
        ["Subset \"Caros\" (Poco + Mucho)", "217", "54.0%"],
    ]
    add_table(doc, headers, rows, col_widths=[7, 2.5, 2.5])

    add_para(doc, "El subset \"Igual o menos caro\" está dominado por quienes dicen \"Igual\" "
                  "(66.5% del subset). La percepción favorable es mayoritariamente de PARIDAD, "
                  "no de ventaja. Solo 33.5% del subset ve a Gama genuinamente como más "
                  "económica.")

    add_h3(doc, "4.8.2  Mix de categorías Gama — subset vs Total")

    add_chart(doc, os.path.join(CHARTS_V6, "C17_subset_categorias.png"), width_in=6.8,
              caption="Hábito de compra en Gama por categoría: subset favorable vs Total. El subset compra más en Gama en casi todas las categorías.")

    headers = ["Categoría", "Total %", "Subset %", "Diferencia pp"]
    rows = [
        ["Gaseosas, jugos y aguas", "7.2%", "11.9%", "+4.7"],
        ["Productos básicos", "5.5%", "8.6%", "+3.1"],
        ["Salsas y Enlatados", "7.5%", "10.3%", "+2.8"],
        ["Galletas y confitería", "7.0%", "9.7%", "+2.7"],
        ["Cuidado Personal", "4.0%", "6.5%", "+2.5"],
        ["Frutas, legumbres", "4.0%", "5.9%", "+1.9"],
        ["Carne de res", "3.7%", "5.4%", "+1.7"],
        ["Congelados", "8.0%", "9.7%", "+1.7"],
        ["Cuidado y limpieza hogar", "5.0%", "6.5%", "+1.5"],
        ["Farmacia", "4.2%", "4.3%", "+0.1 (excepción)"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 2.5, 2.5, 3])

    add_para(doc, "El subset compra MÁS en Gama en todas las categorías comparables. La "
                  "excepción es Farmacia — donde el hábito de compra es prácticamente igual "
                  "entre el subset favorable y el Total. Esto confirma que el hábito en "
                  "Farmacia se explica por conveniencia, no por percepción de precio: el "
                  "shopper va a Gama por farmacia porque está de paso, no porque la considere "
                  "más económica.", italic=False, size=10)

    add_h3(doc, "4.8.3  Cruce con preferencia y misiones")

    headers = ["Indicador", "Total (n=402)", "Subset (n=185)", "Diferencia"]
    rows = [
        ["% que prefiere Gama (P21)", "8.0% (n=32)", "11.4% (n=21)", "+3.4 pp"],
        ["Misión: Abastecimiento general", "23.4%", "29.2%", "+5.8 pp"],
        ["Misión: Reabastecimiento parcial", "67.2%", "60.5%", "-6.7 pp"],
        ["Misión: Urgencia", "7.7%", "8.6%", "+0.9 pp"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 3, 3, 2.5])

    add_para(doc, "Hallazgo importante de misiones:", bold=True)
    add_para(doc, "El subset favorable tiene mayor proporción de misiones de Abastecimiento "
                  "General (+5.8pp) y menor proporción de Reabastecimiento Parcial (-6.7pp) "
                  "que el Total. Esto es relevante: quienes perciben a Gama como competitiva "
                  "en precio tienden a ir a hacer el mercado COMPLETO — no solo a completar. "
                  "Es decir, la percepción de precio favorable habilita la misión de \"mercado "
                  "grande\", no solo la de reposición.")

    add_h3(doc, "4.8.4  Perfil demográfico del subset")

    headers = ["NSE", "n en Total", "n en Subset", "% del NSE que percibe Gama favorable"]
    rows = [
        ["C+/C", "104", "41", "39.4%"],
        ["D", "127", "64", "50.4%"],
        ["E", "171", "80", "46.8%"],
    ]
    add_table(doc, headers, rows, col_widths=[3, 3, 3, 6])

    add_para(doc, "El NSE D tiene la mayor proporción de su segmento que percibe a Gama como "
                  "igual o más económica (50.4%). El NSE C+/C tiene la menor (39.4%) — en el "
                  "segmento de mayor poder adquisitivo, más de 6 de cada 10 ven a Gama como "
                  "más cara. Sin diferencia relevante por género (53.5% femenino en subset vs "
                  "52.5% en Total).")

    add_caveat_box(doc,
        "La edad no está procesada en este reporte. La BBDD 2026 tiene la variable PD2 pero "
        "no estaba en formato numérico procesable para construir rangos etarios. "
        "Recomendación: re-exportar BBDD con PD2 codificada en rangos antes de ola 2027. "
        "(CV-18 del CU-8 v2)")

    add_metodologia_box(doc,
        "Subset definido como respondientes que en P33 eligieron Igual + Poco más económico + "
        "Mucho más económico (n=185). Complemento del subset = respondientes que eligieron "
        "Poco + Mucho más caros (n=217). Cruces con P30 (hábito por categoría), P32 (mejor "
        "precio por categoría), P21 (preferencia), P25 (misión última compra), y NSE. "
        "Las cifras del subset por categoría tienen n entre 8-22 — todos REFERENCIALES (CV-11 "
        "del CU-8 v2).")


def build_sec_4_9_misiones(doc):
    add_h2(doc, "4.9  Comportamiento y misiones de compra")

    add_callout_box(doc,
        "Gama es el líder absoluto del mercado en la misión de \"urgencia / pocos productos\" "
        "(12.2%), superando a Páramo (9.5%). Es el único territorio donde Gama es el primer "
        "nombre que viene a la mente. En las misiones de volumen (mercado grande), Gama está "
        "en posiciones bajas — hay una brecha estructural entre dónde Gama gana (urgencia) y "
        "dónde está el gasto mayor por viaje (mercado grande).")

    add_h3(doc, "4.9.1  Distribución de misiones de última compra (P25)")

    add_chart(doc, os.path.join(CHARTS_V6, "C20_misiones_gama.png"), width_in=6.8,
              caption="Misiones de compra: distribución en el mercado y posición de Gama por misión.")

    headers = ["Misión", "n", "%", "IC95"]
    rows = [
        ["Reabastecimiento parcial", "270", "67.2%", "[62.5, 71.6]"],
        ["Abastecimiento general completo", "94", "23.4%", "[19.4, 27.8]"],
        ["Urgencia (pocos productos)", "31", "7.7%", "[5.4, 10.8]"],
        ["Evento / celebración", "6", "1.5%", "[0.7, 3.2]"],
        ["No alimentos (electro., ropa)", "1", "0.2%", "—"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 2, 2, 3])

    add_para(doc, "Dos tercios de las últimas compras fueron de reabastecimiento parcial. El "
                  "shopper típico NO va al supermercado a \"hacer el mercado\" sino a "
                  "\"completar lo que falta\". Esto redefine el ticket promedio esperado y el "
                  "tipo de oferta que Gama debe tener lista en cada visita.")

    add_h3(doc, "4.9.2  Mejor supermercado por misión de compra (P26)")

    headers = ["Misión", "Líder", "Líder %", "Gama %", "Gap Gama-Líder", "Ninguno %"]
    rows = [
        ["Abastecimiento general completo", "Páramo", "21.6%", "7.2%", "-14.4 pp", "5.0%"],
        ["Reabastecimiento parcial", "Páramo", "16.4%", "8.7%", "-7.7 pp", "—"],
        ["Evento / celebración", "Páramo", "13.7%", "9.2%", "-4.5 pp", "—"],
        ["URGENCIA (pocos productos)", "GAMA", "12.2%", "12.2%", "0 pp (líder)", "24.6%"],
        ["No alimentos", "Hiper Líder", "6.0%", "0.5%", "-5.5 pp", "—"],
    ]
    add_table(doc, headers, rows, col_widths=[5.5, 2.5, 2, 2, 3, 2.5])

    add_callout_box(doc,
        "Hallazgo central: Gama es el supermercado LÍDER para la misión de \"compras de "
        "urgencia / pocos productos\" con 12.2% de las menciones. Es el único territorio "
        "donde Gama gana. Posicionamiento coherente con perfil de \"tienda de proximidad\".")

    add_h3(doc, "4.9.3  Supermercado habitual por categoría (P30) — Gama")

    add_para(doc, "Esta es la pregunta que mide hábito de compra real por categoría. Gama "
                  "tiene presencia transversal pero solo lidera en Farmacia.")

    headers = ["Categoría", "Gama %", "n Gama", "Líder", "Líder %"]
    rows = [
        ["Congelados", "8.0%", "32", "Páramo", "11.4%"],
        ["Salsas y Enlatados", "7.5%", "30", "Páramo", "20.4%"],
        ["Gaseosas, jugos y aguas", "7.2%", "29", "Páramo", "10.7%"],
        ["Galletas y confitería", "7.0%", "28", "Central Madeirense", "13.9%"],
        ["Productos básicos", "5.5%", "22", "Páramo", "19.7%"],
        ["Cuidado y limpieza hogar", "5.0%", "20", "Páramo", "11.2%"],
        ["FARMACIA (Gama líder)", "4.2%", "17", "Gama (empate)", "4.2%"],
        ["Cuidado Personal", "4.0%", "16", "—", "—"],
        ["Frutas, legumbres", "4.0%", "16", "Central Madeirense", "9.2%"],
        ["Proteínas (carne/pollo/charcut.)", "<3%", "<12", "Páramo", "~30%+"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 2, 2, 3.5, 2])

    add_para(doc, "Gama empata el liderazgo de hábito en Farmacia (4.2%), por encima de Plan "
                  "Suárez, Forum y Páramo. Base n=17 referencial pero patrón consistente con "
                  "el hallazgo de urgencia — Gama es el supermercado al que el shopper acude "
                  "para resolver necesidades puntuales.")

    add_h3(doc, "4.9.4  Dónde se compra (PF10) — supermercados nombrados")

    headers = ["Cadena", "Menciones espontáneas"]
    rows = [
        ["Páramo", "119"],
        ["Forum", "87"],
        ["Rio", "75"],
        ["Central Madeirense", "71"],
        ["Plazas", "59"],
        ["Luz", "52"],
        ["Gama", "50"],
        ["Plan Suárez", "35"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 4])

    add_para(doc, "Gama aparece en posición #7 en menciones espontáneas de cadenas habituales "
                  "(50 menciones de 367 cadena-goers = 13.6%). La distancia con Páramo (119) "
                  "es considerable. El ecosistema es de canales mixtos: 20% va a abastos/bodegas "
                  "y 19% a chinos en paralelo con cadenas — la competencia no es solo entre "
                  "supermercados.")

    add_metodologia_box(doc,
        "P25 razón única de la última compra (single-response). P26 mejor supermercado por "
        "misión (5 misiones × cadenas, single-response por misión). P30 supermercado habitual "
        "por categoría (15 categorías × cadenas, multi-response posible). PF10 multi-mention "
        "de tipos de establecimiento. Todas las cifras de Gama por categoría individual son "
        "REFERENCIALES (n<30 en la mayoría) — el patrón general es robusto, las cifras puntuales "
        "no son proyectables individualmente.")


def build_sec_4_10_comunicacion(doc):
    add_h2(doc, "4.10  Comunicación: PTL vs DTLS — análisis completo")

    add_callout_box(doc,
        "V6 incorpora por primera vez el análisis completo de la frase \"De Tu Lado Siempre\" "
        "(DTLS, preguntas P40-P42), que no había sido procesada en V3, V4 ni V5. El hallazgo "
        "central: PTL y DTLS tienen el mismo recall (ambas muy bajas, ~11-12%, sin diferencia "
        "significativa), pero DTLS activa territorio relacional (26%) que PTL prácticamente no "
        "toca (2%). Dirección clara para futuras decisiones comunicacionales.")

    add_h3(doc, "4.10.1  Gap de recall: la campaña no registra")

    headers = ["Métrica", "Resultado 2026", "n base", "Certeza"]
    rows = [
        ["Recall espontáneo \"PRECIOS DE TU LADO\" (PTL)", "10.7% — 43 personas", "402", "✅ Alta"],
        ["Recall espontáneo \"DE TU LADO SIEMPRE\" (DTLS)", "12.4% — 50 personas", "402", "✅ Alta"],
        ["Algún slogan de Gama recordado espontáneo", "4.2% — 17 personas", "402", "✅ Alta"],
        ["Recall asistido PTL (con estímulo)", "~11%", "17-50", "⚠ Base pequeña"],
        ["Interpretación PTL como mensaje de precio", "65%", "43", "⚠ Base referencial"],
        ["Interpretación PTL como solidaridad/apoyo", "35%", "43", "⚠ Base referencial"],
    ]
    add_table(doc, headers, rows, col_widths=[6.5, 3, 2, 3])

    add_chart(doc, os.path.join(CHARTS_V6, "C15_ptl_vs_dtls.png"), width_in=7.0,
              caption="Comparativo PTL vs DTLS: recall espontáneo similar, pero interpretación temática muy diferente.")

    add_h3(doc, "4.10.2  Test de diferencia PTL vs DTLS")

    headers = ["Estadístico", "Valor", "Interpretación"]
    rows = [
        ["z-test diferencia recall PTL vs DTLS", "z=-0.754, p=0.451", "NO significativo — empate técnico"],
        ["Recall combinado (cualquier slogan)", "4.2% espontáneo", "Recall publicitario general bajo"],
        ["Cobertura del mensaje", "88-89% no conoce ninguna frase", "Distribución del mensaje insuficiente"],
    ]
    add_table(doc, headers, rows, col_widths=[6.5, 3, 6])

    add_para(doc, "La diferencia de recall entre PTL (10.7%) y DTLS (12.4%) es de 1.7pp y no es "
                  "estadísticamente significativa. Ambas frases son prácticamente desconocidas "
                  "para el mercado. El problema central no es cuál frase es mejor — es que "
                  "ninguna ha tenido distribución suficiente para registrarse en la memoria de "
                  "los shoppers.")

    add_h3(doc, "4.10.3  Interpretación temática — diferencia clara")

    headers = ["Tema de interpretación", "DTLS (n=50)", "PTL (n=43)"]
    rows = [
        ["Precio / economía", "42%", "72%"],
        ["Acompañamiento / apoyo relacional", "26%", "2%"],
        ["Bienestar del cliente", "16%", "—"],
        ["Fidelidad / confianza", "8%", "—"],
        ["Otros", "8%", "26%"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 3, 3])

    add_callout_box(doc,
        "DTLS activa territorio relacional (26% acompañamiento + 16% bienestar + 8% fidelidad = "
        "50% en clave relacional). PTL es leída casi exclusivamente como promesa de precio (72%). "
        "Si el objetivo estratégico es posicionarse en territorio relacional — donde Gama tiene "
        "su activo más fuerte (Atención OR=5.73) — DTLS es la frase con mayor potencial semántico. "
        "(Hallazgo direccional — bases referenciales, decisión final depende de claim testing "
        "cuantitativo y de presupuesto de distribución.)", color=BLUE_INFO)

    add_h3(doc, "4.10.4  Agrado de cada frase (P38 y P41)")

    headers = ["Frase", "Base", "Top2 Agrado %"]
    rows = [
        ["PTL \"Precios de Tu Lado\"", "43", "53.5%"],
        ["DTLS \"De Tu Lado Siempre\"", "50", "42.0%"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 3, 4])

    add_para(doc, "PTL genera mayor agrado declarado entre quienes la recuerdan (53.5% vs 42.0%), "
                  "una diferencia de 11.5pp en magnitud — pero NO estadísticamente significativa "
                  "dado el tamaño de las bases. Hallazgo direccional, no concluyente.")

    add_para(doc, "Tensión y decisión:", bold=True)
    add_para(doc, "Hay tensión entre los hallazgos. DTLS gana en interpretación relacional pero "
                  "pierde en agrado declarado. PTL tiene mejor agrado pero está atrapada en el "
                  "territorio del precio (donde Gama no compite bien). La decisión sobre cuál "
                  "frase priorizar depende de:")
    add_bullet(doc, "Objetivo estratégico: ¿posicionar relacional (DTLS) o profundizar el claim "
                    "de precio (PTL, que requeriría tener los precios para sostenerlo)?")
    add_bullet(doc, "Presupuesto de distribución: si no hay inversión para llevar cualquier frase "
                    "a recall masivo, la elección es secundaria.")
    add_bullet(doc, "Claim testing cuantitativo en ola 2027: validar cuál frase resuena mejor en "
                    "el target con bases suficientes.")

    add_metodologia_box(doc,
        "P35 recall espontáneo de cualquier publicidad de Gama. P37 recall asistido específico "
        "de PTL. P38 agrado de PTL (5 puntos Likert). P39 interpretación abierta de PTL. "
        "P40-P42 idénticos para DTLS. z-test Newcombe-Wilson para diferencias de proporciones. "
        "Bases bajas (n=43 y n=50) hacen que las diferencias direccionales no alcancen "
        "significancia al 95%. Interpretación temática codificada manualmente desde respuestas "
        "abiertas — categorías inductivas. (CV-3, CV-4, CV-7 del CU-8 v1.)")


def build_sec_4_11_wow(doc):
    add_h2(doc, "4.11  Evolución comparativa 2025→2026 (referencia)")

    add_callout_box(doc,
        "Este bloque usa los datos 2025 exclusivamente como REFERENCIA para entender qué cambió. "
        "Los datos 2026 son el análisis primario del documento. El mercado venezolano de "
        "supermercados de cadena tuvo movimiento significativo en 2026: 9 de los 10 cambios "
        "significativos son aumentos de competidores. Gama es la única gran cadena con posición "
        "estable.")

    add_h3(doc, "4.11.1  Gama: posición estable (0/8 indicadores significativos)")

    headers = ["Indicador", "Delta WoW", "Significancia"]
    rows = [
        ["TOM", "+2.2 pp", "no significativo"],
        ["Asistida", "+2.2 pp", "no significativo"],
        ["Consideración", "+4.3 pp", "no significativo"],
        ["Compra 3m", "-0.2 pp", "no significativo"],
        ["Preferida", "-1.7 pp", "no significativo"],
        ["Última compra", "-1.2 pp", "no significativo"],
        ["Habitual", "+0.8 pp", "no significativo"],
        ["Misiones", "+4.9 pp", "no significativo"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 3, 4])

    add_para(doc, "Lectura: Gama es la única gran cadena del mercado que mantiene posición "
                  "estadísticamente estable entre 2025 y 2026. En un mercado donde 9 de 10 "
                  "cambios significativos son aumentos de competidores, la estabilidad es una "
                  "fortaleza relativa — no un estancamiento.")

    add_h3(doc, "4.11.2  Movimiento competitivo 2025→2026 (Total)")

    add_chart(doc, os.path.join(CHARTS_V6, "C21_wow_movimiento.png"), width_in=6.5,
              caption="Movimientos significativos 2025→2026 — Rio y Páramo crecen, CM cae, Gama estable.")

    headers = ["Cadena", "Métrica", "Delta", "Significancia"]
    rows = [
        ["Rio", "TOM", "+17.0 pp", "sig 99% (p_adj<0.0001)"],
        ["Rio", "Consideración", "+19.6 pp", "sig 99%"],
        ["Rio", "Compra 3m", "+12.5 pp", "sig 99%"],
        ["Rio", "Preferida", "+4.0 pp", "tendencia (p_adj=0.059)"],
        ["Páramo", "TOM", "+12.1 pp", "sig 99% (p_adj=0.0003)"],
        ["Páramo", "Asistida", "+10.8 pp", "sig 99%"],
        ["Páramo", "Consideración", "+17.3 pp", "sig 99%"],
        ["Central Madeirense", "Compra 3m", "-7.7 pp", "sig 95% (p_adj=0.033)"],
        ["Central Madeirense", "Preferida", "-5.5 pp", "tendencia (p_adj=0.053)"],
        ["Gama", "(todos los 8)", "~0 pp", "no significativo (0/8)"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 2.5, 4])

    add_chart(doc, os.path.join(CHARTS_V5, "D9_2x2_competitivo.png"), width_in=6.5,
              caption="Mapa 2x2 movimiento competitivo: Rio y Páramo ganan terreno, CM pierde, Gama estable.")

    add_caveat_box(doc,
        "Factor de ponderación muestral 2025 no disponible (variable @PONDERAR_1 era todo cero "
        "en la base recibida). Composición geográfica difiere significativamente entre olas "
        "(Libertador sobrerrepresentado en 2025). n_2025=785, n_2026=402, corrección "
        "Benjamini-Hochberg FDR. (CV-WOW-001 y CV-WOW-002)")

    add_metodologia_box(doc,
        "Análisis wave-over-wave con z-test Newcombe-Wilson para diferencias de proporciones, "
        "corrección BH-FDR (q<0.05) por comparaciones múltiples sobre 57 ítems comparables. "
        "Test de comparabilidad muestral previo (chi-cuadrado de homogeneidad NSE, geografía, "
        "género, edad) — muestras consideradas suficientemente comparables aunque con caveats "
        "geográficos documentados. Documento de soporte: CU-7 wave-over-wave v1.")


# ============================================================================
# PARTE II — CUALITATIVO
# ============================================================================

def build_parte2_header(doc):
    add_h1(doc, "PARTE II — Insights Cualitativos (hipótesis y señales)")
    add_callout_box(doc,
        "REGLA DE LECTURA OBLIGATORIA DE ESTA PARTE: todo hallazgo aquí es insight cualitativo "
        "o hipótesis. El corpus base es 12 documentos de focus groups (8 sesiones presenciales "
        "+ 4 online), 6 segmentos cubiertos, ~42.094 caracteres totales. Máximo ~42 personas "
        "presenciales. El cualitativo establece mecanismos, barreras y territorios de marca — "
        "NO prevalencia estadística. Una conclusión del cuantitativo con n=402 tiene mayor peso "
        "para decisiones estratégicas que un verbatim de 1-2 personas. El peso de los insights "
        "cuali es proporcional al número de grupos / participantes que los expresan.")


def build_sec_5_regla(doc):
    add_h2(doc, "5.1  Regla de lectura del bloque cualitativo")

    add_para(doc, "Esta parte del documento aplica las siguientes reglas, ratificadas por el "
                  "Owner el 2026-05-18 después del feedback de la consultora del lado Gama:")

    add_bullet(doc, "Los hallazgos cualitativos se presentan como HIPÓTESIS, no como conclusiones "
                    "trianguladas co-iguales al cuantitativo.")
    add_bullet(doc, "Cada insight lleva el ícono 💡 que indica explícitamente su naturaleza "
                    "cualitativa.")
    add_bullet(doc, "Las afirmaciones cuantitativas que aparecen en esta parte (con ícono ✅) "
                    "vienen del análisis cuantitativo y se citan aquí solo como contexto para "
                    "los insights cualitativos.")
    add_bullet(doc, "No se usa el lenguaje \"Triangulado\" como en V4 — el cualitativo se ofrece "
                    "como insight para diseño de activaciones y como hipótesis para confirmar "
                    "en ola 2027.")

    add_para(doc, "Los 5 insights cualitativos de esta parte:", bold=True)
    add_bullet(doc, "5.2 Mecanismo del driver de atención: hipótesis \"acompañamiento-guía\" "
                    "(6 de 7 grupos FG — patrón robusto)")
    add_bullet(doc, "5.3 Barrera del precio: hipótesis del \"sifrinaje\" (1 verbatim central + "
                    "patrón consistente)")
    add_bullet(doc, "5.4 Datos digitales como contexto (3 de 4 docs digitales + 5 de 7 segmentos)")
    add_bullet(doc, "5.5 Arquetipo femenino emergente (6 de 7 grupos FG, no inducido)")

    add_metodologia_box(doc,
        "Análisis cualitativo realizado por analista único siguiendo el protocolo de análisis "
        "temático de Braun & Clarke (2006). Codificación deductivo-inductiva en dos pasadas: "
        "(1) deductiva sobre el framework de drivers cuantitativos V3 (9 categorías a priori); "
        "(2) inductiva sobre patrones emergentes. 52 códigos totales. Sin coeficiente Kappa "
        "inter-codificador (analista único — práctica estándar en investigación aplicada "
        "cuando recursos no permiten doble codificación). Para ola 2027 se recomienda doble "
        "codificación en un 20% del material.")


def build_sec_5_2_acompanamiento(doc):
    add_h2(doc, "5.2  Mecanismo del driver de atención: \"acompañamiento-guía\"")

    add_callout_box(doc,
        "Cuantitativo (alta certeza): OR=5.73*** convergente en 4 métodos. Cualitativo "
        "(💡 6/7 grupos FG, patrón robusto): Gama como \"guía experta de confianza\" que "
        "acompaña la compra. El mecanismo no sería el reconocimiento nominal sino el "
        "acompañamiento activo.", color=BLUE_INFO)

    add_h3(doc, "5.2.1  Lo que el cuantitativo establece (✅ alta certeza)")

    add_para(doc, "OR=5.73 — quien asocia Gama con atención tiene 5.7 veces más odds de "
                  "preferirla. Convergencia de 4 métodos independientes (logit, Random Forest, "
                  "SHAP, razón espontánea citada por el 53% de los preferentes). Resultado "
                  "cuantitativo robusto.")

    add_h3(doc, "5.2.2  La hipótesis cualitativa (💡 6/7 grupos FG)")

    add_para(doc, "El corpus de focus groups documenta espontáneamente en 6 de 7 grupos la "
                  "imagen de Gama como \"guía experta de confianza\" que acompaña la compra. "
                  "El mecanismo no sería el reconocimiento nominal (\"te conocen por tu nombre\") "
                  "sino el acompañamiento activo (\"siempre hay alguien que te orienta\").")

    add_verbatim(doc,
        "Una mujer joven, pero con experiencia. Que sea como una guía, que te diga: "
        "'Mira, hoy llegó el pan árabe fresco, llévatelo'.",
        "Mary Francis Bossia, Ocasionales 18-30")
    add_verbatim(doc,
        "Alguien de Caracas que te oriente y sea como una guía de compras, que no se "
        "quede atrás con la tecnología.",
        "Moisés Torrealba, Ocasionales 18-30")

    add_h3(doc, "5.2.3  Implicación para comunicación")

    add_para(doc, "Si la hipótesis es correcta, el mensaje correcto es \"en Gama siempre hay "
                  "alguien que te ayuda\" (creíble para el shopper), NO \"en Gama te conocen "
                  "por tu nombre\" (hipérbole difícil de verificar). La diferencia es relevante "
                  "para el desarrollo creativo de cualquier campaña que quiera capitalizar el "
                  "activo de atención.")

    add_caveat_box(doc,
        "Hipótesis cualitativa robusta (6/7 FG, emergente no inducido por el moderador). "
        "Requiere claim testing cuantitativo para confirmar como plataforma comunicacional. "
        "No se debe ejecutar campaña basada solo en este insight sin validación previa.")

    add_metodologia_box(doc,
        "Codificación deductiva categoría D1 (Atención/Servicio/Personal del driver V3) + "
        "categoría inductiva I1 (Reconocimiento personal/acompañamiento). Frecuencia: 6 de 7 "
        "grupos de focus group mencionan el patrón espontáneamente, sin ser preguntados "
        "directamente por personificación o por estilo de atención. Verbatims literales del "
        "corpus en español venezolano (no parafraseados).")


def build_sec_5_3_sifrinaje(doc):
    add_h2(doc, "5.3  Barrera del precio: hipótesis del \"sifrinaje\"")

    add_callout_box(doc,
        "Cuantitativo (alta certeza): precio OR=1.03 NS, 54% percibe Gama como cara, paradoja "
        "documentada (la percepción de precio no afecta la preferencia entre quienes ya la "
        "conocen). Cualitativo (💡 1 verbatim central + patrón consistente): la barrera no es "
        "un cálculo económico sino una autoexclusión simbólica — \"Gama es para otro tipo de "
        "persona\".", color=BLUE_INFO)

    add_h3(doc, "5.3.1  Lo que el cuantitativo establece (✅ alta certeza)")

    add_bullet(doc, "Precio OR=1.03 (NS, SHAP #10): el precio NO predice preferencia.")
    add_bullet(doc, "54% percibe Gama como cara (P33 NETO caro).")
    add_bullet(doc, "Pero el cruce P33×P21 muestra que la percepción de precio sí predice "
                    "preferencia a nivel agregado (OR=2.4, sig). Esto resuelve la paradoja: "
                    "el precio no es driver de elección entre quien ya conoce Gama, pero la "
                    "percepción agregada sí actúa como barrera de entrada para quien nunca la "
                    "ha probado.")

    add_h3(doc, "5.3.2  La hipótesis cualitativa (💡 patrón en varios grupos)")

    add_para(doc, "El corpus FG sugiere que la barrera no es un cálculo económico comparativo "
                  "sino una autoexclusión simbólica — \"Gama es para otro tipo de persona\". El "
                  "shopper se autoexcluye antes de entrar a verificar precios.")

    add_verbatim(doc,
        "A mí me parece que es un poco costoso... el sifrinaje y tal, pero para hacer un "
        "mercado grande... es mejor ir al Lux o al Páramo.",
        "Azahara Betancourt, Ocasionales 18-30")

    add_caveat_box(doc,
        "\"Sifrinaje\" es vocabulario espontáneo del informante — no es una caracterización "
        "de Gama hecha por el equipo analítico. Refleja cómo segmentos de no-usuarios codifican "
        "la barrera de entrada percibida. (Caveat literal obligatorio C-LIT-DW1 del BR-2 V4.)")

    add_h3(doc, "5.3.3  Implicación para activación del Segmento 2")

    add_para(doc, "Si la barrera es identitaria, los descuentos de precio no la mueven — "
                  "comunicar precio bajo refuerza la distancia simbólica (\"esto es para otra "
                  "gente, baja el precio para acomodarla\"). La activación correcta del "
                  "Segmento 2 (Pragmáticos Convertibles, 33% del mercado, 0% preferencia "
                  "actual) sería una primera experiencia que resignifique la pertenencia — "
                  "no una oferta económica.")

    add_para(doc, "El cuali V4 enriquece el perfil de los Pragmáticos Convertibles (33% del "
                  "mercado) con una dimensión que el cuestionario cuantitativo no podía "
                  "capturar: la barrera de conversión no es solo el precio como cálculo "
                  "económico, también es el precio como señal de pertenencia social. Esta "
                  "capa adicional refina la estrategia de activación: el anzuelo correcto es "
                  "una razón específica para venir, no una oferta de descuento.", italic=True)

    add_metodologia_box(doc,
        "Categoría inductiva I2 (Barrera identitaria / sifrinaje), emergente del análisis del "
        "corpus. Patrón consistente con menciones en múltiples grupos del concepto de "
        "auto-exclusión por código social. Hipótesis cualitativa — requiere validación "
        "cuantitativa estructurada en ola 2027 (cuestionario de barreras con escala que "
        "distinga precio-como-cálculo vs precio-como-identidad).")


def build_sec_5_4_digital(doc):
    add_h2(doc, "5.4  Datos digitales como contexto (sin recomendaciones)")

    add_callout_box(doc,
        "Esta sección presenta hallazgos sobre el canal digital de Gama (Gama Club, "
        "inventario online, frescos online) que emergieron del análisis cualitativo. Se "
        "incluyen como CONTEXTO porque exceden el brief original del estudio (notoriedad y "
        "brand health). NO se incluyen como recomendaciones — esa es una decisión del Owner "
        "(CO-3) basada en la consideración de que estos hallazgos abren conversación de scope "
        "adicional para una segunda etapa de consultoría sobre el producto digital, no "
        "constituyen entregable del scope original contratado.")

    add_h3(doc, "5.4.1  Gama Club: interés universal + opacidad digital")

    add_para(doc, "Hallazgo cualitativo (💡 6/7 grupos FG): el Gama Club genera interés genuino "
                  "y transversal en todos los segmentos, pero su saldo es completamente "
                  "invisible en el canal digital. El usuario no puede ver cuánto tiene "
                  "acumulado, qué puede canjear, ni cuándo.")

    add_verbatim(doc,
        "Las millas... yo ni sabía que eso se podía usar para pagar hasta hace poco. Yo "
        "pensaba que era como los puntos del banco que uno nunca usa.",
        "Gregory, Frecuentes 18-30")

    add_h3(doc, "5.4.2  Canal digital — dos barreras documentadas")

    add_para(doc, "Hallazgo cualitativo (💡 corpus digital + 5/7 segmentos FG): dos barreras "
                  "operacionales distintas:")

    add_para(doc, "Barrera 1 — Inconsistencia de inventario (3/4 docs uso digital):", bold=True)
    add_verbatim(doc,
        "Yo dejé de usar la página porque a veces pides algo que dice que hay y a la "
        "media hora te llaman para decirte que no tienen.",
        "Mujer, 42 años, segmento Frecuentes 31-50 (sesión online cualitativa)")

    add_para(doc, "Barrera 2 — Desconfianza en frescos online (5/7 segmentos FG):", bold=True)
    add_para(doc, "El shopper no confía en que los perecederos online tengan la misma calidad "
                  "que en tienda. Solución conceptual indicada por el propio shopper: empezar "
                  "la propuesta digital por no-perecederos.")

    add_h3(doc, "5.4.3  Nota de scope")

    add_para(doc, "Estos hallazgos son consistentes y de alta accionabilidad, pero su "
                  "tratamiento como recomendaciones estratégicas excede el brief original del "
                  "estudio (notoriedad y brand health de marca). El Owner (CO-3) los incluye "
                  "como contexto en este documento para que Gama tenga visibilidad del insight, "
                  "y sugiere abrir conversación de scope adicional para una segunda etapa de "
                  "consultoría enfocada en producto digital, si Gama tiene interés.", italic=True)

    add_metodologia_box(doc,
        "Insights del corpus digital (4 docs de sesiones online) + 5 de 7 segmentos FG "
        "presenciales. Atribuciones de verbatims por segmento/edad (no por nombre/seudónimo "
        "del informante) según wording obligatorio C-LIT-DW4-c del BR-2 V4. Decisión CO-3 "
        "del Owner aplicada: hallazgos como contexto, no como recomendaciones del estudio.")


def build_sec_5_5_arquetipo(doc):
    add_h2(doc, "5.5  Arquetipo femenino emergente — hipótesis de territorio disponible")

    add_callout_box(doc,
        "Hipótesis cualitativa emergente (💡 6 de 7 grupos FG, no inducido por moderador): sin "
        "ninguna pregunta directa sobre personificación, 6 de 7 grupos convergieron "
        "espontáneamente en el mismo imaginario: Gama como mujer de 35-45 años, caraqueña, "
        "profesional, activa, guía experta de confianza. Este arquetipo no está apropiado por "
        "ningún competidor directo (Rio, Páramo, CM).", color=BLUE_INFO)

    add_h3(doc, "5.5.1  El patrón emergente")

    add_verbatim(doc,
        "Yo me la imagino como una mujer, como de unos 35 años, super activa, líder, "
        "entusiasta, que sabe lo que hace.",
        "Moisés Torrealba, Ocasionales 18-30")

    add_para(doc, "Coherencia notable: el arquetipo \"guía experta\" es la personificación "
                  "exacta de lo que el OR=5.73 mide a nivel cuantitativo. El cualitativo "
                  "espontáneamente personifica el activo que el cuantitativo mide.")

    add_h3(doc, "5.5.2  Implicación para comunicación")

    add_para(doc, "Si la hipótesis se confirma cuantitativamente, podría ser un territorio "
                  "diferenciador de comunicación. Pero antes de cualquier desarrollo creativo "
                  "se requieren validaciones obligatorias.")

    add_caveat_box(doc,
        "Este hallazgo describe el imaginario de marca que el shopper ya proyecta "
        "espontáneamente — no propone una campaña específica de personificación. La "
        "materialización creativa (figura concreta, voz, presencia visual) requiere prueba de "
        "concepto cuantitativa con el target antes de cualquier ejecución. (Caveat literal "
        "obligatorio C-LIT-DW2-a del BR-2 V4.)")

    add_para(doc, "Prerrequisitos obligatorios antes de cualquier desarrollo creativo en esta "
                  "dirección:", bold=True)
    add_bullet(doc, "Prueba de concepto cuantitativa con el target para validar accionabilidad "
                    "como plataforma de comunicación.")
    add_bullet(doc, "Verificación previa de Gama sobre compatibilidad con contratos de imagen "
                    "vigentes.")
    add_bullet(doc, "Verificación de compatibilidad con plataforma de marca global o "
                    "corporativa (si aplica).")

    add_metodologia_box(doc,
        "Categoría inductiva I5 (Personificación femenina espontánea). Patrón emergente "
        "robusto: 6 de 7 grupos de focus group converger espontáneamente en el mismo "
        "imaginario sin pregunta directa de personificación. La consistencia inter-grupos "
        "(diferentes edades, frecuencias de compra, ubicaciones) refuerza la robustez del "
        "patrón. Sin embargo, sigue siendo un hallazgo cualitativo — la decisión de "
        "ejecutar una campaña basada en este arquetipo requiere validación cuantitativa "
        "previa con concept testing.")


# ============================================================================
# PARTE III — HIPÓTESIS ESPECIALES
# ============================================================================

def build_parte3_header(doc):
    add_h1(doc, "PARTE III — Hipótesis especiales")


def build_sec_6_express(doc):
    add_h2(doc, "6.1  Hipótesis del formato Express — validación parcial")

    add_callout_box(doc,
        "Hipótesis propuesta por la consultora del lado Gama: \"Las misiones de compra de "
        "reposición son fuertes en Gama porque el formato de tienda al que van son Express "
        "(más pequeñas, menos variedad de categorías y/o marcas)\". Status: COMPATIBLE con "
        "los datos pero NO PROBADA directamente — la BBDD 2026 no contiene la variable de "
        "formato/tamaño de tienda necesaria para confirmación formal.")

    add_h3(doc, "6.1.1  Lo que sí se puede afirmar (✅ confirmado en datos)")

    add_bullet(doc, "Gama lidera la misión de urgencia con 12.2% (P26) — confirmado en CU-8 v1. "
                    "Es el único territorio donde Gama gana en el mercado.")
    add_bullet(doc, "Gama empata el liderazgo de hábito en Farmacia (4.2%, P30) — coherente con "
                    "perfil de tienda de resolución de necesidades puntuales.")
    add_bullet(doc, "El shopper típico de Gama es residente de la zona (62%) y decisor único o "
                    "principal (100%) — perfil de \"vecino que va a completar el mercado\".")

    add_h3(doc, "6.1.2  Cruce nuevo: urgencia × NSE")

    headers = ["NSE", "n base", "Gama urgencia %", "Líder urgencia en NSE", "Líder %", "¿Gama es líder?"]
    rows = [
        ["C+/C", "104", "16.3%", "Gama", "16.3%", "SÍ"],
        ["D", "127", "15.0%", "Gama", "15.0%", "SÍ"],
        ["E", "171", "7.6%", "Páramo", "13.5%", "NO"],
    ]
    add_table(doc, headers, rows, col_widths=[2.5, 2, 3, 3.5, 2, 2.5])

    add_chart(doc, os.path.join(CHARTS_V6, "C19_urgencia_nse.png"), width_in=6.5,
              caption="Liderazgo de urgencia por NSE — Gama lidera en C+/C y D, no en E.")

    add_para(doc, "Hallazgo: Gama lidera urgencia en NSE C+/C (16.3%) y NSE D (15.0%) — los "
                  "dos segmentos de mayor ingreso de la muestra. En NSE E, Páramo lidera con "
                  "13.5% vs Gama 7.6%. Patrón compatible con la hipótesis Express: los NSE más "
                  "altos son más proclives a usar tiendas de conveniencia (\"vale más mi tiempo "
                  "que la diferencia de precio\").")

    add_h3(doc, "6.1.3  Cruce nuevo: urgencia × municipio")

    headers = ["Municipio", "Gama urgencia %", "Líder urgencia", "Gama es líder"]
    rows = [
        ["Chacao", "17.1%", "Luz", "NO"],
        ["Libertador", "15.0%", "Gama", "SÍ"],
        ["Baruta", "11.5%", "Plazas", "NO"],
        ["El Hatillo", "9.7%", "Central M.", "NO (base baja)"],
        ["Sucre", "8.9%", "Rio", "NO (base baja)"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 3])

    add_para(doc, "Gama tiene su mayor participación de urgencia en Chacao (17.1%) y Libertador "
                  "(15.0%) — zonas de mayor presencia física de la marca. En Libertador Gama "
                  "es el líder de urgencia. La correlación urgencia-presencia física es "
                  "compatible con la hipótesis Express (a mayor cobertura de formato pequeño "
                  "en una zona, mayor liderazgo en urgencia) — pero no la prueba.")

    add_h3(doc, "6.1.4  Interpretación de la hipótesis")

    headers = ["Evidencia A FAVOR", "Evidencia que FALTARÍA para confirmar"]
    rows = [
        ["Gama lidera urgencia a nivel total (12.2%, líder)",
         "Variable de formato operacional (Express/convencional/grande) NO disponible en BBDD 2026"],
        ["Patrón transversal a NSE C+/C y D (segmentos que usarían formato de conveniencia)",
         "Datos de radio de cobertura por sucursal NO disponibles en encuesta"],
        ["En Libertador (zona de mayor presencia Gama), Gama lidera urgencia",
         "Para validar formal: cruzar base operacional de Gama (lista de sucursales con formato) con georreferenciación de respondientes"],
        ["Misión urgencia + liderazgo en Farmacia (hábito) son coherentes con perfil de tienda de proximidad",
         "Concept testing cuantitativo: preguntar directo al shopper qué tipo de tienda visita"],
    ]
    add_table(doc, headers, rows, col_widths=[7, 8])

    add_callout_box(doc,
        "Conclusión: la hipótesis del formato Express es COMPATIBLE con todos los datos "
        "disponibles, pero NO PROBADA formalmente. Para confirmación se requiere cruzar la "
        "base operacional de Gama (lista de sucursales con su formato) con datos de "
        "georreferenciación de respondientes — análisis fuera del alcance de la encuesta. "
        "Decisión sugerida: si Gama puede compartir esa base operacional, se puede cerrar la "
        "hipótesis en una ronda corta de análisis adicional.")

    add_metodologia_box(doc,
        "Verificación de variables BBDD: el script de detección encontró que la única variable "
        "relacionada con \"formato\" es P45 (disposición a comprar si Gama abre sucursal) — que "
        "no es una variable de tipo de tienda. No hay en los 295 campos de la BBDD ninguna "
        "variable que indique si el respondiente fue a un Gama Express, Gama convencional, "
        "etc. Para validación formal de la hipótesis, se necesitan datos operacionales internos "
        "de Gama no incluidos en este estudio. Reportar como \"compatible con la hipótesis\" — "
        "nunca como \"confirmado que el formato Express explica la urgencia\". (CV-16 del CU-8 v2.)")


# ============================================================================
# PARTE IV — SÍNTESIS Y RECOMENDACIONES
# ============================================================================

def build_parte4_header(doc):
    add_h1(doc, "PARTE IV — Síntesis y Recomendaciones")


def build_sec_7_1_tesis(doc):
    add_h2(doc, "7.1  Tesis V6 — diagnóstico central")

    add_callout_box(doc,
        "Gama posee el activo relacional más diferenciador del mercado (Atención driver #1, "
        "OR=5.73*** con convergencia de 4 métodos), su posición es estadísticamente estable y "
        "su servicio 24h es una ventaja funcional sin equivalente en los competidores. Pero "
        "el crecimiento de Gama es cero mientras Rio y Páramo penetran el segmento C+/C — el "
        "segmento donde Gama tiene su mayor fortaleza. La estrategia correcta requiere tres "
        "movimientos simultáneos:")

    add_bullet(doc, "1. Comunicar el activo que mueve la preferencia (atención), no el precio "
                    "que no la mueve.")
    add_bullet(doc, "2. Activar el 33% convertible (Segmento 2) con el mecanismo correcto: "
                    "primera experiencia guiada, no descuento.")
    add_bullet(doc, "3. Responder al avance en C+/C antes de la ola 2027 — mientras la "
                    "posición sea ventajosa.")


def build_sec_7_2_argumentos(doc):
    add_h2(doc, "7.2  Los 4 argumentos MECE")

    add_h3(doc, "Argumento 1 — El activo es real, robusto y sin equivalente competitivo ✅")
    add_para(doc, "OR=5.73 convergente en 4 métodos independientes. 53% de preferentes citan "
                  "atención como razón espontánea #1. Rio y Páramo no tienen equivalente "
                  "documentado. Limpieza es driver secundario (OR=3.99*). El activo existe y "
                  "es defensivo.")

    add_h3(doc, "Argumento 2 — El activo no está siendo comunicado ✅")
    add_para(doc, "Recall espontáneo PTL = 0/17 = 0%. 65% interpreta PTL como mensaje de "
                  "precio (atributo OR=1.03 NS). El recurso comunicacional está invertido en el "
                  "territorio equivocado.")

    add_h3(doc, "Argumento 3 — El segmento de conversión existe, es grande y tiene mecanismo de barrera conocido (hipótesis) ✅ + 💡")
    add_para(doc, "Segmento 2 (33% del mercado): 0% preferencia actual, menor resistencia al "
                  "precio (3.44 vs 3.66/5 cuantitativo). Hipótesis cualitativa: la barrera es "
                  "identitaria, la activación correcta es primera experiencia guiada. El "
                  "mercado disponible de CM (-7.7pp compra, sig 95%) complementa esta oportunidad.")

    add_h3(doc, "Argumento 4 — La presión competitiva en C+/C se intensifica (hipótesis WoW pendiente confirmación 2027) ⚠")
    add_para(doc, "Rio +25.8pp TOM en C+/C (exploratorio sig 99%). Páramo +22.3pp TOM en C+/C "
                  "(exploratorio sig 99%). Gama +2.2pp (estable). La posición actual es sólida; "
                  "la velocidad de los competidores en el segmento natural de Gama es la señal "
                  "de alerta más importante del estudio.")


def build_sec_7_3_recomendaciones(doc):
    add_h2(doc, "7.3  Recomendaciones priorizadas")

    add_para(doc, "Las recomendaciones son evaluación analítica — no un dato empírico. La "
                  "prioridad asume continuidad de condiciones de mercado y está ordenada por "
                  "combinación de certeza de evidencia + accionabilidad estimada. Las "
                  "recomendaciones de canal digital (Gama Club, inventario, frescos online) "
                  "están EXCLUIDAS por decisión Owner CO-3 — exceden el brief original. La "
                  "señal sobre Rio en proteínas está EXCLUIDA por decisión Owner CO-1 — no "
                  "tiene sustento cuantitativo suficiente.",
             italic=True, size=10, color=GRAY_MID)

    add_h3(doc, "Recomendación 1 — Quick win: Reorientar el anzuelo promocional")

    add_para(doc, "Qué:", bold=True)
    add_para(doc, "Mantener stickers/Cashea/Club como palanca táctica (OR=3.64** confirmado), "
                  "pero cambiar el mensaje de \"precio bajo\" a \"razón específica para venir "
                  "esta semana\".")
    add_para(doc, "Por qué:", bold=True)
    add_para(doc, "Las promociones son el Driver #3 (OR=3.64** — cuantitativo confirmado). "
                  "Pero el mensaje \"precio bajo\" refuerza la distancia simbólica del Seg 2 "
                  "(hipótesis cualitativa del sifrinaje). El mismo activo, con el mensaje "
                  "correcto, no activa la barrera identitaria.")
    add_para(doc, "Mensaje correcto:", bold=True)
    add_para(doc, "\"Esta semana hay una razón para venir a Gama\" (acceso a experiencia/"
                  "producto específico) vs \"Precios de tu lado\" (mensaje de precio bajo que "
                  "OR=1.03 NS confirma irrelevante para preferencia).")
    add_para(doc, "Evidencia: ✅ Cuantitativo OR=3.64** + 💡 Hipótesis cualitativa del mecanismo", size=10)

    add_h3(doc, "Recomendación 2 — Mid-term: Diseño de activación para Segmento 2 (33% del mercado)")

    add_para(doc, "Secuencia de activación:", bold=True)
    add_bullet(doc, "Anzuelo: razón específica para la primera visita (no precio bajo).")
    add_bullet(doc, "Primera experiencia guiada que active el activo de atención (OR=5.73).")
    add_bullet(doc, "Retención: Gama Club visible + saldo acumulado comunicado.")

    add_para(doc, "Por qué:", bold=True)
    add_para(doc, "El Seg 2 tiene 0% preferencia actual + menor resistencia al precio en "
                  "cuantitativo (3.44 vs 3.66/5). La hipótesis cualitativa indica que la "
                  "barrera es identitaria — se neutraliza con la primera experiencia de "
                  "acompañamiento, no con descuentos. El motor de conversión es resignificar "
                  "la pertenencia.")
    add_para(doc, "Evidencia: ✅ Cuantitativo k-means Seg 2 + 💡 Hipótesis cualitativa mecanismo identitario", size=10)

    add_h3(doc, "Recomendación 3 — Mid-term: Comunicar precio en categorías de \"easy win\"")

    add_para(doc, "Qué:", bold=True)
    add_para(doc, "Trabajar comunicación de precio en categorías donde Gama está cerca del "
                  "líder con bases de percepción muy diluidas: Farmacia (gap -0.7pp, 85% "
                  "\"Ninguno\"), Licores (gap -1.2pp, 66% \"Ninguno\"). El gap es mínimo y la "
                  "vacante de percepción es alta — comunicar precio aquí puede dar liderazgo "
                  "perceptual sin recortar precio realmente.")
    add_para(doc, "Por qué:", bold=True)
    add_para(doc, "Estas son las dos únicas categorías donde Gama tiene posición cercana al "
                  "top en P32. En Farmacia ya es líder de hábito (P30) — convertir hábito en "
                  "convicción de precio es accionable a bajo costo.")
    add_para(doc, "Evidencia: ✅ Cuantitativo P30 y P32", size=10)

    add_h3(doc, "Recomendación 4 — Long-term: Estrategia ofensiva en C+/C antes de ola 2027")

    add_para(doc, "Diseñar y ejecutar respuesta al avance de Rio y Páramo en C+/C mientras la "
                  "posición de Gama (TOM 60.6%) aún es ventajosa. La posición actual es "
                  "sólida; la ventana de respuesta ofensiva se cierra si Rio/Páramo completan "
                  "la penetración del segmento.")
    add_para(doc, "Evidencia: ⚠ Hipótesis WoW C+/C — pendiente confirmación ola 2027", size=10)

    add_h3(doc, "Recomendación 5 — Long-term: Decidir entre PTL y DTLS con claim testing")

    add_para(doc, "DTLS activa territorio relacional (26% en cuali) que PTL prácticamente no "
                  "toca (2%). Si el objetivo estratégico es posicionarse en territorio "
                  "relacional — donde Gama tiene su activo más fuerte — DTLS tiene mayor "
                  "potencial semántico. Pero PTL tiene mayor agrado declarado (53% vs 42% "
                  "top2). La decisión requiere claim testing cuantitativo en ola 2027 con "
                  "bases suficientes para distinguir las dos.")
    add_para(doc, "Evidencia: ⚠ Hipótesis interpretativa (bases referenciales)", size=10)

    add_h3(doc, "Recomendación 6 — Long-term sujeta a CO-3 + DW-2: Explorar plataforma \"acompañamiento-guía\"")

    add_para(doc, "Si la hipótesis cualitativa del arquetipo se confirma con concept testing "
                  "cuantitativo, explorar el arquetipo femenino espontáneo (35-45 años, "
                  "guía experta) como plataforma de comunicación.")
    add_para(doc, "Prerrequisitos obligatorios antes de cualquier desarrollo creativo:", bold=True)
    add_bullet(doc, "Prueba de concepto cuantitativa con el target para validar accionabilidad.")
    add_bullet(doc, "Verificación de Gama sobre restricciones de contratos de imagen vigentes.")
    add_bullet(doc, "Verificación de compatibilidad con plataforma de marca global (si aplica).")


def build_sec_7_4_lecturas(doc):
    add_h2(doc, "7.4  Dos lecturas válidas — decisión Owner CO-2")

    add_callout_box(doc,
        "Gama 2026 tiene dos lecturas válidas — y la decisión sobre cuál enfatizar en la "
        "presentación a la Junta de Gama es estratégica, no estadística. Ambas lecturas se "
        "sostienen con los mismos datos. La decisión del Owner: presentar AMBAS lecturas "
        "lado a lado con datos de soporte, para que la Junta decida cuál pesar más.")

    add_h3(doc, "Lectura defensiva — \"Liderazgo defensivo sólido\"")

    add_para(doc, "\"0/8 métricas del embudo de Gama variaron significativamente. Gama es la "
                  "única gran cadena del mercado que mantuvo posición en un año de movimiento. "
                  "El activo relacional (OR=5.73) existe y es robusto. La posición en C+/C "
                  "(TOM 60.6%) es la más fuerte del mercado. El servicio 24h es una ventaja "
                  "funcional sin equivalente competitivo documentado.\"",
             italic=True)
    add_para(doc, "Datos de soporte: todos los datos son sig 99% o alta certeza cuantitativa.", size=10, color=GRAY_MID)

    add_h3(doc, "Lectura ofensiva — \"Señal de alerta competitiva\"")

    add_para(doc, "\"Rio +17pp TOM (sig 99%). Páramo +12pp TOM (sig 99%). Posiblemente +25pp "
                  "y +22pp en C+/C — el segmento natural de Gama. La conversión TOM→Preferida "
                  "es 44.3%→8.0%. El 92% que no prefiere a Gama no está siendo activado. El "
                  "anzuelo comunicacional está en el territorio equivocado (OR=1.03 NS).\"",
             italic=True)
    add_para(doc, "Datos de soporte: WoW C+/C como hipótesis exploratoria (CV-WOW-005).", size=10, color=GRAY_MID)

    add_para(doc, "Ambas lecturas son correctas estadísticamente. La diferencia es de tono y de "
                  "implicación estratégica:", bold=True)
    add_bullet(doc, "La lectura defensiva genera tranquilidad y validación de la estrategia "
                    "actual de Gama.")
    add_bullet(doc, "La lectura ofensiva genera urgencia y abre conversación sobre acciones "
                    "concretas.")
    add_bullet(doc, "La elección depende del estado actual de la conversación entre Gama y la "
                    "consultora y de qué tipo de respuesta organizacional se busca generar.")


# ============================================================================
# PARTE V — AGENDA 2027
# ============================================================================

def build_parte5_header(doc):
    add_h1(doc, "PARTE V — Agenda Ola 2027")


def build_sec_8_agenda(doc):
    add_h2(doc, "8.1  Hipótesis abiertas que requieren ola 2027 para confirmación")

    headers = ["Hipótesis", "Evidencia actual", "Módulo recomendado para ola 2027"]
    rows = [
        ["Imagen experiencial en ascenso (+6.3pp 'atractiva', +5.8pp 'seguro')",
         "Solo tendencia p_adj 0.058-0.089",
         "Repetir batería de atributos + corrección BH-FDR con muestra ampliada"],
        ["Debilitamiento imagen precio Gama (-3.9pp, -5.1pp)",
         "Solo tendencia p_adj 0.053-0.091",
         "Repetir P31-P34 + comparar con Rio/Páramo a nivel atributo"],
        ["Gap comunicacional PTL",
         "Sin confirmación directa con estímulos visuales",
         "Módulo publicidad con estímulos de campaña + benchmark Rio + ad recognition"],
        ["Arquetipo femenino accionable",
         "Hipótesis cualitativa emergente 6/7 FG",
         "Concept testing cuantitativo con target"],
        ["WoW en NSE C+/C",
         "Exploratorio, no pre-registrado en 2025",
         "Pre-registrar en diseño 2027 + booster muestra C+/C n=100+"],
        ["Hipótesis Express formato",
         "Compatible con datos pero variable de formato no disponible",
         "Cruzar base operacional Gama (sucursales × formato) con georreferenciación"],
        ["Mecanismo \"acompañamiento-guía\" del driver atención",
         "6/7 FG cualitativo emergente",
         "Claim testing cuantitativo de mensajes alternativos"],
        ["Barrera identitaria \"sifrinaje\"",
         "1 verbatim cualitativo central + patrón consistente",
         "Cuestionario de barreras estructurado con escala identidad vs cálculo"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 5, 6])

    add_h2(doc, "8.2  5 mejoras metodológicas obligatorias para ola 2027")

    headers = ["Mejora", "Justificación"]
    rows = [
        ["MaxDiff reemplaza Likert saturado (P22)",
         "Eliminar saturación T2B >90% en atributos que impide discriminación entre marcas"],
        ["NPS + switching explícito (3 preguntas)",
         "Medir fidelidad activa, no solo preferencia declarada"],
        ["CEPs expandidos (15-20 ocasiones vs 5 misiones genéricas)",
         "Capturar fragmentación de misiones documentada en focus groups"],
        ["Penetración 12m + frecuencia + ticket promedio",
         "Construir modelo de share of wallet que V6 no puede calcular"],
        ["Booster campo Pref-Gama n=80 + módulo comunicación con estímulos + Rio como benchmark",
         "Ampliar base referencial (actualmente n=32) + medir recall con estímulos de campaña + comparar con líder competitivo"],
    ]
    add_table(doc, headers, rows, col_widths=[7, 9])

    add_h3(doc, "8.3  Horizonte 2028 (inversión mayor, opcional)")

    add_bullet(doc, "Conjoint/DCE para medir la barrera de precio de forma experimental (USD 10-20K).")
    add_bullet(doc, "Distinctive Brand Assets battery básica (colores, logo, slogan sin marca) para "
                    "medir activos de identidad visual.")
    add_bullet(doc, "Van Westendorp PSM en categorías KVI (lácteos/charcutería, carne/pollo, "
                    "granos/enlatados) para calcular rangos de precio aceptable por categoría.")


# ============================================================================
# COMPARATIVO V1-V5
# ============================================================================

def build_sec_9_comparativo(doc):
    add_h1(doc, "9. Análisis comparativo V1-V5 + lecciones aprendidas")

    add_para(doc, "Esta sección documenta la evolución del análisis a lo largo de las 5 "
                  "versiones previas del estudio, los aprendizajes aplicados en V6, y los "
                  "patrones de feedback del Owner que estructuraron las decisiones de "
                  "metodología y presentación.")

    add_h2(doc, "9.1  Tabla comparativa de versiones")

    headers = ["Versión", "Fecha", "Eje narrativo", "Aporte único", "Limitación"]
    rows = [
        ["V1", "Pre-2026-05-12", "Tablas Cora — sin narrativa",
         "Datos crudos disponibles", "Sin estructura, sin recomendaciones"],
        ["V2", "2026-05-16", "Drivers básicos + segmentos",
         "Drivers logit + k-means inicial",
         "Sin cualitativo, sin Plan Suárez, sin precios profundos"],
        ["V3", "2026-05-17", "Activo real + Comunicación equivocada + Conversión perdida",
         "Cualitativo IN-4 triangulación + Plan Suárez C+/C + Argumentos MECE",
         "Sin WoW longitudinal"],
        ["V4", "2026-05-17", "Capa complementaria a V3 (WoW + cuali profundo)",
         "Wave-over-wave 2025-2026 + corpus FG completo + triangulación V4",
         "\"Triangulado\" co-iguala cualitativo con cuantitativo"],
        ["V5", "2026-05-18", "Fresh start — activo + comunicación + activación",
         "Iconografía explícita + C+/C central + 9 charts + cuali como insight",
         "Perdió Plan Suárez C+/C, profundidad precios, DTLS"],
        ["V6", "2026-05-18", "Consolidación con todas las preguntas + tema precios profundo",
         "CU-8 (12 preguntas faltantes) + P33×P21 + capítulos NEW-3/4/Express + comparativo V1-V5",
         "Word-first (PPTX vendrá después)"],
    ]
    add_table(doc, headers, rows, col_widths=[1.5, 2.5, 3.5, 4.5, 4])

    add_h2(doc, "9.2  Lo que cada versión aportó y lo que perdió")

    headers = ["Transición", "Aportó", "Perdió"]
    rows = [
        ["V2 → V3", "Cualitativo + Plan Suárez C+/C + triangulación", "—"],
        ["V3 → V4", "WoW longitudinal + cualitativo profundo Sinta + segmentos confirmados", "—"],
        ["V4 → V5", "Cualitativo downgrade a insight + C+/C central + 9 charts",
         "Plan Suárez C+/C, profundidad precios, DTLS análisis, estilo presentable"],
        ["V5 → V6", "Las 9 preguntas faltantes (P26, P30, P32, P40-P42, etc.) + P33×P21 + "
                    "capítulos nuevos + comparativo + recomendaciones limpias",
         "Nada (V6 es superconjunto de V5 corregido)"],
    ]
    add_table(doc, headers, rows, col_widths=[3, 7, 7])

    add_h2(doc, "9.3  Cómo V6 retroalimenta las versiones anteriores")

    add_bullet(doc, "Recupera de V3: Plan Suárez en C+/C, estilo ejecutivo, tabla \"precio "
                    "como driver\" con OR=1.03 destacado.")
    add_bullet(doc, "Mantiene de V4: capa WoW completa con tests, corpus FG completo.")
    add_bullet(doc, "Mantiene de V5: iconografía clara, cualitativo como insight separado, "
                    "C+/C como capítulo central, charts inline.")
    add_bullet(doc, "Agrega como único de V6: CU-8 (9 preguntas faltantes), P33×P21 gradiente, "
                    "subset shoppers favorables, balance óptimo de categorías, DTLS análisis "
                    "completo, validación parcial hipótesis Express, comparativo histórico.")
    add_bullet(doc, "Saca de versiones previas: recomendaciones digitales (CO-3 — fuera de scope), "
                    "señal Rio en proteínas (CO-1 — sin sustento cuantitativo).")

    add_h2(doc, "9.4  Patrones de feedback del Owner (5 patrones identificados)")

    add_h3(doc, "Patrón 1 — Simplicidad y claridad")
    add_para(doc, "El Owner prefiere redacción sencilla (\"por sentido común\") sobre técnica. "
                  "Insiste en que caveats sean explícitos, no ocultos. Evalúa accionabilidad "
                  "antes de incluir contenido en deliverables. Cita literal de transcripción "
                  "de reunión 27-04-2026: \"Mientras más sencillo mejor. Si yo me pongo a "
                  "redactar algo súper complicado, pasa demasiado largo.\"")

    add_h3(doc, "Patrón 2 — Epistemología rigurosa")
    add_para(doc, "El Owner exige convergencia de métodos antes de declarar conclusión. "
                  "Diferencia explícitamente entre conclusiones (cuantitativo) e hipótesis "
                  "(cualitativo). Pide etiquetado explícito de certeza. Este patrón llevó al "
                  "reposicionamiento en V5 (cuali como insight, no triangulado), profundizado "
                  "en V6.")

    add_h3(doc, "Patrón 3 — Honestidad técnica")
    add_para(doc, "El Owner documenta bugs/limitaciones abiertamente. Prefiere transparencia "
                  "sobre embellecer datos. Ofrece trabajo futuro como solución a limitaciones "
                  "actuales (ola 2027). En V6 esto se manifiesta en la documentación explícita "
                  "de las preguntas que no se pudieron procesar y los caveats acumulados en "
                  "cada sección.")

    add_h3(doc, "Patrón 4 — Decisión sobre entrega")
    add_para(doc, "El Owner ofrece opciones (A/B) en lugar de imponer. Pide confirmación antes "
                  "de ejecuciones costosas. Define él la decisión final pero después de ofrecer "
                  "alternativas. Las decisiones CO-1/CO-2/CO-3 en V5 y V6 son ejemplos directos: "
                  "el Owner deja explícitas las decisiones tomadas y la lógica de cada una.")

    add_h3(doc, "Patrón 5 — Estructura narrativa")
    add_para(doc, "V1-V3: sumativa (agrega capas sobre V2). V4: complementaria (ancla en V3). "
                  "V5: fresh start (2026 como eje — ruptura deliberada). V6: consolidación "
                  "(integra todo el material disponible + corrige omisiones de V5 + agrega "
                  "preguntas faltantes). Cada cambio de versión refleja una decisión "
                  "consciente sobre cómo posicionar el análisis.")

    add_para(doc, "Síntesis del patrón Owner aplicado a V6:", bold=True)
    add_para(doc, "V6 modela que el rigor y la simplicidad no son opuestos. Es más extenso que "
                  "V5 (porque incluye las preguntas faltantes y los análisis nuevos pedidos por "
                  "Cora), pero más claro (porque distingue explícitamente conclusiones de "
                  "hipótesis, organiza el tema precios en un capítulo dedicado en vez de "
                  "fragmentarlo, y muestra los datos en tablas visibles en vez de solo "
                  "conclusiones interpretativas).",
             italic=True)


# ============================================================================
# ANEXOS
# ============================================================================

def build_sec_10_anexo_meto(doc):
    add_h1(doc, "10. Anexo metodológico")

    add_para(doc, "Este anexo consolida la metodología completa del estudio. Cada sección "
                  "incluye además su nota metodológica específica al final, dentro del "
                  "cuerpo del documento.")

    add_h2(doc, "10.1  Diseño del estudio (ME-1, ME-3, ME-5)")

    add_bullet(doc, "Universo: shoppers regulares (≥1 compra/mes) en supermercados de cadena, "
                    "Caracas + Altos Mirandinos.")
    add_bullet(doc, "Muestra principal 2026: n=402 entrevistas face-to-face. Margen de error "
                    "±4.89% al 95% (Wilson).")
    add_bullet(doc, "Muestra referencia 2025: n=785.")
    add_bullet(doc, "Cuotas: NSE (C+/C / D / E), género, edad, geografía (Baruta, Libertador, "
                    "Sucre, Chacao, El Hatillo, Altos Mirandinos).")
    add_bullet(doc, "Cuestionario: 8 bloques (perfil, embudo, atributos, comportamiento, "
                    "precio, publicidad, El Recreo). 45 preguntas en total (PF1-PF10 perfil + "
                    "P16-P42 cuerpo + P43-P45 anexo).")

    add_h2(doc, "10.2  Métodos estadísticos (CU-1 a CU-8)")

    add_bullet(doc, "Embudo: tablas de frecuencias + Newcombe-Wilson para IC95.")
    add_bullet(doc, "Drivers de preferencia: regresión logística (statsmodels) sobre "
                    "Preferida=Gama (1/0) con 10 atributos de P22 como predictores. Pseudo R² = "
                    "0.4371, AUC = 0.929.")
    add_bullet(doc, "Random Forest (sklearn) + SHAP (shap) para validar contribución no-lineal "
                    "de cada feature.")
    add_bullet(doc, "K-means (sklearn) k=3 con silhouette validado contra k=2..6 (máximo en "
                    "k=3) + BayesianGaussianMixture (BIC mínimo en k=3) como confirmación.")
    add_bullet(doc, "Z-test Newcombe-Wilson para diferencias de proporciones (preferido sobre "
                    "z-test clásico para proporciones extremas).")
    add_bullet(doc, "Corrección Benjamini-Hochberg FDR (q<0.05) para múltiples tests.")
    add_bullet(doc, "Wave-over-wave 2025-2026: 57 ítems comparables. Test de comparabilidad "
                    "muestral previo (chi-cuadrado de homogeneidad NSE/geo/género/edad) — "
                    "muestras consideradas suficientemente comparables aunque con caveats.")

    add_h2(doc, "10.3  Análisis cualitativo (IN-1 a IN-8 + IN-7/8 v2)")

    add_bullet(doc, "Corpus: 12 documentos de focus groups Gama (8 sesiones presenciales + 4 "
                    "online), ~42.094 caracteres totales, 6 segmentos cubiertos.")
    add_bullet(doc, "Análisis temático según protocolo Braun & Clarke (2006).")
    add_bullet(doc, "Codificación deductivo-inductiva: (1) deductiva sobre framework cuantitativo "
                    "V3 (9 categorías a priori); (2) inductiva sobre patrones emergentes.")
    add_bullet(doc, "52 códigos totales. 6 temas (Nivel 2). 3 narrativas overarching (Nivel 3).")
    add_bullet(doc, "Single-analyst: sin coeficiente Kappa inter-codificador (limitación "
                    "documentada — práctica estándar en investigación aplicada cuando recursos "
                    "no permiten doble codificación).")
    add_bullet(doc, "Posicionamiento epistémico: insights como hipótesis, no como conclusiones "
                    "co-iguales al cuantitativo (decisión V5/V6).")

    add_h2(doc, "10.4  Gate de Bruna (gobernanza de claims)")

    add_para(doc, "Todos los claims publicables del documento pasaron por el gate de la "
                  "responsable de gobernanza de riesgos/claims (Bruna) en la revisión BR-2 V4 "
                  "el 2026-05-17. La revisión generó:")
    add_bullet(doc, "Decisiones de wording (DW-1 a DW-5): wording obligatorio para 5 claims "
                    "sensibles (verbatim Azahara con caveat literal, arquetipo femenino con "
                    "lenguaje condicional, downgrade gap comunicacional, etc.).")
    add_bullet(doc, "Caveats literales obligatorios para incorporar en el deck público (12 "
                    "strings exactos).")
    add_bullet(doc, "Decisiones condicionadas al Owner (CO-1, CO-2, CO-3) — pendientes en V4, "
                    "resueltas por Owner el 2026-05-18 y aplicadas en V6.")
    add_para(doc, "El gate de Bruna BR-2 V4 sigue vigente para V6 — los datos no cambian, "
                  "los wordings DW-1..DW-5 se aplican.", italic=True)


def build_sec_11_anexo_recreo(doc):
    add_h1(doc, "11. Anexo táctico — P43-P45 El Recreo")

    add_callout_box(doc,
        "Esta sección es ANEXA al estudio principal y se incluye por solicitud del Owner "
        "como verificación táctica. Las preguntas P43-P45 aplican exclusivamente a "
        "respondientes de la parroquia El Recreo (n=21), zona donde Gama tiene sucursal "
        "activa. NO forma parte del análisis general del estudio porque la base es muy baja "
        "y los hallazgos no son proyectables al mercado total. Se presenta de forma "
        "descriptiva solo.", color=ORANGE_W)

    add_h2(doc, "11.1  P44 — Percepción de Gama vs cadenas de la zona (n=21)")

    headers = ["Percepción", "n", "%"]
    rows = [
        ["Igual que las cadenas de la zona", "16", "76.2%"],
        ["Mucho mejor", "2", "9.5%"],
        ["Mejor", "1", "4.8%"],
        ["Mucho peor", "2", "9.5%"],
    ]
    add_table(doc, headers, rows, col_widths=[7, 2, 2])

    add_para(doc, "Lectura descriptiva: la paridad perceptual (76%) es el hallazgo dominante. "
                  "Gama no tiene imagen de superioridad clara ni inferioridad grave en su "
                  "zona de operación actual.", italic=True)

    add_h2(doc, "11.2  P45 — Disposición a comprar si Gama abre sucursal (n=21)")

    headers = ["Nivel", "n", "%"]
    rows = [
        ["Muy Dispuesto", "3", "14.3%"],
        ["Dispuesto", "10", "47.6%"],
        ["Algo dispuesto", "5", "23.8%"],
        ["Poco dispuesto", "2", "9.5%"],
        ["Nada dispuesto", "1", "4.8%"],
        ["Alta disposición (Muy + Dispuesto)", "13", "61.9%"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 2, 2])

    add_para(doc, "Lectura descriptiva: 62% de los residentes de El Recreo tendría alta "
                  "disposición a comprar en Gama si se abre más cerca. Coherente con el "
                  "hecho de que Gama ya opera ahí.", italic=True)

    add_caveat_box(doc,
        "BLOQUE ANEXO — base n=21 (solo El Recreo). Datos no proyectables al mercado total. "
        "Cualquier interpretación es estrictamente indicativa. Se presenta solo para "
        "verificación táctica del Owner. NO incorporar a deck público sin caveat literal "
        "explícito. (BLOCKER-CU8-P44/P45 del CU-8 v1)")

    add_para(doc, "")
    add_separator(doc)
    add_para(doc, "FIN DEL DOCUMENTO V6 — Notoriedad Gama 2026",
             bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=RED_GAMA)
    add_para(doc, "Cora Urrea + Raoul Bermúdez  ·  2026-05-18  ·  Confidencial NDA",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=GRAY_MID)


# ============================================================================
# MAIN
# ============================================================================

def main():
    print("Iniciando construccion del Word V6...")
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = GRAY_DARK

    print("  Portada...")
    build_portada(doc)
    print("  Indice...")
    build_toc(doc)
    print("  Resumen ejecutivo...")
    build_executive_summary(doc)
    print("  Marco V6...")
    build_marco(doc)
    print("  Ficha tecnica...")
    build_ficha_tecnica(doc)

    print("  PARTE I - Cuantitativo...")
    build_parte1_header(doc)
    build_sec_4_1_cpc(doc)
    build_sec_4_2_salud_gama(doc)
    build_sec_4_3_posicion(doc)
    build_sec_4_4_drivers(doc)
    build_sec_4_5_precios(doc)
    build_sec_4_6_lugar_precio_preferencia(doc)
    build_sec_4_7_balance(doc)
    build_sec_4_8_subset(doc)
    build_sec_4_9_misiones(doc)
    build_sec_4_10_comunicacion(doc)
    build_sec_4_11_wow(doc)

    print("  PARTE II - Cualitativo...")
    build_parte2_header(doc)
    build_sec_5_regla(doc)
    build_sec_5_2_acompanamiento(doc)
    build_sec_5_3_sifrinaje(doc)
    build_sec_5_4_digital(doc)
    build_sec_5_5_arquetipo(doc)

    print("  PARTE III - Hipotesis especiales...")
    build_parte3_header(doc)
    build_sec_6_express(doc)

    print("  PARTE IV - Sintesis...")
    build_parte4_header(doc)
    build_sec_7_1_tesis(doc)
    build_sec_7_2_argumentos(doc)
    build_sec_7_3_recomendaciones(doc)
    build_sec_7_4_lecturas(doc)

    print("  PARTE V - Agenda 2027...")
    build_parte5_header(doc)
    build_sec_8_agenda(doc)

    print("  Seccion 9 - Comparativo V1-V5...")
    build_sec_9_comparativo(doc)

    print("  Seccion 10 - Anexo metodologico...")
    build_sec_10_anexo_meto(doc)

    print("  Seccion 11 - Anexo P43-P45 El Recreo...")
    build_sec_11_anexo_recreo(doc)

    doc.save(OUTPUT)
    print(f"\nGuardado: {OUTPUT}")


if __name__ == '__main__':
    main()
