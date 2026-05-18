"""Documento metodológico V8 — Notoriedad Gama 2026
Glossary/manual de los análisis estadísticos y prácticas de IM usados en el estudio.
Para Cora (instrucción) + cliente (referencia educativa).
Lenguaje profesional NO técnico estadístico.
V8 agrega: §13 Top Box puro vs Top2Box + §14 Regresión logística + §15 Mundo de marca P23.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V8\2026-05-18_Notoriedad-Gama-2026_V8_Documento-Metodologico.docx"

# Paleta Gama corregida V8 — per brand-kit.md V0.1 (rojo saturado #E30613, NO vino tinto)
RED_GAMA  = RGBColor(0xE3, 0x06, 0x13)
RED_DARK  = RGBColor(0xB8, 0x05, 0x10)
GRAY_DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_MID  = RGBColor(0x6B, 0x6B, 0x6B)
BLUE_INFO = RGBColor(0x4A, 0x6F, 0xA5)
ORANGE_W  = RGBColor(0xF2, 0xA9, 0x00)
GREEN_OK  = RGBColor(0x2D, 0x8F, 0x47)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color='808080', size='4'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_para(doc, text, size=11, bold=False, italic=False, color=GRAY_DARK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    run.font.size = Pt(20); run.bold = True
    run.font.color.rgb = RED_GAMA


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(15); run.bold = True
    run.font.color.rgb = RED_DARK


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(12); run.bold = True
    run.font.color.rgb = GRAY_DARK


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11); run.bold = True; run.italic = True
    run.font.color.rgb = BLUE_INFO


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("•  ")
    run.font.size = Pt(size); run.bold = True; run.font.color.rgb = RED_GAMA
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = GRAY_DARK


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, '7A1212')
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h)); run.bold = True
        run.font.size = Pt(10); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            set_cell_borders(cell)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'F5F5F5')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()


def add_callout(doc, text, color=RED_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("★ ")
    run.font.size = Pt(11); run.bold = True; run.font.color.rgb = color
    run = p.add_run(text)
    run.font.size = Pt(11); run.bold = True; run.font.color.rgb = color


# ========== ESTRUCTURA ==========

def build_portada(doc):
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "DOCUMENTO METODOLÓGICO V7",
             size=22, bold=True, color=RED_GAMA, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Notoriedad y Preferencia de Marca — Gama 2026",
             size=18, bold=True, color=RED_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, "Explicación de los análisis estadísticos y prácticas de investigación de mercado utilizados en el informe principal",
             size=12, italic=True, color=GRAY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "Para uso del equipo analítico y como referencia educativa al cliente",
             size=11, italic=True, color=GRAY_MID, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        doc.add_paragraph()
    metadata = [
        ("Fecha", "2026-05-18"),
        ("Versión", "V7 — acompaña al Word V7 principal y al PPTX V7"),
        ("Audiencia primaria", "Cora Urrea (consultoría)"),
        ("Audiencia secundaria", "Cliente Gama (cuando solicite profundidad metodológica)"),
        ("Estilo", "Profesional, NO técnico estadístico"),
        ("Cómo usar", "Como glosario — leer entrada por entrada según necesidad"),
    ]
    t = doc.add_table(rows=len(metadata), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(metadata):
        c1 = t.rows[i].cells[0]; c2 = t.rows[i].cells[1]
        c1.width = Cm(5); c2.width = Cm(11)
        p = c1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(k + ":"); run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RED_DARK
        p = c2.paragraphs[0]
        run = p.add_run(v); run.font.size = Pt(11)


def build_intro(doc):
    add_h1(doc, "1. Cómo usar este documento")
    add_para(doc, "Este documento explica, en lenguaje profesional pero accesible, los análisis estadísticos y las "
                  "prácticas de investigación de mercado utilizadas en el informe Notoriedad Gama 2026 V7. Está "
                  "diseñado para dos audiencias:")
    add_bullet(doc, "Cora Urrea (consultora a cargo del proyecto): para instruirse sobre cada análisis y poder "
                    "explicarle al cliente cuando pregunte.")
    add_bullet(doc, "Cliente Gama: como referencia educativa cuando quiera profundidad metodológica sobre un punto "
                    "específico del informe principal.")
    add_para(doc, "Estructura del documento:", bold=True, space_before=10)
    add_bullet(doc, "Sección 2 — Diseño del estudio y muestra")
    add_bullet(doc, "Sección 3 — Análisis de proporciones e intervalos de confianza")
    add_bullet(doc, "Sección 4 — Tests de diferencia y comparaciones múltiples")
    add_bullet(doc, "Sección 5 — Modelos predictivos (logit, RF, SHAP)")
    add_bullet(doc, "Sección 6 — Segmentación (K-means)")
    add_bullet(doc, "Sección 7 — Análisis perceptual (DNA z-scores, MDS)")
    add_bullet(doc, "Sección 8 — Wave-over-wave (comparativo 2025-2026)")
    add_bullet(doc, "Sección 9 — Análisis cualitativo (Braun & Clarke)")
    add_bullet(doc, "Sección 10 — Conceptos transversales (referencial, sig, IC95, etc.)")


def build_diseno(doc):
    add_h1(doc, "2. Diseño del estudio y muestra")

    add_h2(doc, "2.1 Universo y muestra")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "El universo es el conjunto total de personas sobre el cual queremos sacar conclusiones. La muestra "
                  "es el subconjunto que efectivamente entrevistamos. La calidad de las conclusiones depende de que "
                  "la muestra sea representativa del universo.")
    add_para(doc, "Universo de este estudio:", bold=True)
    add_para(doc, "Shoppers regulares (≥1 compra mensual) en supermercados de cadena en Caracas + Altos Mirandinos.")
    add_para(doc, "Muestra:", bold=True)
    add_bullet(doc, "n=402 entrevistas face-to-face en 2026.")
    add_bullet(doc, "n=785 referencia 2025 (con caveats de comparabilidad).")
    add_bullet(doc, "Cuotas controladas: NSE (C+/C, D, E), género, edad, municipio.")

    add_h2(doc, "2.2 Margen de error y nivel de confianza")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "El margen de error indica cuánto pueden diferir los resultados de la muestra respecto al universo "
                  "real. El nivel de confianza (95%) indica que en 95 de cada 100 estudios similares, los resultados "
                  "estarían dentro de ese margen.")
    add_para(doc, "En este estudio:", bold=True)
    headers = ["Subconjunto", "n", "Margen de error al 95%"]
    rows = [
        ["Total", "402", "±4.89%"],
        ["NSE C+/C (clase media)", "104", "±9.8%"],
        ["NSE D", "127", "±8.7%"],
        ["NSE E", "171", "±7.5%"],
        ["Pref-Gama (REFERENCIAL)", "32", "±17 pp"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 3, 5])
    add_para(doc, "Cuándo aparece en el informe:", bold=True)
    add_para(doc, "Cada cifra reportada lleva implícito su margen de error según el tamaño de la base usada para calcularla. "
                  "Las cifras de Pref-Gama (n=32) están marcadas como REFERENCIAL en todo el informe — el margen es tan "
                  "amplio (±17pp) que no se pueden considerar proyectables al universo, solo descriptivos del grupo medido.")


def build_proporciones(doc):
    add_h1(doc, "3. Análisis de proporciones e intervalos de confianza")

    add_h2(doc, "3.1 Método Newcombe-Wilson")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Un método para calcular el intervalo de confianza alrededor de una proporción (por ejemplo, el 8% "
                  "que prefiere Gama). Es preferible al método clásico (Wald) porque funciona mejor cuando las "
                  "proporciones son cercanas a 0% o 100%, o cuando la muestra es pequeña.")
    add_para(doc, "Por qué se usa en este estudio:", bold=True)
    add_para(doc, "Muchas cifras del estudio son proporciones (% TOM, % preferida, % asociación, etc.). Newcombe-Wilson "
                  "produce intervalos más realistas, especialmente para los subgrupos pequeños (NSE, preferentes por marca, "
                  "recordadores publicitarios). El método clásico podría producir intervalos que se extienden fuera del "
                  "rango 0-100%, lo cual es lógicamente imposible.")
    add_para(doc, "Dónde aparece en el informe:", bold=True)
    add_bullet(doc, "Capítulo 1 — Embudo de marcas (todas las cifras de TOM, Consideración, Compra, Preferida)")
    add_bullet(doc, "Capítulo 3 — Percepción de precios (P33, P34) con IC95")
    add_bullet(doc, "Capítulo 4 — Recall publicitario (PTL, DTLS)")
    add_para(doc, "Cómo leer un IC95:", bold=True)
    add_para(doc, "Si reportamos \"% preferida Gama = 8.0% [IC95: 5.7%, 11.0%]\", quiere decir que la mejor estimación "
                  "puntual es 8.0%, pero el valor real en el universo podría estar entre 5.7% y 11.0% con 95% de confianza.")


def build_tests_diferencia(doc):
    add_h1(doc, "4. Tests de diferencia y comparaciones múltiples")

    add_h2(doc, "4.1 Z-test para diferencia de proporciones (Newcombe)")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Un test estadístico que indica si la diferencia observada entre dos proporciones (por ejemplo, "
                  "TOM Gama 2025 vs TOM Gama 2026) es genuina o podría deberse al azar de la muestra.")
    add_para(doc, "Cómo se lee:", bold=True)
    add_bullet(doc, "p-value < 0.01 — diferencia altamente significativa (***), prácticamente seguro que no es azar")
    add_bullet(doc, "p-value < 0.05 — diferencia significativa (**), confianza estándar al 95%")
    add_bullet(doc, "p-value < 0.10 — tendencia (*), señal direccional pero no concluyente")
    add_bullet(doc, "p-value > 0.10 — no significativo, la diferencia podría ser azar")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 1 — Comparativo 2025-2026 (Rio +17pp TOM sig 99%, Páramo +12pp sig 99%, Gama 0/8 sig)")
    add_bullet(doc, "Capítulo 3 — Cruce P33×P21 (chi² = 21.94, p<0.001)")

    add_h2(doc, "4.2 Corrección de comparaciones múltiples (Benjamini-Hochberg FDR)")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Cuando se hacen muchos tests estadísticos al mismo tiempo (por ejemplo, 57 tests WoW comparando "
                  "cada cifra entre 2025 y 2026), aumenta la probabilidad de obtener \"falsos positivos\" — diferencias "
                  "que parecen significativas pero son solo azar. La corrección Benjamini-Hochberg (BH-FDR) ajusta los "
                  "p-values para mantener controlada esa tasa de falsos positivos.")
    add_para(doc, "Por qué se usa:", bold=True)
    add_para(doc, "Más confiable que reportar \"hubo cambios significativos\" sin corrección. Permite afirmar con "
                  "rigor que de los 10 cambios significativos detectados WoW, todos son genuinos al 95% de confianza.")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 1 (sección 1.1) — comparativo embudo Gama 2025→2026")
    add_bullet(doc, "Anexo metodológico interno (CU-7) — todos los tests WoW con corrección BH-FDR")

    add_h2(doc, "4.3 Chi-cuadrado de independencia")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Un test que evalúa si dos variables categóricas están relacionadas o son independientes. Por "
                  "ejemplo, ¿la percepción de precio (P33: cara/igual/económica) está relacionada con la preferencia "
                  "por Gama (P21: preferida sí/no)?")
    add_para(doc, "Cómo se lee:", bold=True)
    add_para(doc, "Si el chi² produce p-value bajo (<0.05), las dos variables están asociadas. La fuerza de la "
                  "asociación se mide con la V de Cramér (0 = nula, 1 = perfecta).")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 3 (sección 3.3) — cruce P33 × P21: chi²=21.94, p<0.001, V=0.234 (asociación pequeña-moderada)")


def build_modelos(doc):
    add_h1(doc, "5. Modelos predictivos (logit, Random Forest, SHAP)")

    add_h2(doc, "5.1 Regresión logística (Logit)")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Un modelo estadístico que predice una variable binaria (sí/no, como \"prefiere Gama\") a partir "
                  "de varias variables explicativas (las asociaciones de atributos con Gama). Para cada atributo "
                  "calcula un \"Odds Ratio (OR)\" que indica cuánto cambia la probabilidad de preferir Gama si la "
                  "persona asocia ese atributo con la marca.")
    add_para(doc, "Cómo se lee el OR:", bold=True)
    add_bullet(doc, "OR > 1: el atributo aumenta la probabilidad (por ejemplo, OR=5.73 atención significa 5.7 "
                    "veces más probabilidad de preferir Gama si se asocia con atención)")
    add_bullet(doc, "OR = 1: el atributo no afecta (por ejemplo, OR=1.03 precio = el precio no predice preferencia)")
    add_bullet(doc, "OR < 1: el atributo disminuye la probabilidad")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 2 (sección 2.4) — DNA y drivers de preferencia: Atención OR=5.73***, Limpieza OR=3.99*, "
                    "Promociones OR=3.64**, Precio OR=1.03 NS")
    add_para(doc, "Calidad del modelo:", bold=True)
    add_bullet(doc, "Pseudo R² = 0.4371 (buena capacidad explicativa)")
    add_bullet(doc, "AUC = 0.929 (excelente capacidad de discriminación)")

    add_h2(doc, "5.2 Random Forest + SHAP")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Random Forest es un modelo predictivo basado en cientos de árboles de decisión. SHAP es una técnica "
                  "que descompone la predicción del modelo en la contribución específica de cada variable. Permite "
                  "validar (o cuestionar) los resultados de la regresión logística usando un método completamente "
                  "diferente.")
    add_para(doc, "Por qué se usa:", bold=True)
    add_para(doc, "Si dos métodos completamente diferentes (logit y RF/SHAP) llegan al mismo ranking de drivers, la "
                  "conclusión es robusta. En este estudio Atención es #1 en ambos métodos — es la convergencia más "
                  "fuerte del análisis.")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 2 (sección 2.4) — convergencia de 4 métodos en Atención como driver #1")


def build_segmentacion(doc):
    add_h1(doc, "6. Segmentación (K-means)")

    add_h2(doc, "6.1 K-means")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Una técnica de segmentación que agrupa a los respondentes en \"clusters\" basados en similitud de "
                  "sus respuestas en múltiples variables (importancia de atributos, percepción de precio, comportamiento "
                  "de compra). El \"k\" indica cuántos segmentos se identifican (en este estudio k=3).")
    add_para(doc, "Por qué k=3:", bold=True)
    add_para(doc, "El número óptimo de segmentos se valida con dos criterios estadísticos: silhouette score (mide qué "
                  "tan bien definido está cada cluster) y BIC (mide qué tan bien el modelo se ajusta a los datos). "
                  "En este estudio ambos criterios apuntan a k=3 como óptimo.")
    add_para(doc, "Resultados:", bold=True)
    headers = ["Segmento", "Tamaño", "Características principales"]
    rows = [
        ["Seg 1 — Mayoría Exigente", "59%", "Alta exigencia, no exclusivos Gama, misión múltiple"],
        ["Seg 2 — Pragmáticos Convertibles", "33%", "Menor resistencia precio, 0% pref-Gama actual"],
        ["Seg 3 — Núcleo Leal", "8% (n=32)", "100% pref-Gama, alta frecuencia, vínculo afectivo"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 3, 8])
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 5 (sección 5.4) — recomendaciones diferenciadas por segmento")
    add_callout(doc, "Caveat importante: el silhouette score de este estudio es ~0.20 (moderado). Esto significa que "
                     "los segmentos son tendencias interpretativas, no categorías perfectamente discretas. Los "
                     "perfiles se solapan en los márgenes — se debe leer como \"perfiles dominantes\" más que como "
                     "\"tipos cerrados\".")


def build_perceptual(doc):
    add_h1(doc, "7. Análisis perceptual (DNA z-scores, MDS)")

    add_h2(doc, "7.1 Z-scores DNA de marca")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Un z-score indica cuántas desviaciones estándar está una cadena por encima o por debajo de la "
                  "media del mercado en un atributo determinado. Es una forma de identificar las fortalezas y "
                  "debilidades relativas (\"sobreíndices\" y \"subíndices\") de cada marca vs el promedio.")
    add_para(doc, "Cómo se lee:", bold=True)
    add_bullet(doc, "z > +1.0: sobreíndice fuerte — fortaleza diferenciadora de la marca")
    add_bullet(doc, "z entre +0.5 y +1.0: sobreíndice moderado")
    add_bullet(doc, "z entre -0.5 y +0.5: neutral — atributo comoditizado")
    add_bullet(doc, "z entre -0.5 y -1.0: subíndice moderado — debilidad relativa")
    add_bullet(doc, "z < -1.0: subíndice fuerte — debilidad estructural")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 2 (sección 2.4) — DNA de Gama: 6 atributos con sobreíndice (Tienda atractiva +1.10, "
                    "Calidad +1.01, Seguro +0.84, Limpieza +0.81, Rapidez +0.62, Atención +0.52)")
    add_bullet(doc, "Capítulo 2 (sección 2.4) — DNA Gama Total vs C+/C (segmento natural ve aún más fuerte)")

    add_h2(doc, "7.2 Multidimensional Scaling (MDS)")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Una técnica que toma una matriz de muchas variables (en este estudio: 10 atributos × 10 cadenas "
                  "= 100 dimensiones) y la reduce a 2 dimensiones que se pueden visualizar como un mapa. Las marcas "
                  "que aparecen cerca son percibidas como similares; las que aparecen lejos son percibidas como "
                  "diferentes.")
    add_para(doc, "Cómo se interpreta el mapa:", bold=True)
    add_bullet(doc, "Posiciones relativas, no absolutas: lo que importa es la vecindad entre marcas")
    add_bullet(doc, "Las dimensiones del mapa no tienen un nombre predefinido — se interpretan a posteriori "
                    "(en este estudio: dimensión 1 = precio vs experiencia; dimensión 2 = surtido vs conveniencia)")
    add_bullet(doc, "Stress de Kruskal < 0.20 = representación aceptable")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 2 (sección 2.3) — 3 mapas perceptuales: Total, C+/C, Pref-Gama")


def build_wow(doc):
    add_h1(doc, "8. Wave-over-wave (comparativo 2025-2026)")

    add_h2(doc, "8.1 Qué es y por qué se hace")
    add_para(doc, "El análisis wave-over-wave compara los resultados de una ola del estudio (2026) contra los de la "
                  "ola anterior (2025) para identificar cambios significativos. Permite distinguir entre cambios "
                  "reales del mercado y variaciones aleatorias de la muestra.")
    add_para(doc, "Dónde aparece:", bold=True)
    add_bullet(doc, "Capítulo 1 (sección 1.1) — embudo Gama 2025 vs 2026")
    add_bullet(doc, "Capítulo 1 — movimiento competitivo (Rio +17pp TOM sig 99%, Páramo +12pp, CM -7.7pp compra)")

    add_h2(doc, "8.2 Caveats importantes (CV-WOW-001, CV-WOW-002, CV-WOW-005)")
    add_para(doc, "Estos caveats acompañan a todas las cifras comparativas 2025-2026 del informe:")
    add_callout(doc, "CV-WOW-001 — sin ponderación 2025: la BBDD 2025 recibida no incluía el factor de ponderación "
                     "muestral (la variable @PONDERAR_1 era todo cero). Los estimados 2025 pueden tener sesgo de diseño "
                     "si la muestra original era estratificada.", color=ORANGE_W)
    add_callout(doc, "CV-WOW-002 — composición geográfica difiere: Libertador está sobrerrepresentado en 2025 vs 2026. "
                     "Los cambios WoW deben interpretarse con precaución, particularmente en variables que correlacionan "
                     "con geografía.", color=ORANGE_W)
    add_callout(doc, "CV-WOW-005 — análisis WoW por NSE no pre-registrados en 2025: los cruces por NSE en WoW (por "
                     "ejemplo, Rio +25.8pp TOM en C+/C) deben interpretarse como hipótesis a confirmar en ola 2027, "
                     "no como conclusiones definitivas.", color=ORANGE_W)


def build_cuali(doc):
    add_h1(doc, "9. Análisis cualitativo (Braun & Clarke)")

    add_h2(doc, "9.1 Análisis temático")
    add_para(doc, "Qué es:", bold=True)
    add_para(doc, "Un método sistemático para identificar patrones (\"temas\") en datos cualitativos (entrevistas, "
                  "focus groups, respuestas abiertas). Sigue el protocolo de Braun & Clarke (2006), una de las guías "
                  "más usadas en investigación aplicada de mercado.")
    add_para(doc, "Cómo funciona:", bold=True)
    add_bullet(doc, "Codificación deductiva: aplicar categorías a priori basadas en hipótesis del estudio")
    add_bullet(doc, "Codificación inductiva: identificar patrones emergentes que no estaban en el plan original")
    add_bullet(doc, "Tres niveles: códigos (Nivel 1) → temas (Nivel 2) → narrativas overarching (Nivel 3)")
    add_para(doc, "En este estudio:", bold=True)
    add_bullet(doc, "Corpus: 12 documentos de focus groups Gama (~42.094 caracteres)")
    add_bullet(doc, "52 códigos identificados")
    add_bullet(doc, "6 temas (Nivel 2)")
    add_bullet(doc, "3 narrativas overarching (Nivel 3)")

    add_h2(doc, "9.2 Caveat importante: rol del cualitativo en este estudio")
    add_callout(doc, "Por decisión explícita de Cora (reunión 18/05), el análisis cualitativo del 2026 NO se incluye "
                     "como conclusión en el informe principal — era de otro estudio (App Gama + Gama Club), no del "
                     "tracker general. Solo se utilizan verbatims puntuales para CONFIRMAR posicionamiento cuantitativo, "
                     "nunca como conclusión propia.", color=BLUE_INFO)
    add_para(doc, "Dónde aparece en el informe principal:", bold=True)
    add_bullet(doc, "Capítulo 2 (sección 2.5) — gap verbal vs asociativo: verbatims como ilustración del fenómeno cuantitativo")

    add_h2(doc, "9.3 Limitación: analista único (sin Kappa inter-codificador)")
    add_para(doc, "El análisis cualitativo de este estudio fue realizado por un único analista. La práctica ideal "
                  "incluiría un segundo codificador independiente para calcular el coeficiente de acuerdo Kappa de "
                  "Cohen — pero es práctica estándar en investigación aplicada usar analista único cuando los recursos "
                  "no permiten doble codificación. La limitación está documentada y para ola 2027 se recomienda doble "
                  "codificación en un 20% del material.")


def build_conceptos(doc):
    add_h1(doc, "10. Conceptos transversales")

    add_h2(doc, "10.1 REFERENCIAL (base pequeña)")
    add_para(doc, "Cuando una cifra se calcula sobre una base de menos de 30 personas, se marca como REFERENCIAL. "
                  "Esto significa que el dato es indicativo de la tendencia del grupo medido pero NO es proyectable "
                  "al universo más amplio. El margen de error es demasiado grande para sacar conclusiones definitivas.")
    add_para(doc, "Ejemplos en este estudio:", bold=True)
    add_bullet(doc, "Pref-Gama (n=32): TODAS las cifras de los preferentes de Gama son referenciales (margen ±17pp).")
    add_bullet(doc, "Categorías P30/P32 para Gama (n<30 en cada categoría): tendencias, no proyectables.")
    add_bullet(doc, "Recordadores PTL (n=43) y DTLS (n=50): bases aceptables a nivel total, pero los cruces por NSE "
                    "producen n<30 (referenciales).")

    add_h2(doc, "10.2 Significancia estadística vs significancia práctica")
    add_para(doc, "Que un dato sea \"estadísticamente significativo\" significa que es muy probable que el patrón "
                  "observado sea real (no azar). PERO no necesariamente es importante estratégicamente. Por ejemplo:")
    add_bullet(doc, "Diferencia significativa pero pequeña: cambio de +0.5pp en una métrica puede ser sig al 95% si "
                    "n es grande, pero estratégicamente irrelevante.")
    add_bullet(doc, "Diferencia importante pero no significativa: cambio de +5pp en una variable con n=30 puede no "
                    "ser sig pero ser una señal direccional valiosa.")
    add_para(doc, "El criterio final es: ¿este dato cambia algo que el cliente va a decidir? Si sí, importa "
                  "aunque sea referencial. Si no, no importa aunque sea significativo.")

    add_h2(doc, "10.3 Glosario rápido")
    headers = ["Término", "Significado en una línea"]
    rows = [
        ["IC95", "Intervalo de Confianza al 95% — rango donde está el valor real con 95% probabilidad"],
        ["p-value", "Probabilidad de que el resultado observado sea por azar (menor = más significativo)"],
        ["OR (Odds Ratio)", "Cuánto cambia la probabilidad de un evento por una variable explicativa"],
        ["TOM (Top of Mind)", "Primera marca que viene a la mente cuando se pregunta espontáneamente"],
        ["T2B (Top 2 Box)", "% que elige las 2 opciones más altas de una escala (ej: Importante + Muy importante)"],
        ["NSE", "Nivel Socio-Económico (C+/C, D, E en este estudio)"],
        ["BH-FDR", "Corrección Benjamini-Hochberg para múltiples tests — controla falsos positivos"],
        ["Silhouette", "Métrica de calidad de segmentación (0-1, más alto = mejor)"],
        ["MDS", "Multidimensional Scaling — reduce muchas variables a un mapa 2D"],
        ["SHAP", "Técnica que descompone la predicción de un modelo en contribución de cada variable"],
        ["Newcombe-Wilson", "Método para calcular IC95 de proporciones (preferible a Wald clásico)"],
        ["Stress de Kruskal", "Medida de calidad de un MDS (<0.20 = aceptable)"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 12])


def build_top_box(doc):
    add_h1(doc, "12. Top Box puro vs Top2Box (TB vs T2B) [NUEVO V8]")
    add_para(doc, "¿Qué es?", bold=True)
    add_para(doc, "Cuando una pregunta usa escala de 5 puntos (1=Nada importante / 5=Muy importante), hay dos formas "
                  "habituales de reportar la importancia agregada:")
    add_bullet(doc, "**Top Box puro (TB):** solo % que respondió 5 (\"Muy importante\"). Métrica más estricta y "
                    "discriminadora — solo cuenta la intensidad máxima.")
    add_bullet(doc, "**Top2Box (T2B):** % que respondió 4 + 5 (\"Importante\" + \"Muy importante\"). Métrica más "
                    "permisiva — suaviza diferencias entre atributos porque casi todos llegan al techo.")

    add_para(doc, "¿Cuándo usar cuál?", bold=True)
    add_bullet(doc, "Si los atributos tienen poca varianza en T2B (efecto techo): usar TB puro para discriminar mejor.")
    add_bullet(doc, "Si los atributos están más distribuidos: T2B puede ser suficiente.")
    add_bullet(doc, "El TB puro tiene menor poder estadístico (n más pequeño) pero mayor capacidad discriminante "
                    "entre atributos cercanos al techo.")

    add_para(doc, "Aplicación en V8:", bold=True)
    add_para(doc, "En V7 reportamos T2B de la batería P22 (10 atributos importancia). Cora pidió en su email del "
                  "18/05 19:25 reprocesar con TB puro. El resultado reordena el top-5: Menor precio sube 6 "
                  "posiciones (de #8 a #2 con 62.4% TB) y Rapidez sube 4 (de #7 a #3 con 58.7% TB). Atención y Valer "
                  "dinero bajan 3 posiciones cada uno. El T2B suavizaba la importancia del precio porque la mayoría "
                  "responde \"importante\" sin distinguir intensidad — al filtrar solo los que dicen \"MUY importante\", "
                  "las prioridades funcionales emergen con mayor claridad.")

    add_callout(doc,
        "Buena práctica de IM: reportar AMBAS métricas en metodología (para auditoría) y usar TB puro como métrica "
        "principal cuando el T2B tiene efecto techo. La decisión metodológica de Cora fue correcta.")


def build_regresion(doc):
    add_h1(doc, "13. Regresión logística para drivers de preferencia [NUEVO V8]")
    add_para(doc, "¿Qué es?", bold=True)
    add_para(doc, "La regresión logística es un modelo estadístico que estima la probabilidad de un resultado binario "
                  "(sí/no, prefiere/no prefiere) a partir de una serie de variables predictoras (atributos, "
                  "percepciones, características del consumidor). El output principal son:")
    add_bullet(doc, "**Coeficientes (B):** efecto de cada variable sobre el logaritmo de las odds del resultado.")
    add_bullet(doc, "**Odds Ratio (OR):** versión interpretable del coeficiente. OR=2.4 significa que cuando esa "
                    "variable aumenta una unidad, las odds del resultado se multiplican por 2.4.")
    add_bullet(doc, "**P-value:** probabilidad de que el efecto observado sea ruido. P<0.05 = significativo.")
    add_bullet(doc, "**Pseudo-R² (McFadden):** medida de ajuste del modelo. >0.20 considerado bueno. <0.05 pobre.")
    add_bullet(doc, "**LLR p-value:** prueba si el modelo global es mejor que un modelo nulo. P<0.05 = modelo "
                    "estadísticamente útil.")

    add_para(doc, "¿Por qué logística y NO lineal en V8?", bold=True)
    add_para(doc, "Cora pidió \"regresión lineal\" en su email del 18/05. Aplicamos una corrección técnica: la "
                  "variable dependiente (\"prefiere Gama vs no\") es BINARIA, no continua. La regresión lineal con DV "
                  "binaria produce probabilidades fuera del rango [0,1] (predicciones imposibles) y residuos no-"
                  "normales (viola los supuestos del modelo). La regresión logística es el método correcto para DV "
                  "binaria y produce probabilidades interpretables.")

    add_para(doc, "¿Por qué solo Páramo fue viable en V8?", bold=True)
    add_para(doc, "Regla práctica: regresión logística confiable requiere ~50 observaciones del resultado positivo "
                  "(en este caso, preferentes de la cadena). Tamaños de muestra por cadena en este estudio:")
    add_bullet(doc, "Páramo: n=85 preferentes ✅ viable")
    add_bullet(doc, "Las otras 7 cadenas: n entre 28-45 ✗ insuficientes para regresión")
    add_para(doc, "Para las cadenas sin base suficiente, se reportan diferencias descriptivas de importancia entre "
                  "preferentes y no-preferentes (Cap 4.4.5 del Word V8). No es equivalente a regresión — no controla "
                  "por covariables — es indicativo.")

    add_para(doc, "VIF — multicolinealidad", bold=True)
    add_para(doc, "Antes de correr una regresión con múltiples predictores, se verifica que no estén altamente "
                  "correlacionados entre sí (multicolinealidad). El **Variance Inflation Factor (VIF)** mide esta "
                  "correlación. Umbrales:")
    add_bullet(doc, "VIF < 5: sin problema")
    add_bullet(doc, "VIF 5-10: aceptable, monitorear")
    add_bullet(doc, "VIF > 10: problemático, considerar métodos alternativos (KDA, Shapley)")
    add_para(doc, "En V8, el VIF máximo entre las 10 importancias P22 fue 1.35 — sin multicolinealidad. La regresión "
                  "logística era el método correcto, no se requirió KDA ni Shapley.")

    add_para(doc, "Resultado de la regresión Páramo V8", bold=True)
    add_para(doc, "Pseudo-R²=0.021, LLR p-value=0.574 — modelo NO significativo. Solo Mayor calidad fue marginalmente "
                  "significativo con OR=0.648 (un punto más de importancia de calidad reduce 35% las odds de preferir "
                  "Páramo — coherente: Páramo se elige por precio, no calidad).")

    add_callout(doc,
        "El hallazgo metodológico V8 más importante: cuando las importancias declaradas no discriminan (todos las "
        "marcan altas, efecto techo), la regresión logística confirma que NO predicen preferencia. El modelo conceptual "
        "correcto para este mercado es PERCEPCIÓN P23 → PREFERENCIA, no IMPORTANCIA P22 → PREFERENCIA. La campaña "
        "debe activar la percepción de atributos en Gama, no reclamar que los atributos son importantes.")


def build_mundo_marca(doc):
    add_h1(doc, "14. Mundo de marca: análisis P23 × Experiencia [NUEVO V8]")
    add_para(doc, "¿Qué es \"mundo de marca\"?", bold=True)
    add_para(doc, "El mundo de marca de una cadena es el conjunto de asociaciones perceptuales que el consumidor le "
                  "atribuye. En el cuestionario, se mide vía:")
    add_bullet(doc, "**P21 — Marca preferida:** \"¿Cuál es tu cadena preferida?\" (espontáneo, una sola).")
    add_bullet(doc, "**P21.1 — Razón de preferencia:** \"¿Por qué?\" (espontáneo, abierto multi-respuesta).")
    add_bullet(doc, "**P23 — Asociación inducida:** \"¿Qué cadenas asocias con [atributo X]?\" (cerrado, multi-cadena).")
    add_bullet(doc, "**P24 — Experiencia reciente:** \"¿En cuál cadena hiciste tu última compra?\" (una sola, hechos).")
    add_bullet(doc, "**P25/P26 — Experiencia habitual:** \"¿En cuáles cadenas compras habitualmente para X misión?\".")

    add_para(doc, "Diferencia crítica entre P21.1 y P23:", bold=True)
    add_bullet(doc, "P21.1 es ESPONTÁNEO: lo que la persona dice sin estímulo — refleja prominencia y peso mental.")
    add_bullet(doc, "P23 es INDUCIDO: pregunta cerrada con lista de cadenas — refleja imagen amplia, incluso de "
                    "cadenas que no son top-of-mind.")
    add_para(doc, "El deck V8 usa P23 como métrica de IMAGEN (qué percibe el mercado) y P21/P21.1 como métrica de "
                  "PREFERENCIA REVELADA (qué eligen). Cruzarlos permite identificar dónde la imagen sostiene la "
                  "preferencia y dónde existe brecha.")

    add_para(doc, "¿Por qué importa filtrar por experiencia?", bold=True)
    add_para(doc, "La imagen de marca en P23 puede ser MUY distinta entre quienes conocen la cadena (han comprado) "
                  "vs quienes no. Para Gama, la diferencia es dramática: solo 21.9% del mercado total la asocia con "
                  "Atención, pero 83.3% de quienes compraron recientemente en Gama sí lo hacen (+62.5 pp). Este "
                  "patrón identifica la \"barrera de visibilidad\" — el consumidor que entra valida el atributo, pero "
                  "el que no entra no lo percibe.")

    add_para(doc, "Atributo de SOMBRA — definición operacional V8", bold=True)
    add_para(doc, "El \"atributo de sombra\" de una cadena es aquel atributo donde la brecha de asociación P23 entre "
                  "preferentes y mercado total es máxima. Indica el atributo más \"sub-comunicado\" — alto entre "
                  "quienes la conocen, bajo entre quienes no. En V8:")
    add_bullet(doc, "Gama → ATENCIÓN (+62.5 pp). Este es el insight más importante de V8.")
    add_bullet(doc, "Plazas → ATENCIÓN (+68 pp). Comparte el atributo sombra con Gama — son vecinos perceptuales.")
    add_bullet(doc, "Rio → CALIDAD (+69 pp). Atributo sombra distinto a Gama — diferenciador potencial.")

    add_para(doc, "Vecindad perceptual — definición operacional V8", bold=True)
    add_para(doc, "Dos cadenas son \"perceptualmente vecinas\" si comparten: (a) estructura de top-3 atributos en P23 "
                  "total, (b) atributo de sombra, (c) perfil de debilidades. Gama y Plazas cumplen los 3 criterios. "
                  "Es la primera vez que el estudio identifica un competidor perceptual cercano a Gama — implicación "
                  "estratégica clave para la diferenciación de campaña 2027.")

    add_callout(doc,
        "Buena práctica de IM: cuando se analiza imagen de marca, SIEMPRE cruzar con experiencia. La imagen sin "
        "experiencia es expectativa cultural / publicitaria. La imagen tras experiencia es percepción real. La "
        "brecha entre ambas es la \"deuda de comunicación\" o \"sobre-prometida\" según el signo.")


def build_cierre(doc):
    add_h1(doc, "15. Para profundizar más")
    add_para(doc, "Si el cliente requiere profundidad adicional sobre algún punto específico, los siguientes documentos "
                  "técnicos internos están disponibles bajo solicitud:")
    add_bullet(doc, "ME-1 — Research design retroactivo (Methos, 2026-05-17)")
    add_bullet(doc, "ME-3 — Methodology rationale (Methos)")
    add_bullet(doc, "ME-4 — Vigilance Q2 2026 (Methos)")
    add_bullet(doc, "ME-5 — Methodology Plan V4 (Methos, 2026-05-17)")
    add_bullet(doc, "CU-1 a CU-10 — Análisis cuantitativos completos (Cuanti). CU-10 nuevo en V8.")
    add_bullet(doc, "IN-1 a IN-8 — Análisis cualitativos completos (Sinta)")
    add_bullet(doc, "BR-2 V4 — Gate de claims publicables (Bruna, 2026-05-17). Vigente para V8 (datos no cambiaron).")
    add_para(doc, "")
    add_para(doc, "—  FIN DEL DOCUMENTO METODOLÓGICO V8  —",
             bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RED_GAMA, size=12)
    add_para(doc, "Cora Urrea + Raoul Bermúdez  ·  2026-05-18  ·  Confidencial NDA",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY_MID, size=10)


def main():
    print("Iniciando documento metodologico V8...")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    print("  Portada..."); build_portada(doc)
    print("  Intro..."); build_intro(doc)
    print("  2. Diseño estudio..."); build_diseno(doc)
    print("  3. Proporciones..."); build_proporciones(doc)
    print("  4. Tests diferencia..."); build_tests_diferencia(doc)
    print("  5. Modelos predictivos..."); build_modelos(doc)
    print("  6. Segmentación..."); build_segmentacion(doc)
    print("  7. Análisis perceptual..."); build_perceptual(doc)
    print("  8. Wave-over-wave..."); build_wow(doc)
    print("  9. Cualitativo..."); build_cuali(doc)
    print("  10. Conceptos transversales..."); build_conceptos(doc)
    print("  12. Top Box puro vs T2B (NUEVO V8)..."); build_top_box(doc)
    print("  13. Regresion logistica (NUEVO V8)..."); build_regresion(doc)
    print("  14. Mundo de marca P23 (NUEVO V8)..."); build_mundo_marca(doc)
    print("  15. Cierre..."); build_cierre(doc)

    doc.save(OUTPUT)
    print(f"\nGuardado: {OUTPUT}")


if __name__ == '__main__':
    main()
