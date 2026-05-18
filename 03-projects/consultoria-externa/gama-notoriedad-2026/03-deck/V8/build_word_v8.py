"""Word V8 — Notoriedad Gama 2026
Tras revisión V7 por Cora 18/05 17:04 y 19:25 + CU-10 (TB puro + drivers + mundo de marca).
Bifurcación sistemática Total + C+/C + Pref-Gama.
8 ajustes V7 aprobados + 4 nuevos bloques V8 desde CU-10:
  - Cap 1.4: P22 Top Box puro reranking (precio +6, rapidez +4, atención -3)
  - Cap 2.5: Mundo de marca por cadena (P23) + atributo sombra Gama (Atención +62.5pp)
  - Cap 2.6: Vecindad perceptual Gama-Plazas (mismo cluster experiencial)
  - Cap 4: Drivers reformulados (imagen P23 → preferencia, no importancia → preferencia)
PALETA CORREGIDA: rojo Gama saturado (#E30613) per brand-kit.md V0.1 (vs vino tinto V7).
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V8\2026-05-18_Notoriedad-Gama-2026_V8.docx"
CHARTS_V5 = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V5\charts"
CHARTS_V6 = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V6\charts"
CHARTS_V7 = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\03-deck\V7\charts"
CHARTS_CU9 = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\02-analysis\cuanti\outputs\plots"
CHARTS_CU10 = r"C:\rAUL\03-projects\consultoria-externa\gama-notoriedad-2026\02-analysis\cuanti\outputs\plots"

# Paleta Gama corregida — per brand-kit.md V0.1 (2026-05-18)
# V7 usaba #7A1212 (vino tinto) — Cora corrigió: "rojo Gama saturado, NO vino tinto"
RED_GAMA  = RGBColor(0xE3, 0x06, 0x13)  # Primario Gama saturado vibrante
RED_DARK  = RGBColor(0xB8, 0x05, 0x10)  # Variante dark del rojo saturado, no granate
GRAY_DARK = RGBColor(0x1A, 0x1A, 0x1A)  # Negro corporativo brand-kit
GRAY_MID  = RGBColor(0x6B, 0x6B, 0x6B)  # Gris medio brand-kit
BLUE_INFO = RGBColor(0x4A, 0x6F, 0xA5)  # Páramo brand-kit (azul familia competidores)
ORANGE_W  = RGBColor(0xF2, 0xA9, 0x00)  # Ámbar alerta brand-kit
GREEN_OK  = RGBColor(0x2D, 0x8F, 0x47)  # Verde validación brand-kit
PURPLE_PEND = RGBColor(0x6A, 0x1B, 0x9A)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color='808080', size='4'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_para(doc, text, size=11, bold=False, italic=False, color=GRAY_DARK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    run.font.size = Pt(22); run.bold = True
    run.font.name = 'Calibri'; run.font.color.rgb = RED_GAMA


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16); run.bold = True
    run.font.name = 'Calibri'; run.font.color.rgb = RED_DARK


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13); run.bold = True
    run.font.name = 'Calibri'; run.font.color.rgb = GRAY_DARK


def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11); run.bold = True; run.italic = True
    run.font.name = 'Calibri'; run.font.color.rgb = BLUE_INFO


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("•  ")
    run.font.size = Pt(size); run.bold = True
    run.font.color.rgb = RED_GAMA; run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = GRAY_DARK; run.font.name = 'Calibri'


def add_table(doc, headers, rows, col_widths=None, header_color='E30613'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, header_color)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.name = 'Calibri'
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            set_cell_borders(cell)
            if r_idx % 2 == 0:
                set_cell_bg(cell, 'F5F5F5')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9); run.font.name = 'Calibri'
    doc.add_paragraph()


def add_chart(doc, chart_path, width_in=6.5, caption=None):
    if not os.path.exists(chart_path):
        add_para(doc, f"[Chart faltante: {os.path.basename(chart_path)}]", italic=True, color=ORANGE_W)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = GRAY_MID; run.font.name = 'Calibri'
    doc.add_paragraph()


def add_metodologia_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("📋 Nota metodológica  ·  ")
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = BLUE_INFO; run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(9); run.italic = True; run.font.color.rgb = GRAY_MID; run.font.name = 'Calibri'


def add_caveat_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("⚠ Caveat  ·  ")
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = ORANGE_W; run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(9); run.italic = True; run.font.color.rgb = ORANGE_W; run.font.name = 'Calibri'


def add_callout_box(doc, text, color=RED_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("★ ")
    run.font.size = Pt(11); run.bold = True; run.font.color.rgb = color
    run = p.add_run(text)
    run.font.size = Pt(11); run.bold = True; run.font.color.rgb = color; run.font.name = 'Calibri'


def add_dif_marker(doc, text):
    """Marcador purple visible para [DIF pendiente Owner/Cora]."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run("⚙ DECISIÓN PENDIENTE  ·  ")
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = PURPLE_PEND; run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(11); run.bold = True; run.font.color.rgb = PURPLE_PEND; run.font.name = 'Calibri'


def add_piramide_separator(doc, descriptivo_o_analitico):
    """Marcador de Pirámide Minto: lámina descriptiva vs analítica."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    if descriptivo_o_analitico == 'descriptivo':
        run = p.add_run("📊 DATOS — ")
        run.font.size = Pt(11); run.bold = True; run.font.color.rgb = BLUE_INFO
        run = p.add_run("Lo que muestran las cifras")
        run.font.size = Pt(11); run.bold = True; run.italic = True; run.font.color.rgb = BLUE_INFO
    else:
        run = p.add_run("🔍 ANÁLISIS — ")
        run.font.size = Pt(11); run.bold = True; run.font.color.rgb = RED_DARK
        run = p.add_run("Lo que los datos significan")
        run.font.size = Pt(11); run.bold = True; run.italic = True; run.font.color.rgb = RED_DARK


def add_verbatim(doc, quote, attribution):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'"{quote}"')
    run.italic = True; run.font.size = Pt(10); run.font.color.rgb = GRAY_DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.0)
    p2.paragraph_format.space_after = Pt(8)
    run = p2.add_run(f"— {attribution}")
    run.font.size = Pt(9); run.font.color.rgb = GRAY_MID; run.italic = True


# =================== SECCIONES ===================

def build_portada(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOTORIEDAD Y PREFERENCIA DE MARCA")
    run.font.size = Pt(26); run.bold = True; run.font.color.rgb = RED_GAMA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gama  ·  2026")
    run.font.size = Pt(32); run.bold = True; run.font.color.rgb = RED_GAMA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nVersión V8  ·  Bifurcación Total + C+/C + Pref-Gama + Top-Box puro")
    run.font.size = Pt(16); run.italic = True; run.font.color.rgb = GRAY_DARK

    for _ in range(3):
        doc.add_paragraph()

    metadata = [
        ("Fecha", "2026-05-18"),
        ("Equipo analítico", "Cora Urrea + Raoul Bermúdez"),
        ("Tipo", "Reporte de Investigación de Mercado"),
        ("Base", "n=402 (2026) — Total y NSE C+/C (clase media, n=104)"),
        ("Margen de error", "±4.89% (Total) · ±9.8% (C+/C)"),
        ("Estructura", "Bifurcación sistemática Total + C+/C + Pref-Gama"),
        ("Confidencialidad", "NDA"),
        ("Status", "V8 — TB puro + mundo de marca + drivers reformulados (Cora email 18/05 19:25)"),
    ]
    t = doc.add_table(rows=len(metadata), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(metadata):
        cell_k = t.rows[i].cells[0]
        cell_v = t.rows[i].cells[1]
        cell_k.width = Cm(5); cell_v.width = Cm(10)
        p = cell_k.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(k + ":"); run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RED_DARK
        p = cell_v.paragraphs[0]
        run = p.add_run(v); run.font.size = Pt(11); run.font.color.rgb = GRAY_DARK

    for _ in range(3):
        doc.add_paragraph()

    add_para(doc, "V7 incorpora la reunión Cora + Owner del 18/05 + las preguntas adicionales + el plan aprobado. "
                  "Contiene bifurcación sistemática Total / C+/C / Pref-Gama en cada capítulo, recupera análisis V3 perdidos "
                  "(DNA z-scores + modelo mental + 3 segmentos), incorpora capítulo nuevo de Posicionamiento + Mundo de marca, "
                  "y capítulo de perfil del recordador publicitario. Se entrega como borrador para revisión antes de "
                  "documento metodológico separado + HTML interactivo + PPTX final.",
             size=10, italic=True, color=GRAY_MID, align=WD_ALIGN_PARAGRAPH.CENTER)


def build_executive(doc):
    add_h1(doc, "1. Resumen Ejecutivo")
    add_callout_box(doc,
        "Gama tiene el DNA experiencial más fuerte y consolidado del mercado: sobreíndice en 6 atributos "
        "(vs 4 en V3), incluyendo Atención (z=+0.52) y Rapidez (+0.62). El segmento natural (C+/C) lo ve incluso "
        "mejor en lo experiencial. Rio sigue siendo precio-dominante — su crecimiento +17pp TOM se explica por "
        "cobertura física, no migración perceptual. El Top-Box puro de P22 (corrección Cora 18/05) revela que el "
        "mercado venezolano tiene prioridades funcionales más duras de lo que el T2B sugería: precio sube +6 "
        "posiciones, rapidez sube +4. Y el análisis del mundo de marca (P23 × experiencia) confirma que la "
        "ATENCIÓN es el atributo de sombra de Gama — invisible en el mercado total pero dramáticamente alto entre "
        "quienes la prefieren o han comprado en ella (+62.5 pp). La tarea de comunicación: trasladar la imagen de "
        "los leales al mercado amplio.")

    add_h3(doc, "5 hallazgos centrales del estudio 2026")
    add_para(doc, "1. DNA experiencial de Gama consolidado en 6 atributos.", bold=True, size=11)
    add_para(doc, "Los 4 atributos del V3 (Tienda atractiva +1.10, Calidad +1.01, Seguro +0.84, Limpieza +0.81) se "
                  "confirman con valores iguales o más altos. Ahora se suman Rapidez (+0.62) y Atención (+0.52). "
                  "El subíndice de precio se mantiene (-0.76) y el de \"valer dinero\" mejora levemente (-0.45 vs -0.67 en V3).", size=10)

    add_para(doc, "2. NSE C+/C ve a Gama incluso más fuerte en experiencia.", bold=True, size=11)
    add_para(doc, "En el segmento natural de Gama, la asociación con Atención (+9.8pp vs Total), Rapidez (+8.7pp) y "
                  "Seguridad (+6.9pp) es sustancialmente más alta. La imagen de Gama en C+/C es la del DNA, amplificada.", size=10)

    add_para(doc, "3. Rio sigue siendo precio-dominante — no migró desde experiencia, sigue compitiendo en precio.", bold=True, size=11)
    add_para(doc, "Hallazgo crítico de validación V3: el dato V3 reportaba 51% atención para Rio, pero la lectura "
                  "completa de la BBDD 2026 muestra que 68% de los preferentes de Rio citan precio como razón "
                  "principal. Rio creció +17pp TOM en el último año via PRECIO + cobertura física (apertura de "
                  "tiendas, hipótesis Cora compatible con datos). Implicación: Gama es la ÚNICA cadena del mercado "
                  "venezolano donde el precio NO es la razón dominante de preferencia (40.6% vs 51-84% en las otras 7).", size=10)

    add_para(doc, "4. Recall publicitario llega al segmento equivocado.", bold=True, size=11)
    add_para(doc, "Las frases PTL y DTLS son recordadas por ~12% del mercado. Pero el perfil de quien recordó es: 58% "
                  "NSE E, 23% D, solo 19% C+/C. El segmento natural de Gama está SUBREPRESENTADO entre los recordadores "
                  "(19% vs 25.9% que es C+/C en la muestra). La campaña actual está alcanzando al segmento que más "
                  "percibe a Gama como cara, no al segmento objetivo.", size=10)

    add_para(doc, "5. Categorías \"ofrecer valor\" identificadas.", bold=True, size=11)
    add_para(doc, "Congelados (hábito 8.0%, brecha precio +3.3pp), Gaseosas (7.2%, +4.0pp), Salsas y Enlatados (7.5%, "
                  "+3.0pp). En estas categorías el shopper ya elige Gama aunque no la percibe como la más económica — "
                  "evidencia directa de que la propuesta de valor funciona sin liderazgo de precio. Son territorios "
                  "donde la palanca es comunicar diferencial, no reducir precio.", size=10)

    add_para(doc, "6. Top-Box puro de P22 revela presión de precio más severa de lo que T2B sugería [NUEVO V8].", bold=True, size=11)
    add_para(doc, "La corrección metodológica de Cora (filtrar solo \"Muy Importante\" en lugar de Top2Box) reordena "
                  "el top-5: Limpieza (63.9%), Menor precio (62.4%, +6 posiciones), Rapidez caja (58.7%, +4 posiciones), "
                  "Mayor calidad (58.7%), Seguro (57.7%). El T2B suavizaba la importancia del precio. En el Núcleo Leal "
                  "Gama (n=32, REF), Rapidez es #1 con 84.4% y Menor precio cae al #9 con 43.8%, confirmando que el "
                  "leal de Gama es el segmento que MENOS exige precio del mercado. La presión de precio existe en el "
                  "mercado amplio (donde Gama tiene su brecha de captación), no en el segmento que ya prefiere Gama.", size=10)

    add_para(doc, "7. ATENCIÓN es el atributo de sombra de Gama — la palanca de comunicación principal [NUEVO V8].", bold=True, size=11)
    add_para(doc, "En la imagen del mercado total (P23), Gama es asociada con Limpieza (31.1%), Seguro (29.6%) y "
                  "Atractiva (28.9%). La Atención solo aparece en 21.9% del total — está SUBCOMUNICADA. Pero quienes "
                  "han comprado en Gama recientemente la asocian con Atención en 83.3% (+62.5 pp vs total) — la brecha "
                  "más grande del estudio. La regresión logística no encuentra que la importancia DECLARADA prediga "
                  "preferencia (pseudo-R² 0.021, no significativo) porque todos marcan los atributos como importantes. "
                  "Lo que predice preferencia es la PERCEPCIÓN de que la cadena tiene el atributo. La campaña debe "
                  "trasladar la imagen experiencial de los leales (atención + rapidez + valor) al mercado amplio.", size=10)

    add_metodologia_box(doc,
        "Resumen ejecutivo sintetiza hallazgos cuantitativos de CU-1 a CU-10 (n=402, m.e. ±4.89% al 95%). CU-10 (V8) "
        "agregó análisis Top Box puro de P22, regresión logística para drivers de preferencia (Páramo único caso "
        "viable, n=85), y mundo de marca por cadena (P23 × experiencia P24). Las cifras de Pref-Gama (n=32) son "
        "referenciales. Las cifras de Gama por categoría son referenciales (n<30 en casi todos los casos). El "
        "análisis cualitativo del 2026 (corpus App Gama + Gama Club) NO se incluye como conclusión (era de otro "
        "estudio); solo se usan verbatims para confirmar posicionamiento cuantitativo, según decisión Cora 18/05.")


def build_marco(doc):
    add_h1(doc, "2. Marco del documento V8")

    add_h3(doc, "2.1 Qué es V8 (y qué cambia vs V7)")
    add_para(doc, "V8 incorpora la corrección metodológica y el análisis adicional pedido por Cora en su email del "
                  "18/05 19:25, tras la entrega de V7. Mantiene los 8 ajustes aprobados de V7 y agrega 3 bloques "
                  "nuevos derivados del CU-10:")
    add_bullet(doc, "Top-Box puro de P22 (reemplaza Top2Box): reranking del top-5 de importancia de atributos.")
    add_bullet(doc, "Mundo de marca por cadena (P23 × Total × Experiencia P24): perfil perceptual de las 8 cadenas.")
    add_bullet(doc, "Drivers de preferencia reformulados: imagen P23 → preferencia, NO importancia P22 → preferencia. "
                    "Con evaluación de viabilidad de regresión (VIF + n por cadena) + corrida logística para Páramo "
                    "(único caso viable n=85) + análisis descriptivo para las demás.")
    add_bullet(doc, "Paleta corregida — rojo Gama saturado (#E30613) en headers y bordes, no el vino tinto del V7.")
    add_para(doc, "Lo que se mantiene de V7: bifurcación sistemática Total + C+/C + Pref-Gama, Pirámide Minto "
                  "explícita por bloque, capítulos de Posicionamiento + DNA + Perfil recordador, las 5 vías creativas "
                  "para abordar precio (Cap 5.6), y el prerrequisito metodológico de análisis de sensibilidad (Cap 5.7).",
                  size=10, italic=True, color=GRAY_MID)

    add_h3(doc, "2.2 Iconografía de certeza")
    headers = ["Ícono", "Significado"]
    rows = [
        ["✅", "Conclusión cuanti: evidencia significativa al 95%, base ≥30"],
        ["⚠", "Hipótesis apoyada: tendencia o soporte cualitativo robusto"],
        ["💡", "Insight cualitativo: hallazgo de focus group — NO proyectable"],
        ["📊", "Referencia evolutiva: dato 2025 con caveats de comparabilidad"],
        ["⚙", "DECISIÓN PENDIENTE: marcador purple para revisión Owner/Cora antes de versión final"],
    ]
    add_table(doc, headers, rows, col_widths=[2, 14])

    add_h3(doc, "2.3 Estructura de cada capítulo (Pirámide Minto)")
    add_para(doc, "Cada capítulo sustantivo se organiza en dos niveles, según pedido explícito de Cora en la reunión "
                  "18/05 (minuto 28:15):")
    add_bullet(doc, "📊 DATOS — lámina descriptiva con tablas y charts. \"Aquí están los números\".")
    add_bullet(doc, "🔍 ANÁLISIS — lámina analítica con interpretación. \"Aquí está lo que esos números significan\".")
    add_para(doc, "Esto permite al cliente (numérico, visual, exigente) ver primero los datos crudos y después "
                  "el análisis sobre ellos.", italic=True, size=10, color=GRAY_MID)

    add_h3(doc, "2.4 Lo que NO va en este documento (decisión Cora 18/05)")
    add_bullet(doc, "Cualitativo 2026 como conclusión (era de otro estudio: App Gama + Gama Club). Solo se usa para confirmar posicionamiento.")
    add_bullet(doc, "Recomendaciones digitales (Gama Club, inventario, frescos online) — fuera de scope original.")
    add_bullet(doc, "Mensajes alternativos para PTL/DTLS — Cora pidió validación, no propuesta de nuevos mensajes.")
    add_bullet(doc, "Análisis geográfico/municipios como capítulo central — queda como Anexo A.")
    add_bullet(doc, "Comparativo 2025-2026 en precios o comunicación — solo en embudo + posicionamiento (lo único realmente comparable).")
    add_bullet(doc, "Hipótesis Express formato — queda como Anexo B (compatible con datos pero variable no en BBDD).")

    add_metodologia_box(doc,
        "V7 reusa los mismos datos cuantitativos de las versiones previas (BBDD 2026 n=402, BBDD 2025 n=785). Lo "
        "que cambia es la profundidad analítica y la estructura narrativa. Los gates de Bruna BR-2 V4 siguen vigentes; "
        "los nuevos claims derivados de CU-9 requerirán revisión Bruna específica antes del PPTX final.")


def build_ficha(doc):
    add_h1(doc, "3. Ficha técnica del estudio")
    headers = ["Dimensión", "Especificación"]
    rows = [
        ["Universo", "Shoppers regulares (≥1 compra/mes) en supermercados de cadena, Caracas + Altos Mirandinos"],
        ["Muestra principal 2026", "n=402 entrevistas face-to-face"],
        ["Margen de error Total", "±4.89% al 95% (Wilson)"],
        ["NSE C+/C combinado (\"clase media\")", "n=104 (25.9%) · m.e. ±9.8%"],
        ["NSE D", "n=127 (31.6%)"],
        ["NSE E", "n=171 (42.5%)"],
        ["Pref-Gama (subgrupo, REFERENCIAL)", "n=32 (8.0%) · m.e. ±17pp"],
        ["Geografía", "Baruta (122) · Libertador (80) · Sucre (79) · Chacao (70) · El Hatillo (31) · Altos Mirandinos (20)"],
        ["Referencia 2025", "n=785 — comparable solo en embudo + posicionamiento, con caveats"],
        ["Modelos", "Logit (AUC=0.929) · RF/SHAP · K-means k=3 · Z-test BH-FDR · Newcombe-Wilson IC95 · MDS · Z-scores DNA"],
        ["Cualitativo", "12 docs FG (~42k chars) — usado SOLO para confirmar posicionamiento (no como conclusión)"],
        ["Cuestionario", "PF1-PF10 perfil · P16-P42 cuerpo · P43-P45 anexo El Recreo"],
        ["Procesado en V7", "100% de las preguntas con datos válidos + bifurcación Total/C+C/PrefGama"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 11])

    add_h3(doc, "Caveats metodológicos generales")
    add_bullet(doc, "C+/C combinado: en la BBDD no son separables C+ y C. n=104 → m.e. ±9.8% al 95%.")
    add_bullet(doc, "Sin ponderación 2025: factor @PONDERAR_1 era todo cero. Posible sesgo de diseño.")
    add_bullet(doc, "Composición geográfica difiere entre olas: Libertador sobrerrepresentado en 2025.")
    add_bullet(doc, "Pref-Gama (n=32) referencial. Todas las cifras de este subgrupo con caveats.")
    add_bullet(doc, "Análisis cualitativo single-analyst (sin Kappa). Práctica estándar en investigación aplicada.")

    add_metodologia_box(doc,
        "El documento metodológico separado (entregable adicional) contiene la explicación detallada de cada análisis "
        "estadístico utilizado (qué, por qué, en qué lugares del informe). Cora pidió ese documento para instruirse + "
        "poder explicar al cliente si pregunta.")


# =========== CAPÍTULO 1: EMBUDO + DRIVERS ESPONTÁNEOS ==========

def build_cap1(doc):
    add_h1(doc, "Capítulo 1 — Embudo de marcas + drivers de preferencia espontáneos")

    add_callout_box(doc,
        "Gama es la única gran cadena con posición estadísticamente estable entre 2025 y 2026 (0/8 indicadores del "
        "embudo con variación significativa). En C+/C su TOM es 60.6% (vs 44.3% Total) y su Preferida 13.5% (vs 8.0%). "
        "En el modelo mental, Gama se consolida como atención-dominante (78% de sus preferentes cita atención, vs 53% "
        "en V3) — el Núcleo Leal está cada vez más alineado con el DNA experiencial.")

    add_h2(doc, "1.1 Embudo de marcas")
    add_h3(doc, "📊 DATOS — Total")
    headers = ["Etapa", "Gama 2026", "Páramo 2026", "Rio 2026", "CM 2026", "Plan Suárez"]
    rows = [
        ["TOM (Top of Mind)", "44.3%", "39.1%", "45.0%", "—", "—"],
        ["Asistida", "60.2%", "—", "—", "—", "—"],
        ["Consideración", "31.8%", "—", "—", "—", "—"],
        ["Compra 3m", "17.7%", "—", "—", "—", "—"],
        ["Preferida", "8.0%", "21.1%", "10.2%", "11.2%", "—"],
        ["Habitual", "20.2%", "—", "—", "—", "—"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 3, 2, 2])
    add_chart(doc, os.path.join(CHARTS_V5, "D3_funnel_gama.png"), width_in=6.5,
              caption="Embudo Gama 2026 vs referencia 2025 — 0/8 indicadores con variación significativa.")

    add_h3(doc, "📊 DATOS — C+/C (clase media)")
    headers = ["Etapa", "Gama C+/C", "Gama Total", "Diferencial C+/C - Total"]
    rows = [
        ["TOM", "60.6%", "44.3%", "+16.3 pp"],
        ["Preferida", "13.5%", "8.0%", "+5.5 pp"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 4])

    add_h3(doc, "🔍 ANÁLISIS")
    add_para(doc, "Tres lecturas del embudo de Gama:", bold=True)
    add_bullet(doc, "Estabilidad como fortaleza: en un mercado donde Rio creció +17pp TOM y Páramo +12pp TOM "
                    "(ambos sig 99% — CU-7), Gama mantuvo posición en los 8 indicadores. La estabilidad es un "
                    "logro defensivo en mercado turbulento, no estancamiento.")
    add_bullet(doc, "Sobrerepresentación en C+/C: Gama tiene +16.3pp más TOM en su segmento natural que en el Total. "
                    "Esto valida que el DNA experiencial (que veremos en Cap. 2) está perfectamente alineado con el "
                    "perfil del cliente premium.")
    add_bullet(doc, "Conversión TOM→Preferida baja: 44.3% conoce a Gama, solo 8.0% la prefiere. Este gap es "
                    "el problema central del estudio. El Cap. 3 (drivers) y Cap. 5 (precios) explican los mecanismos.")

    add_caveat_box(doc,
        "Caveats WoW: CV-WOW-001 (sin ponderación 2025), CV-WOW-002 (composición geográfica difiere entre olas). "
        "Comparativo 2025-2026 solo aplica al embudo y posicionamiento (Cap. 1-2). No aplicar a precios ni comunicación.")

    add_h2(doc, "1.2 Posición competitiva 2026 y empate técnico en C+/C")
    add_chart(doc, os.path.join(CHARTS_V5, "D4_posicionamiento.png"), width_in=6.5,
              caption="Mapa de posicionamiento competitivo TOM vs Preferida 2026 con cambio WoW.")

    add_h3(doc, "📊 DATOS — Preferencia C+/C 2026 (n=104)")
    add_para(doc, "Cifras verificadas de la BBDD 2026 sobre preferencia (P21) en NSE C+/C:")
    headers = ["Cadena", "n preferentes C+/C", "% del segmento C+/C"]
    rows = [
        ["Páramo", "18", "17.3%"],
        ["Plan Suárez", "16", "15.4%"],
        ["Gama", "14", "13.5%"],
        ["Luz", "11", "10.6%"],
        ["Central Madeirense", "9", "8.7%"],
        ["Plazas", "9", "8.7%"],
        ["Forum", "7", "6.7%"],
        ["Rio", "7", "6.7%"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 4, 4])

    add_h3(doc, "🔍 ANÁLISIS — empate técnico en el segmento natural")
    add_callout_box(doc,
        "Las tres primeras cadenas (Páramo 17.3%, Plan Suárez 15.4%, Gama 13.5%) están en empate técnico en "
        "preferencia C+/C — los intervalos de confianza se solapan. Esto significa que ninguna tiene una ventaja "
        "perceptual decisiva sobre las otras dos. Plan Suárez es competidor relevante para Gama en C+/C, aunque "
        "está subutilizado en análisis previos.")

    add_chart(doc, os.path.join(CHARTS_V6, "C16_plan_suarez_cpc.png"), width_in=6.0,
              caption="Empate técnico en preferencia C+/C: Páramo / Plan Suárez / Gama con IC95 solapados.")

    add_para(doc, "Implicación: cualquier movimiento competitivo de las tres puede alterar el equilibrio. Pérdidas "
                  "en C+/C cuestan más que pérdidas en Total — la inversión defensiva en este segmento está justificada.")

    add_h2(doc, "1.3 Drivers de preferencia espontáneos (P21.1) por cadena — Modelo mental")
    add_h3(doc, "📊 DATOS — % preferentes que mencionan cada razón (P21.1 razón espontánea)")
    headers = ["Cadena", "n preferentes", "% mencionan PRECIO", "Modelo mental"]
    rows = [
        ["Gama", "32 (ref.)", "40.6%", "ÚNICA cadena ATENCIÓN-dominante"],
        ["Central Madeirense", "45", "51.1%", "Precio-dominante"],
        ["Forum", "44", "63.6%", "Precio-dominante"],
        ["Plazas", "37", "64.9%", "Precio-dominante"],
        ["Rio", "41", "68.3%", "Precio-dominante"],
        ["Plan Suárez", "28", "71.4%", "Precio-dominante"],
        ["Luz", "32", "75.0%", "Precio-dominante"],
        ["Páramo", "85", "84.0%", "Precio-dominante"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 2.5, 3.5, 4.5])

    add_chart(doc, os.path.join(CHARTS_V7, "V7C1_modelo_mental_rio_migrado.png"), width_in=6.5,
              caption="Modelo mental por cadena — Gama es la ÚNICA donde precio NO es la razón dominante de preferencia.")

    add_h3(doc, "🔍 ANÁLISIS — Gama outlier estructural del mercado")
    add_callout_box(doc,
        "Gama es la ÚNICA cadena del mercado venezolano donde el precio NO es la razón dominante de preferencia. "
        "40.6% de sus preferentes citan precio vs 51-84% en los otros 7 competidores. No es un fenómeno reciente — "
        "es diferenciación estructural sostenida.")

    add_para(doc, "Cuatro hallazgos del modelo mental:", bold=True)
    add_bullet(doc, "Gama outlier estructural: las otras 7 cadenas tienen >50% de sus preferentes mencionando precio. "
                    "Gama tiene 40.6% — la diferencia es categórica, no marginal.")
    add_bullet(doc, "Plan Suárez también es precio-dominante (71.4%), consistente con su posicionamiento histórico. "
                    "Su empate técnico con Gama en C+/C (15.4% preferencia) NO se basa en compartir territorio "
                    "atributivo — son perfiles muy distintos.")
    add_bullet(doc, "Rio creció +17pp TOM en el último año pero NO migró desde experiencia: sigue siendo "
                    "precio-dominante (68.3%, era 51% en V3). Su crecimiento se explica mejor por cobertura "
                    "física/apertura de tiendas que por cambio de posicionamiento perceptual.")
    add_bullet(doc, "Implicación estratégica MATIZADA: Gama mantiene un DNA experiencial diferenciado que vale "
                    "defender activamente. Pero esto NO es excluyente con abordar el precio — los no-clientes de "
                    "Gama operan mayoritariamente en el modelo mental precio-dominante, y la realidad socioeconómica "
                    "venezolana hace del precio una palanca importante para captar nuevos clientes. La estrategia "
                    "correcta combina defensa del DNA + abordaje creativo del precio (ver Cap 3 y Cap 5).")

    add_metodologia_box(doc,
        "P21.1 (razón espontánea) se calcula como % de PERSONAS (no menciones) que mencionan cada categoría en alguna "
        "de sus 16 respuestas posibles. Este es el método correcto y el mismo que se usó en V3 — los % son comparables. "
        "Bases pequeñas (CM n=45, Forum n=44, Plan Suárez n=28, etc.) son referenciales. Para Gama n=32 (8.0% de la "
        "muestra) — interpretar como tendencia, no como proyección.")

    # ============ SECCIÓN 1.4 NUEVA V8 — TB puro reranking ============
    add_h2(doc, "1.4 Importancia de atributos — Top Box puro [NUEVO V8]")
    add_piramide_separator(doc, "datos")

    add_para(doc, "A pedido explícito de Cora (email 18/05 19:25), reprocesamos la batería de importancia de "
                  "atributos (P22) usando Top Box puro — solo respondieron \"Muy Importante\" (5 en la escala 1-5). "
                  "El V7 reportaba Top2Box (4+5), que es más permisivo y suaviza diferencias entre atributos. "
                  "El TB puro revela una jerarquía más severa.", size=11)

    headers = ["Rank TB", "Atributo", "TB % Total", "T2B % (V7)", "Rank T2B", "Δ rank"]
    rows = [
        ["1", "Limpieza/orden", "63.9%", "97.5%", "1", "0"],
        ["2", "Menor precio", "62.4%", "94.0%", "8", "+6"],
        ["3", "Rapidez caja", "58.7%", "94.0%", "7", "+4"],
        ["4", "Mayor calidad", "58.7%", "97.5%", "2", "-2"],
        ["5", "Seguro", "57.7%", "96.0%", "4", "-1"],
        ["6", "Valer dinero", "57.7%", "96.5%", "3", "-3"],
        ["7", "Mayor surtido", "57.2%", "95.5%", "6", "-1"],
        ["8", "Mejor atención", "57.0%", "96.0%", "5", "-3"],
        ["9", "Promociones", "50.5%", "89.6%", "9", "0"],
        ["10", "Tienda atractiva", "42.5%", "84.3%", "10", "0"],
    ]
    add_table(doc, headers, rows, col_widths=[1.2, 3.2, 1.5, 1.5, 1.3, 1.3])
    add_chart(doc, os.path.join(CHARTS_CU10, "cu10_tb_vs_t2b_reranking.png"), width_in=6.5,
              caption="Reranking TB vs T2B — los atributos funcionales (precio, rapidez) suben; los relacionales (atención, valor) bajan.")

    add_piramide_separator(doc, "analisis")
    add_para(doc, "El T2B suavizaba la importancia del precio. Al filtrar solo \"Muy Importante\", Menor precio sube "
                  "6 posiciones y aparece como la segunda prioridad del mercado venezolano. Rapidez en caja sube 4 "
                  "posiciones — la urgencia percibida sobre velocidad es subestimada por T2B. En el otro extremo, "
                  "Mejor atención y Valer dinero bajan 3 posiciones cada uno: son ampliamente valorados pero con "
                  "intensidad moderada.", size=11)

    add_callout_box(doc,
        "El T2B inflaba artificialmente la importancia de atributos RELACIONALES (atención, valor) frente a los "
        "FUNCIONALES (precio, rapidez). El TB puro revela que el consumidor venezolano tiene prioridades "
        "funcionales más severas de lo que el T2B sugería. Esta es la lectura correcta del peso de cada atributo "
        "en la decisión de compra del mercado amplio.")

    add_h3(doc, "1.4.1 TB puro en C+/C — el segmento natural exige Atención con intensidad")
    headers = ["Rank", "Atributo", "TB % C+/C", "TB % Total", "Δ pp"]
    rows = [
        ["1", "Mejor atención", "64.4%", "57.0%", "+7.4"],
        ["2", "Limpieza/orden", "63.5%", "63.9%", "-0.4"],
        ["3", "Valer dinero", "62.5%", "57.7%", "+4.8"],
        ["4", "Rapidez caja", "61.5%", "58.7%", "+2.8"],
        ["5", "Seguro", "60.6%", "57.7%", "+2.9"],
        ["6", "Menor precio", "60.6%", "62.4%", "-1.8"],
    ]
    add_table(doc, headers, rows, col_widths=[1, 3.5, 1.7, 1.7, 1.3])
    add_para(doc, "En C+/C la Atención salta al #1 con 64.4% (vs #8 con 57.0% en Total). El segmento natural de Gama "
                  "exige Atención con intensidad superior al mercado total — confirma que el territorio relacional es "
                  "estratégicamente más valioso en C+/C de lo que el promedio general indica.", size=11)

    add_h3(doc, "1.4.2 TB puro en Pref-Gama (n=32, REFERENCIAL) — confirma la bifurcación")
    headers = ["Rank", "Atributo", "TB % Pref-Gama", "TB % Total", "Δ pp"]
    rows = [
        ["1", "Rapidez caja", "84.4% *", "58.7%", "+25.7"],
        ["2", "Mejor atención", "71.9% *", "57.0%", "+14.9"],
        ["2", "Valer dinero", "71.9% *", "57.7%", "+14.2"],
        ["4", "Limpieza/orden", "65.6% *", "63.9%", "+1.7"],
        ["5", "Mayor surtido", "62.5% *", "57.2%", "+5.3"],
        ["9", "Menor precio", "43.8% *", "62.4%", "-18.6"],
        ["10", "Promociones", "25.0% *", "50.5%", "-25.5"],
    ]
    add_table(doc, headers, rows, col_widths=[1, 3.5, 1.7, 1.5, 1.3])
    add_para(doc, "* REFERENCIAL n=32. Confirmado por T2B (96.9%) y media 4.81/5 del CU-4.", size=9, italic=True, color=GRAY_MID)
    add_para(doc, "Los preferentes de Gama son el segmento que MÁS exige Rapidez y Atención (+25.7 pp y +14.9 pp vs "
                  "Total) y el que MENOS exige Precio y Promociones (-18.6 pp y -25.5 pp). La bifurcación atención-"
                  "dominante vs precio-dominante es más extrema en TB puro que en T2B — refuerza la tesis V7 del DNA "
                  "experiencial sostenido.", size=11)

    add_chart(doc, os.path.join(CHARTS_CU10, "cu10_heatmap_tb_puro.png"), width_in=6.8,
              caption="Heatmap TB puro por marca preferida — la fila Gama destaca por exigir intensidad en rapidez+atención y baja exigencia en precio+promociones.")

    add_metodologia_box(doc,
        "Top Box puro = % que respondió 5 (\"Muy Importante\") en la escala 1-5 de P22. T2B (Top2Box) = % que "
        "respondió 4 o 5. Las dos métricas son complementarias, no contradictorias — el TB puro discrimina mejor "
        "entre atributos cuando todos están cerca del techo del T2B (efecto de aquiescencia). Cuando el deck final "
        "se construya, recomendamos reportar TB puro como métrica principal de importancia, con T2B en metodología "
        "para referencia.")


# =========== CAPÍTULO 2: POSICIONAMIENTO + MUNDO DE MARCA + DNA ==========

def build_cap2(doc):
    add_h1(doc, "Capítulo 2 — Posicionamiento, Mundo de marca y DNA competitivo")

    add_callout_box(doc,
        "Este capítulo es nuevo en V7 — recupera análisis de V3 que se habían perdido en V5/V6 y los amplifica con "
        "bifurcación Total + C+/C + Pref-Gama. Hallazgo central: el DNA de Gama se consolidó. Donde en V3 sobreíndecía "
        "en 4 atributos experienciales, ahora sobreíndice en 6. El segmento natural (C+/C) ve a la marca con un perfil "
        "aún más fuerte. El subíndice de precio se mantiene pero Valer-dinero mejora (-0.45 vs -0.67 en V3).")

    add_h2(doc, "2.1 Importancia de atributos (P22) — Total vs C+/C")
    add_h3(doc, "📊 DATOS")
    headers = ["Atributo", "T2B % Total", "T2B % C+/C", "Diferencia pp"]
    rows = [
        ["Rapidez caja", "94.0%", "97.1%", "+3.1 (único positivo)"],
        ["Menor precio", "94.0%", "93.3%", "-0.7"],
        ["Limpieza/orden", "97.5%", "96.2%", "-1.3"],
        ["Valer dinero", "96.5%", "95.2%", "-1.3"],
        ["Mejor atención", "96.0%", "95.2%", "-0.8"],
        ["Seguro", "96.0%", "95.2%", "-0.8"],
        ["Mayor calidad", "97.5%", "94.2%", "-3.3"],
        ["Mayor surtido", "95.5%", "91.3%", "-4.2"],
        ["Tienda atractiva", "84.3%", "79.8%", "-4.5"],
        ["Promociones", "89.6%", "84.6%", "-5.0"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 4])
    add_chart(doc, os.path.join(CHARTS_V7, "V7C2_p22_importancia_diff_cpc_total.png"), width_in=6.5,
              caption="Importancia de atributos C+/C - Total. Rapidez en caja es el único atributo donde C+/C es MÁS exigente.")

    add_h3(doc, "🔍 ANÁLISIS")
    add_para(doc, "Tres observaciones sobre importancia:", bold=True)
    add_bullet(doc, "Saturación generalizada: 9 de 10 atributos tienen T2B >90%. La importancia declarada en abstracto "
                    "NO discrimina entre marcas. El diferenciador real no es \"si importa\" sino \"quién lo tiene\" "
                    "(la asociación de marca — sección 2.2).")
    add_bullet(doc, "Rapidez en caja es el único atributo donde C+/C es más exigente que el Total (+3.1pp). El "
                    "segmento de mayor ingreso valora más el tiempo. Oportunidad de activación específica.")
    add_bullet(doc, "C+/C es MENOS exigente en Tienda atractiva, Promociones, Surtido y Calidad. El segmento "
                    "natural de Gama valora menos lo estético-comercial y más lo funcional-experiencial. Es coherente "
                    "con el DNA de la marca.")

    add_h2(doc, "2.2 Asociación marca × atributo (P23) — Gama y competidores")
    add_h3(doc, "📊 DATOS — Perfil Gama Total")
    headers = ["Atributo", "Gama %", "Páramo %", "Brecha Páramo-Gama"]
    rows = [
        ["Limpieza/orden", "31.1%", "34.6%", "+3.5 pp (cercano)"],
        ["Seguro", "29.6%", "36.1%", "+6.5 pp"],
        ["Tienda atractiva", "28.9%", "31.3%", "+2.4 pp (cercano)"],
        ["Mayor calidad", "26.6%", "26.6%", "0.0 pp (empate)"],
        ["Mejor atención", "21.9%", "32.3%", "+10.4 pp"],
        ["Rapidez caja", "21.1%", "30.3%", "+9.2 pp"],
        ["Mayor surtido", "17.9%", "29.4%", "+11.5 pp"],
        ["Valer dinero", "11.4%", "32.1%", "+20.7 pp"],
        ["Promociones", "9.0%", "30.6%", "+21.6 pp"],
        ["Menor precio", "7.2%", "35.1%", "+27.9 pp"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 4])
    add_chart(doc, os.path.join(CHARTS_CU9, "cu9_heatmap_p23_total.png"), width_in=6.5,
              caption="Heatmap asociación atributo × cadena (P23) Total — Páramo domina, Gama compite en experienciales.")

    add_h3(doc, "📊 DATOS — Gama C+/C vs Total")
    headers = ["Atributo", "Gama Total", "Gama C+/C", "Diferencia"]
    rows = [
        ["Atención", "21.9%", "31.7%", "+9.8 pp"],
        ["Rapidez", "21.1%", "29.8%", "+8.7 pp"],
        ["Seguro", "29.6%", "36.5%", "+6.9 pp"],
        ["Menor precio", "7.2%", "12.5%", "+5.3 pp"],
        ["Calidad", "26.6%", "29.8%", "+3.2 pp"],
        ["Tienda atractiva", "28.9%", "29.8%", "+0.9 pp"],
        ["Limpieza", "31.1%", "29.8%", "-1.3 pp"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 3])

    add_h3(doc, "🔍 ANÁLISIS")
    add_para(doc, "Hallazgos clave de asociación:", bold=True)
    add_bullet(doc, "Páramo es el outlier del mercado: supera a Gama en TODOS los atributos excepto Calidad (empate). "
                    "Las mayores brechas son en Menor precio (+27.9pp) y Promociones (+21.6pp) — exactamente los "
                    "atributos donde Gama subindexa. Páramo es estructuralmente el competidor de precio.")
    add_bullet(doc, "En atributos experienciales (limpieza, tienda atractiva, calidad) las brechas Páramo-Gama son "
                    "mucho más pequeñas (0-7pp). Ahí Gama es competitiva perceptualmente.")
    add_bullet(doc, "C+/C ve a Gama mejor que el Total en TODOS los atributos experienciales: Atención +9.8pp, "
                    "Rapidez +8.7pp, Seguro +6.9pp. El segmento natural está perfectamente alineado con el DNA. "
                    "Incluso en precio, C+/C es +5.3pp más favorable que el Total.")

    add_h2(doc, "2.3 Mapa perceptual (MDS) multi-segmento")
    add_chart(doc, os.path.join(CHARTS_CU9, "cu9_mds_Total.png"), width_in=6.0,
              caption="Mapa perceptual MDS — Total. Páramo outlier; Gama, CM, Forum agrupados en cuadrante experiencial.")
    add_chart(doc, os.path.join(CHARTS_CU9, "cu9_mds_CC.png"), width_in=6.0,
              caption="Mapa perceptual MDS — C+/C. Mismo patrón, posición de Gama más definida.")
    add_chart(doc, os.path.join(CHARTS_CU9, "cu9_mds_PrefGama.png"), width_in=6.0,
              caption="Mapa perceptual MDS — Pref-Gama (n=32, referencial). Validación del cuadrante experiencial.")

    add_para(doc, "Lectura del MDS:", bold=True)
    add_bullet(doc, "Dimensión 1: oposición precio-dominante (Páramo) vs experiencia-dominante (Gama, CM, Hiper Líder).")
    add_bullet(doc, "Dimensión 2: variantes dentro del espacio de experiencia (surtido vs conveniencia).")
    add_bullet(doc, "Gama se posiciona consistentemente en el cuadrante experiencial. Páramo en el cuadrante precio.")

    add_h2(doc, "2.4 DNA de Gama — Z-scores vs media del mercado (recuperado V3)")
    add_h3(doc, "📊 DATOS — DNA Total + comparación V3")
    headers = ["Atributo", "Z-score V7 2026", "Z-score V3 2025", "Tipo"]
    rows = [
        ["Tienda atractiva", "+1.10", "+1.09", "SOBREÍNDICE (+++)"],
        ["Mayor calidad", "+1.01", "+0.97", "SOBREÍNDICE (+++)"],
        ["Seguro", "+0.84", "+0.76", "SOBREÍNDICE (++)"],
        ["Limpieza/orden", "+0.81", "+0.72", "SOBREÍNDICE (++)"],
        ["Rapidez caja", "+0.62", "(nuevo)", "SOBREÍNDICE (+) — NUEVO V7"],
        ["Mejor atención", "+0.52", "(nuevo)", "SOBREÍNDICE (+) — NUEVO V7"],
        ["Mayor surtido", "+0.13", "neutral", "Neutral"],
        ["Valer dinero", "-0.45", "-0.67", "Subíndice leve (mejora vs V3)"],
        ["Promociones", "-0.64", "-0.67", "SUBÍNDICE (--)"],
        ["Menor precio", "-0.76", "-0.72", "SUBÍNDICE (---)"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 3, 3, 4])

    add_chart(doc, os.path.join(CHARTS_CU9, "cu9_dna_gama_total.png"), width_in=6.5,
              caption="DNA de Gama — Z-scores vs media del mercado. 6 atributos sobre-índice (4 en V3).")

    add_chart(doc, os.path.join(CHARTS_V7, "V7C6_dna_total_vs_cpc.png"), width_in=6.8,
              caption="DNA Gama Total vs C+/C — C+/C ve a Gama incluso más fuerte en atributos experienciales.")

    add_h3(doc, "🔍 ANÁLISIS — DNA consolidado y diferenciación estructural")
    add_para(doc, "Cuatro hallazgos del DNA:", bold=True)
    add_bullet(doc, "Los 4 atributos experienciales del V3 se confirman con valores iguales o mayores. Tienda "
                    "atractiva (+1.10) sigue siendo el atributo de mayor sobreíndice — Gama es la cadena más "
                    "claramente reconocida como 'tienda bonita' del mercado.")
    add_bullet(doc, "El DNA AMPLIÓ a 6 atributos: ahora también sobreíndice Rapidez (+0.62) y Atención (+0.52). "
                    "El perfil experiencial de Gama se volvió más denso y consistente.")
    add_bullet(doc, "Valer dinero mejora levemente (-0.45 vs -0.67 en V3) — sigue siendo subíndice pero la brecha "
                    "se acortó. Posible señal de que la propuesta de valor empieza a registrarse.")
    add_bullet(doc, "C+/C amplifica el DNA: el segmento natural ve a Gama más fuerte en Atención, Rapidez, Seguro y "
                    "Tienda atractiva. La brecha de precio (Menor precio) es MENOR en C+/C que en el Total — el "
                    "segmento premium no penaliza tanto a Gama por precio.")

    add_h3(doc, "Lectura estratégica MATIZADA — Gama outlier estructural, no oportunidad coyuntural")
    add_callout_box(doc,
        "El DNA de Gama no es un fenómeno reciente ni una ventaja coyuntural. La combinación de Cap 1.3 (modelo "
        "mental) + Cap 2.4 (DNA z-scores) muestra que Gama es estructuralmente diferente del mercado: única cadena "
        "donde precio NO domina la verbalización de preferencia + sobreíndice consistente en atributos experienciales. "
        "Esta diferenciación vale defenderse activamente como ventaja sostenible.",
        color=BLUE_INFO)

    add_para(doc, "NOTA IMPORTANTE — esto no significa ignorar el precio:", bold=True)
    add_para(doc, "Cora pidió explícitamente matizar esta lectura (email 18/05): el DNA experiencial sostenido NO "
                  "es excluyente de abordar el tema precio. Razones:")
    add_bullet(doc, "El 92% que NO prefiere Gama hoy opera mayoritariamente en el modelo precio-dominante (51-84% "
                    "de los preferentes de las otras 7 cadenas citan precio como razón).")
    add_bullet(doc, "La realidad socioeconómica venezolana hace que el precio sea palanca real de captación de "
                    "nuevos clientes, especialmente en segmentos D y E.")
    add_bullet(doc, "Gama no debería distanciarse mucho de competidores en percepción de precio — esto reduciría "
                    "su base de oportunidad sin compensación clara en valor.")
    add_para(doc, "La estrategia correcta combina: (a) defender el DNA experiencial diferenciado, (b) abordar el "
                  "precio creativamente sin disminuir percepción de valor. Las recomendaciones específicas están en "
                  "Cap 3 (precios estratégicos por categoría) y Cap 5 (vías creativas: promociones, esquemas "
                  "diferenciados, etc.).", italic=True)

    add_h2(doc, "2.5 Mundo de marca — Gap verbal vs asociativo (Correlación P21.1 × P23)")
    add_h3(doc, "📊 DATOS")
    headers = ["Cadena", "Razón espontánea #1", "%", "Atributo asociado #1", "%", "Gap"]
    rows = [
        ["Gama", "Atención", "78%", "Limpieza", "31.1%", "Sistémico"],
        ["Páramo", "Precio", "84%", "Seguro", "36.1%", "Sistémico"],
        ["Rio", "Precio", "68%", "Limpieza", "33.6%", "Sistémico"],
        ["Central Madeirense", "Atención", "23%", "Limpieza", "29.1%", "Sistémico"],
        ["Forum", "Atención", "17%", "Limpieza", "29.6%", "Sistémico"],
    ]
    add_table(doc, headers, rows, col_widths=[3.5, 3, 2, 3.5, 2, 2])

    add_h3(doc, "🔍 ANÁLISIS — Insight estructural del mercado")
    add_para(doc, "Hallazgo de diseño de estudio (no de Gama específicamente):")
    add_bullet(doc, "En NINGUNA cadena la razón espontánea principal coincide con el atributo asociado inducidamente "
                    "más frecuente. El patrón es sistémico: espontáneamente los shoppers hablan de precio o "
                    "atención; cuando se les pregunta directamente, mencionan limpieza/seguridad.")
    add_bullet(doc, "Implicación para Gama: la atención (78% espontánea) es el activo que los leales VERBALIZAN. La "
                    "limpieza (31.1% asociada) es la imagen que el mercado AMPLIO TIENE cuando se le pregunta. Ambas "
                    "son verdad y complementarias — el posicionamiento es bi-dimensional (atención + ambiente físico).")
    add_bullet(doc, "La campaña debería ACTIVAR la verbalización de atención en el mercado amplio. El atributo ya "
                    "existe en la imagen (21.9% en P23 Total) pero solo el Núcleo Leal lo verbaliza espontáneamente. "
                    "La tarea de la campaña es trasladar la asociación inducida a la verbalización espontánea.")

    add_metodologia_box(doc,
        "Z-scores DNA calculados sobre media y desviación estándar de las 10 cadenas incluidas. Normalización "
        "descriptiva — no es un z-test inferencial. Variación de z-scores entre V3 y V7 puede reflejar cambios en el "
        "paisaje competitivo (otras cadenas cambiaron sus z-scores) además de cambios en Gama. MDS sobre matriz P23 "
        "estandarizada con stress de Kruskal aceptable (<0.20). La 'correlación nominal' P21.1×P23 es mapeo "
        "semántico — no test estadístico de correlación.")

    # ============ SECCIÓN 2.5 NUEVA V8 — Mundo de marca por cadena ============
    add_h2(doc, "2.5 Mundo de marca por cadena — P23 imagen total × experiencia [NUEVO V8]")
    add_piramide_separator(doc, "datos")

    add_para(doc, "P23 mide qué atributos asocia el respondente a cada cadena, sin haber visitado necesariamente. "
                  "Es asociación INDUCIDA (vs P21.1 que es razón ESPONTÁNEA de preferencia). La tabla siguiente "
                  "muestra el % del mercado total que asocia cada atributo con cada cadena.", size=11)

    headers = ["Cadena", "Limpieza", "Seguro", "Atractiva", "Calidad", "Atención", "Surtido", "Rapidez", "Precio", "Promoc"]
    rows = [
        ["Gama", "31.1%", "29.6%", "28.9%", "26.6%", "21.9%", "17.9%", "21.1%", "7.2%", "9.0%"],
        ["Páramo", "34.6%", "36.1%", "31.3%", "26.6%", "32.3%", "29.4%", "30.3%", "35.1%", "30.6%"],
        ["Rio", "33.6%", "27.1%", "28.1%", "26.1%", "23.9%", "21.9%", "21.6%", "13.9%", "17.4%"],
        ["Central Mad.", "29.1%", "27.6%", "18.9%", "22.6%", "21.1%", "20.6%", "20.4%", "17.7%", "17.2%"],
        ["Forum", "29.6%", "26.9%", "25.4%", "24.9%", "21.1%", "24.1%", "20.1%", "18.7%", "18.7%"],
        ["Plazas", "30.6%", "28.1%", "19.9%", "21.4%", "21.1%", "17.4%", "15.7%", "10.9%", "15.9%"],
        ["Plan Suárez", "14.9%", "16.2%", "11.2%", "13.4%", "11.4%", "14.2%", "10.4%", "11.9%", "11.2%"],
        ["Luz", "14.9%", "15.2%", "9.7%", "10.9%", "13.4%", "11.4%", "10.0%", "12.2%", "14.9%"],
    ]
    add_table(doc, headers, rows, col_widths=[1.7, 1.0, 0.9, 1.0, 0.9, 1.0, 0.9, 0.95, 0.85, 0.9])

    add_chart(doc, os.path.join(CHARTS_CU10, "cu10_heatmap_mundo_total.png"), width_in=6.8,
              caption="Heatmap P23 imagen total — Páramo es la única cadena con imagen completa en todas las dimensiones.")

    add_piramide_separator(doc, "analisis")
    add_para(doc, "Tres lecturas clave del mundo de marca en el mercado total:", bold=True)
    add_bullet(doc, "Gama tiene perfil experiencial puro — domina en Limpieza (31.1%), Seguro (29.6%), Atractiva "
                    "(28.9%) y Calidad (26.6%). Pero es DÉBIL en Precio (7.2%) y Promociones (9.0%) — las brechas más "
                    "grandes vs Páramo son justamente en precio (+27.9 pp) y promociones (+21.6 pp).")
    add_bullet(doc, "Páramo es la ÚNICA cadena con asociación fuerte en TODOS los atributos. Domina en Seguro "
                    "(36.1%), Precio (35.1%), Limpieza (34.6%), Atención (32.3%), Valer dinero (32.1%), Promociones "
                    "(30.6%). Es la cadena con imagen más rica y completa del mercado venezolano. Competir frontalmente "
                    "con Páramo es imposible para Gama — pero tampoco es necesario.")
    add_bullet(doc, "Todas las cadenas convergen en Limpieza y Seguro como tabla de higiene — son atributos de piso "
                    "que cualquier cadena del segmento medio-alto debe tener para competir. El diferenciador real "
                    "emerge en el 4to-5to atributo y en los atributos de SOMBRA (lo que los preferentes ven que el "
                    "mercado amplio no ve).")

    add_h3(doc, "2.5.1 ATENCIÓN — el atributo de sombra de Gama (descubrimiento clave V8)")
    add_para(doc, "Cuando comparamos la imagen total de Gama (P23) con la imagen entre quienes la prefieren o han "
                  "comprado recientemente, aparece un patrón crítico:", size=11)

    headers = ["Atributo", "Total (n=402)", "Pref-Gama (n=32) *", "Exp. reciente (n=30) *", "Δ Pref vs Total", "Δ Exp vs Total"]
    rows = [
        ["Mejor atención", "21.9%", "84.4%", "83.3%", "+62.5 pp", "+61.4 pp"],
        ["Limpieza/orden", "31.1%", "90.6%", "76.7%", "+59.5 pp", "+45.6 pp"],
        ["Seguro", "29.6%", "84.4%", "86.7%", "+54.8 pp", "+57.1 pp"],
        ["Mayor surtido", "17.9%", "71.9%", "80.0%", "+54.0 pp", "+62.1 pp"],
        ["Rapidez caja", "21.1%", "75.0%", "73.3%", "+53.9 pp", "+52.2 pp"],
        ["Tienda atractiva", "28.9%", "81.2%", "73.3%", "+52.3 pp", "+44.4 pp"],
        ["Mayor calidad", "26.6%", "78.1%", "76.7%", "+51.5 pp", "+50.1 pp"],
        ["Menor precio", "7.2%", "31.2%", "43.3%", "+24.0 pp", "+36.1 pp"],
    ]
    add_table(doc, headers, rows, col_widths=[1.8, 1.1, 1.3, 1.3, 1.3, 1.3])
    add_para(doc, "* REFERENCIAL n<30. Denominador en \"experiencia reciente\" = quienes compraron en Gama en su "
                  "última visita (P24), NO muestra total — interpretar con esa base.", size=9, italic=True, color=GRAY_MID)

    add_chart(doc, os.path.join(CHARTS_CU10, "cu10_gama_brechas_segmentos.png"), width_in=6.8,
              caption="Brechas P23 Gama — la atención tiene la brecha más grande entre la imagen general (21.9%) y la imagen de quienes la prefieren (84.4%).")

    add_callout_box(doc,
        "Gama tiene una IMAGEN BIFURCADA. Para quien no la conoce bien, es una cadena experiencial mediocre — por "
        "debajo de Páramo en casi todo. Para quien la prefiere o ha comprado recientemente, es una cadena de "
        "excelencia experiencial percibida. La brecha más grande (+62.5 pp) es en ATENCIÓN — el atributo de SOMBRA "
        "de Gama. La tarea de comunicación es trasladar la imagen de los leales al mercado amplio.")

    add_para(doc, "La validación por experiencia es contundente: 83.3% de los compradores recientes de Gama la asocian "
                  "con Atención, vs solo 21.9% del total. La experiencia REAL de comprar en Gama confirma y amplifica "
                  "la imagen de atención. La barrera no es el producto/servicio — es la VISIBILIDAD y RECONOCIMIENTO "
                  "de ese atributo en el mercado amplio que no ha entrado todavía.", size=11)

    add_chart(doc, os.path.join(CHARTS_CU10, "cu10_heatmap_mundo_experiencia.png"), width_in=6.8,
              caption="Heatmap P23 entre compradores recientes (P24) — todas las cadenas mejoran dramáticamente; el efecto es universal pero la magnitud diferencia.")

    # ============ SECCIÓN 2.6 NUEVA V8 — Vecindad perceptual Gama-Plazas ============
    add_h2(doc, "2.6 Espacio competitivo — vecindad perceptual Gama vs Plazas vs Rio [NUEVO V8]")
    add_piramide_separator(doc, "datos")

    add_para(doc, "El análisis del mundo de marca revela cuáles cadenas son perceptualmente VECINAS de Gama (mismo "
                  "cluster experiencial) y cuáles son LEJANAS (cluster opuesto):", size=11)

    add_h3(doc, "2.6.1 Tabla resumen — mundo de marca compacto por las 8 cadenas")
    headers = ["Cadena", "Atributos centrales (top 3 P23 Total)", "4to atributo", "Atributo SOMBRA (preferentes)", "n pref"]
    rows = [
        ["Gama", "Limpieza, Seguro, Atractiva", "Calidad", "Atención (+62.5 pp)", "32 REF*"],
        ["Páramo", "Seguro, Precio, Limpieza", "Atención", "Valer dinero (+60 pp)", "85"],
        ["Rio", "Limpieza, Atractiva, Seguro", "Atención", "Calidad (+69 pp)", "41"],
        ["Central Mad.", "Limpieza, Seguro, Calidad", "Atención", "Atención (+63 pp)", "45"],
        ["Forum", "Limpieza, Seguro, Atractiva", "Surtido", "Valer dinero (+60 pp)", "44"],
        ["Plazas", "Limpieza, Seguro, Calidad", "Atención", "Atención (+68 pp)", "37"],
        ["Plan Suárez", "Seguro, Limpieza, Surtido", "Precio", "Valer dinero (+71 pp)", "28 REF*"],
        ["Luz", "Seguro, Promociones, Limpieza", "Atención", "Surtido (+67 pp)", "32 REF*"],
    ]
    add_table(doc, headers, rows, col_widths=[1.4, 2.7, 1.3, 2.0, 1.0])

    add_piramide_separator(doc, "analisis")

    add_h3(doc, "2.6.2 Plazas es el competidor perceptual MÁS CERCANO a Gama")
    add_para(doc, "Plazas comparte con Gama:", bold=True)
    add_bullet(doc, "Misma estructura de atributos centrales: Limpieza, Seguro, Calidad (vs Limpieza, Seguro, Atractiva en Gama).")
    add_bullet(doc, "Mismo atributo SOMBRA: Atención (+68 pp en Plazas, +62.5 pp en Gama).")
    add_bullet(doc, "Mismo perfil de DEBILIDAD en precio: Plazas 10.9% asociación de precio (vs 7.2% Gama).")
    add_para(doc, "Diferencia clave: Plazas tiene n preferentes mayor (37 vs 32) pero menor cobertura geográfica y "
                  "menor notoriedad TOM. En el espacio perceptual del consumidor, son cadenas casi gemelas — Gama "
                  "necesita diferenciarse de Plazas en el atributo de Atención para que el mercado amplio capte la "
                  "diferencia.", size=11)

    add_h3(doc, "2.6.3 Rio comparte cluster experiencial pero con ventaja de CALIDAD")
    add_para(doc, "Rio es el segundo competidor perceptual más cercano a Gama en el cluster experiencial. Pero su "
                  "atributo SOMBRA es CALIDAD (+69 pp entre compradores recientes — 85.1% asocian Rio con calidad vs "
                  "26.1% del total). Si Rio comunicara su calidad de productos al mercado amplio, podría capturar al "
                  "consumidor experiencial que hoy considera Gama. Es la amenaza competitiva latente más relevante "
                  "para Gama en el medio plazo.", size=11)

    add_h3(doc, "2.6.4 Páramo es el opuesto completo — competencia en territorio diferente")
    add_para(doc, "Páramo domina en precio/valor mientras Gama domina en ambiente/experiencia. No compiten en la "
                  "misma dimensión — Páramo captura al consumidor precio-sensible, Gama al consumidor experiencial. "
                  "La estrategia correcta para Gama no es disputar el territorio de Páramo (perdería) sino blindar y "
                  "amplificar su propio territorio + atacar a Plazas y Rio dentro del cluster experiencial.", size=11)

    add_callout_box(doc,
        "Implicación competitiva V8: Gama tiene dos prioridades. (1) DEFENDER el atributo de Atención frente a "
        "Plazas — son vecinos perceptuales y el primero que active el atributo en el mercado amplio se queda con él. "
        "(2) VIGILAR a Rio en calidad — si Rio comunica calidad, podría migrar de su posicionamiento actual precio-"
        "dominante a una propuesta mixta calidad-precio que sería muy difícil de combatir.")

    add_metodologia_box(doc,
        "Vecindad perceptual basada en patrones de asociación P23 (matriz de imagen) y atributos sombra (brechas "
        "Pref vs Total). No es clustering jerárquico formal — es lectura analítica de la matriz P23 con z-scores "
        "DNA del CU-10. Las cadenas con n preferentes <30 (Gama, Plan Suárez, Luz) tienen lecturas indicativas. La "
        "comparación Gama-Plazas-Rio es robusta porque las tres tienen patrones consistentes con el T2B previo + "
        "experiencia reciente.")


# =========== CAPÍTULO 3: PRECIOS + CATEGORÍAS + MISIONES ==========

def build_cap3(doc):
    add_h1(doc, "Capítulo 3 — Precios, categorías y misiones de compra")

    add_callout_box(doc,
        "El precio es el atributo MENOS predictivo de preferencia (OR=1.03 NS, SHAP #10) pero la PERCEPCIÓN de "
        "precio favorable predice preferencia con OR=2.4 (IC95 [1.12, 5.12]). Resolución de la paradoja: el precio "
        "no es driver de elección entre quien ya conoce Gama, pero la percepción agregada actúa como barrera de "
        "entrada para quien nunca la ha experimentado. Por categoría, Gama no lidera precio en ninguna de alta "
        "rotación. Pero las categorías Congelados, Gaseosas y Salsas tienen hábito Gama sin liderazgo de precio — "
        "territorio donde la palanca es ofrecer valor, no reducir precio.")

    add_h3(doc, "Contexto crítico — Gama es el único outlier en un mercado precio-dominante")
    add_para(doc, "Antes de entrar en el detalle de precios, vale recordar el hallazgo del Cap 1.3 (modelo mental):", bold=True)
    add_bullet(doc, "7 de 8 cadenas del mercado son precio-dominantes (51-84% de sus preferentes citan precio como "
                    "razón espontánea de preferencia).")
    add_bullet(doc, "Gama es la ÚNICA excepción (40.6% mencionan precio).")
    add_para(doc, "Este contexto cambia cómo se interpreta el análisis de precios que sigue:", bold=True)
    add_bullet(doc, "Para los preferentes de Gama (el 8% del mercado): el precio NO es prioridad — la propuesta de "
                    "valor es experiencial. Aquí las categorías 'ofrecer valor' (Congelados, Gaseosas, Salsas) son la palanca.")
    add_bullet(doc, "Para los no-preferentes de Gama (el 92% restante): el precio SÍ es palanca real. La realidad "
                    "venezolana hace que la gente se mueva por precio. Aquí el análisis de categorías 'cuidar precio' "
                    "y la búsqueda de vías creativas (Cap 5) son críticos para captación.")
    add_bullet(doc, "Plan Suárez (71.4% precio, 15.4% preferencia en C+/C) es el competidor más relevante para "
                    "captar al no-cliente del segmento medio — Páramo y Rio son más fuertes pero tienen DNA "
                    "estructuralmente similar a Plan Suárez.")

    add_h2(doc, "3.1 Percepción global de precios (P33, P34, P31)")
    add_h3(doc, "📊 DATOS — P33 percepción Gama vs competencia (Total y NSE)")
    headers = ["Grupo", "n", "NETO caro %", "NETO económico %", "Brecha pp"]
    rows = [
        ["Total", "402", "54.0%", "15.4%", "+38.6"],
        ["C+/C", "104", "60.6%", "11.5%", "+49.1"],
        ["D", "127", "49.6%", "18.1%", "+31.5"],
        ["E", "171", "53.2%", "15.8%", "+37.4"],
        ["Pref-Gama (ref.)", "32", "34.4%", "(estimado mayor)", "Menor — distinto al mercado"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 2, 3, 3, 3])
    add_chart(doc, os.path.join(CHARTS_V6, "C10_p33_percepcion_precio.png"), width_in=6.5,
              caption="Percepción de precios de Gama por NSE. Problema transversal — más fuerte en C+/C.")

    add_h3(doc, "📊 DATOS — P34 evolución 6 meses + P31 ranking de precio")
    add_bullet(doc, "P34: 45% percibe que Gama subió precios en último semestre, solo 10.5% percibe bajada.")
    add_bullet(doc, "P31 ranking: Gama es #6 (mean rank 5.77/10, Top-3 21.9%). Páramo lidera (4.24, 47.8%).")
    add_chart(doc, os.path.join(CHARTS_V6, "C11_p31_ranking_precio.png"), width_in=6.8,
              caption="Ranking de precio percibido — Gama es la 6ta cadena más económica del mercado.")

    add_h2(doc, "3.2 Mejor precio por categoría (P32) — 15 categorías")
    add_chart(doc, os.path.join(CHARTS_V6, "C12_p32_categorias.png"), width_in=7.0,
              caption="Mejor precio percibido por categoría: Gama vs líder de cada categoría. Páramo domina 11/15.")
    add_para(doc, "Hallazgo crítico: Gama no lidera precio en ninguna categoría de alta rotación. Sus mejores "
                  "posiciones son Farmacia (#2, gap -0.7pp) y Licores (#4, gap -1.2pp) — ambas con bases de "
                  "percepción muy diluidas (84% y 66% 'Ninguno').")
    add_para(doc, "Proteínas (carne, pollo, charcutería): Páramo lidera con gaps de 33pp vs Gama. En NSE C+/C la "
                  "brecha es de 29-30pp. Son las categorías que más penalizan la imagen de precio total.", bold=True)

    add_h2(doc, "3.3 Cruce percepción precio × preferencia (P33 × P21)")
    add_chart(doc, os.path.join(CHARTS_V6, "C13_p33xp21_gradiente.png"), width_in=6.5,
              caption="Gradiente monótono claro entre percepción precio y preferencia. OR=2.4 sig.")

    add_para(doc, "Resolución de la paradoja precio:", bold=True)
    add_bullet(doc, "Atributo precio NO predice preferencia (OR=1.03 NS, atributo #10 en SHAP) — entre quien ya "
                    "conoce Gama, el precio no es criterio de elección.")
    add_bullet(doc, "Percepción agregada P33 SÍ predice preferencia (OR=2.4 sig, chi²=21.94 p<0.001). Pasar de "
                    "\"Gama es cara\" a \"Gama es igual o económica\" duplica probabilidad de preferencia.")
    add_bullet(doc, "Interpretación: la percepción agregada actúa como barrera de entrada. Quien percibe a Gama como "
                    "cara NUNCA experimenta el activo de atención (OR=5.73). La atención sigue siendo driver más "
                    "potente (5.73 vs 2.4) — pero la percepción de precio define quién llega a experimentarla.")

    add_h2(doc, "3.4 Hábito de compra × percepción precio por categoría (P30 × P32)")
    add_h3(doc, "📊 DATOS — Categorías \"territorio protegido\" Gama")
    headers = ["Categoría", "Hábito Gama %", "Precio Gama %", "Brecha pp", "Tipo"]
    rows = [
        ["Congelados", "8.0%", "4.7%", "+3.3", "Hábito sin precio (ofrecer valor)"],
        ["Salsas y Enlatados", "7.5%", "4.5%", "+3.0", "Hábito sin precio (ofrecer valor)"],
        ["Gaseosas, jugos", "7.2%", "3.2%", "+4.0", "Hábito sin precio (ofrecer valor)"],
        ["Galletas y confitería", "7.0%", "5.7%", "+1.3", "Cuidar precio"],
        ["Cuidado hogar", "5.0%", "3.7%", "+1.3", "Neutro"],
        ["Farmacia (líder hábito)", "4.2%", "3.0%", "+1.2", "Conveniencia 24h"],
        ["Carne", "3.7%", "3.0%", "+0.7", "Terreno Páramo"],
        ["Pollo", "3.5%", "2.5%", "+1.0", "Terreno Páramo"],
        ["Charcutería", "3.2%", "1.7%", "+1.5", "Terreno Páramo"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 2.5, 2.5, 2, 4])

    add_h3(doc, "🔍 ANÁLISIS — Categorías \"ofrecer valor\" vs \"cuidar precio\" (respuesta a Cora)")
    add_chart(doc, os.path.join(CHARTS_V7, "V7C5_estrategia_categorias.png"), width_in=6.5,
              caption="Clasificación estratégica de categorías para Gama (todas referenciales para Gama por n<30).")

    add_para(doc, "Respuesta a la pregunta clave de Cora (reunión 18/05 min 14:43): "
                  "\"¿Cuáles son las categorías donde Gama tiene que cuidar precio y cuáles donde debería ofrecer valor?\"", bold=True)

    add_para(doc, "OFRECER VALOR (mantener precio, comunicar diferencial experiencial):", bold=True, color=GREEN_OK)
    add_bullet(doc, "Congelados, Gaseosas/Jugos, Salsas y Enlatados — en estas 3 categorías el shopper ya elige "
                    "Gama aunque no la percibe como la más económica. Evidencia directa de que la propuesta de "
                    "valor funciona sin liderazgo de precio. Palanca: comunicar por qué vale venir a Gama (servicio, "
                    "conveniencia, ambiente), no reducir precio.")
    add_bullet(doc, "Farmacia — categoría de liderazgo de hábito sin liderazgo de precio. Comunicar 24h/conveniencia.")

    add_para(doc, "CUIDAR PRECIO (defender posición, no perder hábito):", bold=True, color=BLUE_INFO)
    add_bullet(doc, "Galletas/confitería — brecha precio pequeña (+1.3pp), hábito alto (7%). Si Gama cede en "
                    "precio aquí pierde hábito porque hay alternativas competitivas cercanas.")
    add_bullet(doc, "Productos básicos — Páramo lidera precio; Gama tiene hábito moderado que puede erosionarse "
                    "si el gap precio sube.")

    add_para(doc, "SIN ACCIÓN DE PRECIO INMEDIATA (terreno Páramo):", bold=True, color=ORANGE_W)
    add_bullet(doc, "Proteínas (carne, pollo, charcutería) — Páramo domina con brechas enormes (33pp). Competir en "
                    "precio aquí es batalla de alto costo con baja probabilidad de conversión. No es territorio "
                    "natural de Gama.")

    add_caveat_box(doc,
        "Análisis 100% perceptual (P30/P32 de encuesta). Las decisiones de precio deben validarse con datos reales "
        "de pricing operativo. Un shopper que percibe precio alto puede estar equivocado o puede estar en lo correcto. "
        "n de Gama por categoría es REFERENCIAL en todos los casos (CV-9-09).")

    add_h2(doc, "3.5 Misiones de compra (P25, P26)")
    add_chart(doc, os.path.join(CHARTS_V6, "C20_misiones_gama.png"), width_in=6.5,
              caption="Distribución de misiones de última compra y posición Gama por misión.")
    add_bullet(doc, "Misión dominante del mercado: Reabastecimiento parcial (67.2%).")
    add_bullet(doc, "Gama lidera \"urgencia / pocos productos\" con 12.2% — único territorio donde es la opción #1.")
    add_bullet(doc, "Gama está en posición #7 para abastecimiento general (mercado grande): pierde el 34% del share "
                    "of wallet de sus propios preferentes.")

    add_metodologia_box(doc,
        "Cruces P30 (hábito por categoría) × P32 (mejor precio por categoría) × P25 (misión última compra) sobre "
        "n=402. Todas las cifras de Gama por categoría individual son REFERENCIALES (n<30 en casi todos los casos). "
        "Los patrones son indicativos. Cruce P33×P21 con chi² y OR — base agregada Total + subset n=185 (favorables).")


# =========== CAPÍTULO 4: COMUNICACIÓN PTL + DTLS + PERFIL RECORDADOR ==========

def build_cap4(doc):
    add_h1(doc, "Capítulo 4 — Comunicación PTL y DTLS — perfil del recordador")

    add_callout_box(doc,
        "Ambas frases tienen recall bajo (~11-12%) y entre ellas no hay diferencia significativa. Pero hay una "
        "diferencia crítica en la interpretación: DTLS activa territorio relacional (26%) que PTL prácticamente no "
        "toca (2%). El hallazgo más relevante para Cora: las dos frases llegan al segmento equivocado. 58% de los "
        "recordadores son NSE E, solo 19% son C+/C — el segmento natural de Gama está SUBREPRESENTADO entre "
        "quienes recordaron las frases. La campaña activa hábito de compra (+5-7pp) pero NO mueve percepción de precio.")

    add_h2(doc, "4.1 Recall y comparativo PTL vs DTLS")
    add_h3(doc, "📊 DATOS")
    headers = ["Métrica", "Resultado", "Base", "Certeza"]
    rows = [
        ["Recall espontáneo cualquier slogan Gama (P35)", "4.2%", "402", "✅ Alta"],
        ["Recall PTL \"Precios de tu lado\" (P37)", "10.7%", "402", "✅ Alta"],
        ["Recall DTLS \"De tu lado siempre\" (P40)", "12.4%", "402", "✅ Alta"],
        ["Diferencia recall PTL vs DTLS", "z=-0.754, p=0.451", "—", "NO significativo"],
        ["Algún slogan recordado", "23.1%", "402", "Combinado P35+P37+P40"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 3, 2, 4])
    add_chart(doc, os.path.join(CHARTS_V6, "C15_ptl_vs_dtls.png"), width_in=7.0,
              caption="Comparativo PTL vs DTLS: recall similar (n.s.); interpretación temática muy diferente.")

    add_h3(doc, "🔍 ANÁLISIS")
    add_bullet(doc, "PTL y DTLS tienen recall estadísticamente equivalentes (~11-12%). El 88% del mercado no recuerda "
                    "ni una ni otra. El problema NO es cuál frase es mejor sino que ninguna ha tenido distribución suficiente.")
    add_bullet(doc, "DTLS activa territorio relacional (26% interpretación 'acompañamiento/apoyo') que PTL no toca "
                    "(2%). Si el objetivo estratégico es posicionarse en territorio relacional, DTLS tiene mayor "
                    "potencial semántico — aunque PTL genera mayor agrado declarado (53% vs 42%, diferencia n.s.).")

    add_h2(doc, "4.2 Perfil del recordador — NSE")
    add_h3(doc, "📊 DATOS")
    headers = ["NSE", "Muestra total %", "Recuerdan PTL (n=43)", "Recuerdan DTLS (n=50)"]
    rows = [
        ["E", "42.5%", "58.1%", "54.0%"],
        ["D", "31.6%", "23.3%", "26.0%"],
        ["C+/C", "25.9%", "18.6%", "20.0%"],
    ]
    add_table(doc, headers, rows, col_widths=[3, 4, 4, 4])
    add_chart(doc, os.path.join(CHARTS_V7, "V7C3_perfil_recordadores.png"), width_in=6.8,
              caption="Distribución NSE del recordador vs muestra. NSE E sobrerepresentado, C+/C subrepresentado.")

    add_h3(doc, "🔍 ANÁLISIS")
    add_para(doc, "Hallazgo crítico: la campaña llega al segmento equivocado.", bold=True)
    add_bullet(doc, "NSE C+/C representa solo el 19% de los recordadores de PTL y el 20% de los de DTLS — pero es el "
                    "25.9% de la muestra total. SUBREPRESENTADO en 6-7 pp.")
    add_bullet(doc, "NSE E representa el 54-58% de los recordadores — pero es solo el 42.5% de la muestra. "
                    "SOBRERREPRESENTADO en 12-16 pp.")
    add_bullet(doc, "Implicación: el mensaje PTL (\"Precios de tu lado\") llega principalmente al segmento que más "
                    "percibe a Gama como cara y menos compra Gama. La campaña refuerza un atributo (precio) en un "
                    "segmento (E) donde Gama no compite estructuralmente.")
    add_bullet(doc, "El target objetivo de Gama (C+/C, su segmento natural con TOM 60.6%) está siendo MENOS alcanzado "
                    "por la campaña que la muestra promedio del mercado.")

    add_h2(doc, "4.3 Re-segmentación de TODAS las preguntas según recall publicitario")
    add_chart(doc, os.path.join(CHARTS_V7, "V7C4_resegmentacion_recall.png"), width_in=6.8,
              caption="Re-segmentación recall vs no-recall — diferencias relevantes en hábito Gama (+5-7pp).")

    add_h3(doc, "📊 DATOS — Diferencias relevantes (>5pp) entre recordadores y no-recordadores")
    headers = ["Variable", "Recuerda %", "No recuerda %", "Diferencia pp"]
    rows = [
        ["Hábito Gama Gaseosas (P30)", "13.1%", "6.2%", "+6.9"],
        ["Hábito Gama Salsas (P30)", "13.1%", "6.5%", "+6.6"],
        ["Hábito Gama Galletas (P30)", "13.1%", "5.9%", "+7.2"],
        ["Hábito Gama Licores (P30)", "8.2%", "2.9%", "+5.3"],
        ["Misión abastecimiento en Gama (P26)", "13.1%", "6.2%", "+6.9"],
    ]
    add_table(doc, headers, rows, col_widths=[6, 3, 3, 3])

    add_h3(doc, "🔍 ANÁLISIS")
    add_bullet(doc, "La recordación publicitaria SE ASOCIA con mayor hábito de compra en Gama (+5-7pp en 4 categorías). "
                    "Asociación, no causalidad: pueden ser clientes habituales que ven más publicidad.")
    add_bullet(doc, "La misión de abastecimiento en Gama es mayor entre recordadores (+6.9pp). Quienes recuerdan "
                    "la publicidad van a Gama para 'hacer el mercado completo' en mayor proporción — coherente con "
                    "el perfil del Núcleo Leal.")
    add_bullet(doc, "NO se detectan diferencias relevantes en preferencia ni en percepción de precio. La publicidad "
                    "actual NO está generando conversión de preferencia ni moviendo el posicionamiento de precio — "
                    "solo está activando comportamiento entre quienes ya tienen hábito.")

    add_metodologia_box(doc,
        "Recall definido como P35=Sí OR P37=Sí OR P40=Sí. Los dos grupos (93 recuerdan vs 309 no recuerdan) tienen "
        "bases OK para comparación a nivel total. Cruces por NSE para DTLS producen n=10 en C+/C — referenciales. La "
        "diferencia recall-hábito es asociativa, no causal. (CV-9-08: recall-hábito puede reflejar que compradores "
        "frecuentes ven más publicidad, no que la publicidad genera hábito.)")

    # ============ SECCIÓN 4.4 NUEVA V8 — Drivers reformulados ============
    add_h2(doc, "4.4 Drivers de preferencia — reformulación V8 (imagen → preferencia, no importancia → preferencia)")
    add_piramide_separator(doc, "datos")

    add_para(doc, "El CU-10 evaluó la viabilidad de regresión lineal/logística como Cora pidió en su email del 18/05 "
                  "19:25. El resultado es metodológicamente importante para entender qué modelo conceptual rige la "
                  "preferencia de cadena en este mercado.", size=11)

    add_h3(doc, "4.4.1 Evaluación de viabilidad — VIF y tamaños muestrales")
    headers = ["Criterio", "Resultado", "Decisión"]
    rows = [
        ["VIF máximo entre 10 importancias P22", "1.35 (umbral aceptable: <5)", "✅ Sin multicolinealidad"],
        ["Cadenas con n ≥50 preferentes (umbral regresión confiable)", "Solo Páramo (n=85)", "⚠ 7 de 8 cadenas excluidas"],
        ["Método estadístico correcto para DV binaria preferencia 0/1", "Regresión LOGÍSTICA (no lineal)", "✅ Ajuste técnico aplicado"],
    ]
    add_table(doc, headers, rows, col_widths=[4, 4, 3])

    add_para(doc, "Nota técnica para Cora: la pregunta de email mencionaba \"regresión lineal\". La variable "
                  "dependiente (preferencia 0/1) es binaria, lo que invalida los supuestos de la regresión lineal "
                  "(produce probabilidades fuera de [0,1] y residuos no-normales). El método correcto es la regresión "
                  "logística, que produce probabilidades dentro del rango y coeficientes interpretables como odds "
                  "ratios. Aplicamos esa corrección.", size=10, italic=True, color=GRAY_MID)

    add_h3(doc, "4.4.2 Regresión logística Páramo (único caso viable) — DV: prefiere Páramo (1/0)")
    headers = ["Parámetro", "Valor"]
    rows = [
        ["n total", "402"],
        ["n preferentes Páramo", "85"],
        ["Pseudo-R² McFadden", "0.021 (POBRE)"],
        ["LLR p-value (modelo global)", "0.574 (NO significativo)"],
        ["Único coeficiente significativo", "Mayor calidad: OR=0.648, p=0.042"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 6])

    add_chart(doc, os.path.join(CHARTS_CU10, "cu10_coefs_paramo.png"), width_in=6.5,
              caption="Coeficientes logísticos Páramo — solo Mayor calidad sale significativo, y con signo NEGATIVO (más exigencia de calidad → menos preferencia Páramo).")

    add_piramide_separator(doc, "analisis")
    add_callout_box(doc,
        "HALLAZGO METODOLÓGICO CENTRAL V8: las importancias declaradas en P22 NO predicen la preferencia de cadena. "
        "El modelo logístico Páramo (único caso viable) explica menos del 2.1% de la varianza. ¿Por qué? Porque las "
        "importancias son uniformemente altas (TB puro 57-64% para casi todos los atributos) y NO discriminan entre "
        "quienes prefieren distintas cadenas. El verdadero predictor de preferencia es la IMAGEN P23 (asociación "
        "atributo × cadena), no la importancia P22.")

    add_h3(doc, "4.4.3 Reformulación conceptual — el modelo correcto es imagen → preferencia")
    add_para(doc, "El modelo conceptual implícito en V7 era: \"importancia × percepción\" (típico IP map). El CU-10 "
                  "demuestra que ese modelo NO opera en este mercado. El modelo correcto es:", size=11)

    add_para(doc, "PERCEPCIÓN DE QUE LA CADENA TIENE EL ATRIBUTO → PREFERENCIA", bold=True, color=RED_GAMA, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)

    add_para(doc, "Los consumidores TODOS valoran los mismos atributos como importantes (techo de importancia). Lo que "
                  "diferencia su preferencia de cadena es qué cadena PERCIBEN como dueña de esos atributos. Quien "
                  "percibe que Gama tiene Atención (21.9% del total) la considera. Quien no la percibe así (78.1% del "
                  "total) no la considera para ese atributo aunque la atención SÍ le importe.", size=11)

    add_h3(doc, "4.4.4 Implicación para campaña Gama 2026-2027")
    add_para(doc, "La campaña NO debe argumentar \"la atención es importante\" — todos ya lo creen. La campaña debe "
                  "argumentar \"Gama TIENE atención superior\" + demostrarlo. Es un cambio sustantivo en el copy y "
                  "en la dirección creativa:", size=11)
    add_bullet(doc, "Anti-mensaje: \"En Gama te ofrecemos buen servicio\" — genérico, no diferencia.")
    add_bullet(doc, "Mensaje correcto: testimonios + casos concretos + comparativas visuales que muestran que la "
                    "atención en Gama es SUPERIOR. La evidencia es contundente: 83.3% de quienes han comprado "
                    "recientemente la asocian con atención (vs 21.9% del total que no la conoce).")
    add_bullet(doc, "Implicación: la campaña debe activar EXPERIENCIA + RECONOCIMIENTO, no DECLARACIÓN. Una vez que "
                    "alguien pruebe Gama, su asociación con atención sube +62.5 pp. La barrera está en hacer que el "
                    "no-cliente entre por primera vez.")

    add_h3(doc, "4.4.5 Análisis descriptivo Gama (alternativa a regresión por n insuficiente)")
    add_para(doc, "Para las 7 cadenas excluidas de regresión (incluyendo Gama n=32), reportamos la diferencia en TB% "
                  "entre preferentes y no-preferentes de cada cadena. Para Gama:", size=11)
    headers = ["Atributo", "TB% Pref-Gama *", "TB% No-Pref-Gama", "Diferencia"]
    rows = [
        ["Rapidez caja", "84.4%", "55.5%", "+28.9 pp"],
        ["Mejor atención", "71.9%", "54.5%", "+17.4 pp"],
        ["Valer dinero", "71.9%", "55.7%", "+16.2 pp"],
        ["Tienda atractiva", "50.0%", "41.8%", "+8.2 pp"],
        ["Mayor surtido", "62.5%", "56.5%", "+6.0 pp"],
        ["Menor precio", "43.8%", "64.0%", "-20.2 pp"],
        ["Promociones", "25.0%", "52.4%", "-27.4 pp"],
    ]
    add_table(doc, headers, rows, col_widths=[3, 2.5, 3, 2.5])
    add_para(doc, "* REFERENCIAL n=32. Análisis descriptivo NO equivalente a regresión — no controla por covariables. Es indicativo.",
                  size=9, italic=True, color=GRAY_MID)
    add_para(doc, "Los preferentes de Gama son el segmento que MÁS exige Rapidez (+28.9 pp) y Atención (+17.4 pp) y "
                  "el que MENOS exige Precio (-20.2 pp) y Promociones (-27.4 pp) vs el resto del mercado. Esta "
                  "brecha descriptiva es el sello del Núcleo Leal de Gama y refuerza la tesis V7 del DNA experiencial.", size=11)

    add_metodologia_box(doc,
        "Para versiones futuras (ola 2027), si se quiere predecir preferencia con regresión, la variable independiente "
        "correcta es la asociación INDUCIDA P23 (binaria por atributo × cadena) cruzada con P22 como ponderador de "
        "importancia. Esto es un análisis de importancia-rendimiento (IP map) que está fuera del scope CU-10 pero "
        "se puede calcular con los datos actuales si Cora lo solicita.")


# =========== CAPÍTULO 5: SÍNTESIS Y RECOMENDACIONES BIFURCADAS ==========

def build_cap5(doc):
    add_h1(doc, "Capítulo 5 — Síntesis estratégica y recomendaciones (bifurcadas)")

    add_callout_box(doc,
        "Las recomendaciones se presentan bifurcadas según el segmento objetivo, como pidió Cora en la reunión: "
        "qué hacer en Total (toda la base de oportunidad) y qué hacer específicamente en C+/C (el segmento natural "
        "que define el éxito estratégico de Gama). Las palancas son diferentes para cada segmento.")

    add_h2(doc, "5.1 Tesis V7")
    add_para(doc, "Gama es estructuralmente diferente del mercado venezolano de supermercados: la única cadena cuyos "
                  "preferentes NO eligen mayoritariamente por precio (40.6% mencionan precio vs 51-84% en las otras "
                  "7 cadenas). Su DNA experiencial se ha consolidado (6 atributos con sobreíndice vs 4 en V3) y su "
                  "segmento natural (C+/C) lo ve incluso más fuerte en atributos experienciales. La estrategia "
                  "correcta combina dos movimientos simultáneos: (a) DEFENDER el DNA experiencial diferenciado como "
                  "ventaja estructural sostenible, (b) ABORDAR el tema precio de manera creativa para captar al 92% "
                  "que no prefiere Gama hoy — sin disminuir la percepción de valor que sostiene la diferenciación. "
                  "Ambos movimientos no son excluyentes: la realidad socioeconómica venezolana hace que el precio sea "
                  "palanca real de captación, especialmente entre no-clientes y NSE más bajos.",
             bold=True)

    add_h2(doc, "5.2 Recomendaciones — TOTAL (toda la base)")
    add_h3(doc, "REC-T1 — Reorientar el anzuelo comunicacional")
    add_bullet(doc, "Mantener PTL/DTLS como recursos pero cambiar el wording del mensaje promocional: en lugar de "
                    "\"precio bajo\" usar \"razón específica para venir esta semana\" (basado en hallazgo de "
                    "ineficacia del mensaje de precio + atributo precio OR=1.03 NS).")
    add_bullet(doc, "Activar la verbalización de atención en el mercado amplio. El atributo ya existe en la imagen "
                    "(21.9% en P23 Total) pero solo el Núcleo Leal lo verbaliza espontáneamente. Cap. 2.5.")

    add_h3(doc, "REC-T2 — Categorías \"ofrecer valor\" — comunicar diferencial sin reducir precio")
    add_bullet(doc, "Congelados, Gaseosas/Jugos, Salsas y Enlatados: comunicar por qué vale venir a Gama (servicio, "
                    "conveniencia, ambiente, 24h) en estas categorías. NO reducir precio. Cap. 3.4.")
    add_bullet(doc, "Farmacia: comunicar conveniencia y 24h. Gama ya lidera en hábito en farmacia (4.2%) — convertir "
                    "ese hábito en convicción.")

    add_h3(doc, "REC-T3 — Defensa estructural del DNA experiencial")
    add_bullet(doc, "Aprovechar la salida de Rio del cuadrante experiencial. Comunicar diferenciación frente al "
                    "vacío competitivo. (Pendiente DIF-9-1 confirmación Cora.)")

    add_h2(doc, "5.3 Recomendaciones — C+/C (\"clase media\")")
    add_h3(doc, "REC-CC1 — Activación específica con Rapidez en caja")
    add_bullet(doc, "Es el único atributo donde C+/C es MÁS exigente que el Total (+3.1pp en T2B importancia). "
                    "Comunicar rapidez/no-cola en caja para C+/C — atributo subutilizado en la comunicación actual.")

    add_h3(doc, "REC-CC2 — Reasignar peso de la inversión publicitaria hacia C+/C")
    add_bullet(doc, "Actualmente C+/C es el 19-20% de los recordadores de PTL/DTLS, vs 25.9% que representa de la "
                    "muestra. SUBALCANZADO. Revisar canales y medios — quizás la publicidad actual está en medios "
                    "más populares (NSE E sobrerepresentado).")

    add_h3(doc, "REC-CC3 — Amplificar lectura del DNA en C+/C")
    add_bullet(doc, "C+/C ve a Gama más fuerte en Atención (+9.8pp), Rapidez (+8.7pp) y Seguro (+6.9pp). El mensaje "
                    "que activa estos atributos en C+/C tiene mayor reverberación que en el Total.")

    add_h3(doc, "REC-CC4 — Defender empate técnico C+/C vs Plan Suárez + Páramo")
    add_bullet(doc, "Gama, Páramo y Plan Suárez están en empate técnico en preferencia C+/C (IC95 solapados). "
                    "Pérdidas en C+/C cuestan más que pérdidas en Total. Inversión defensiva justificada.")

    add_h2(doc, "5.4 Recomendaciones — Pref-Gama (Núcleo Leal)")
    add_h3(doc, "REC-PG1 — Profundizar basket en el Núcleo Leal")
    add_bullet(doc, "El Seg 3 (n=32, 8% del mercado, 100% pref-Gama) tiene perfil experiencial puro (78% atención, "
                    "alta frecuencia). Estrategia: aumentar categorías compradas por visita + frecuencia. "
                    "Mensaje: \"todo lo que necesitas, en el lugar donde ya confías\".")

    add_h3(doc, "REC-PG2 — Convertirlos en embajadores hacia el Seg 2")
    add_bullet(doc, "El Seg 2 (Pragmáticos Convertibles, 33%) tiene la menor resistencia al precio + 0% preferencia "
                    "Gama actual. La activación correcta: primera experiencia guiada (no descuento). El Núcleo Leal "
                    "puede ser fuente de referencia social.")

    add_h2(doc, "5.5 Dos lecturas válidas del estudio (decisión Owner)")
    add_chart(doc, os.path.join(CHARTS_V6, "C21_wow_movimiento.png"), width_in=6.5,
              caption="Movimientos significativos WoW 2025-2026 — Rio y Páramo crecen, CM cae, Gama estable.")

    add_para(doc, "Como en V5 y V6, el estudio admite dos lecturas estadísticamente correctas:", italic=True)
    add_bullet(doc, "Lectura defensiva: \"Gama es la única gran cadena con posición estable en un mercado turbulento. "
                    "DNA experiencial consolidado en 6 atributos. C+/C alineado.\"")
    add_bullet(doc, "Lectura ofensiva: \"Rio +17pp TOM y Páramo +12pp TOM. El cuadrante experiencial se está "
                    "vaciando — momento de capitalizar antes que algún competidor lo ocupe.\"")
    add_para(doc, "Ambas son correctas. La diferencia es estratégica, no estadística. Owner elige el tono según la "
                  "audiencia y momento de la conversación con Gama.", italic=True, color=GRAY_MID)

    # ===== AJUSTE 7 + 8 (Cora email 17:04 18/05): vías creativas precio + sensibilidad =====
    add_h2(doc, "5.6 Vías creativas para abordar precio — sin disminuir percepción de valor")

    add_callout_box(doc,
        "Esta sección responde a la observación explícita de Cora (email 18/05 17:04): cualquier abordaje del tema "
        "precio en Gama debe preservar la percepción de valor que sostiene su DNA experiencial. Las vías creativas "
        "permiten captar al no-cliente precio-sensible sin convertir a Gama en una cadena precio-dominante más.")

    add_h3(doc, "5.6.1 Principio: precio competitivo, no precio bajo")
    add_para(doc, "Gama NO debería buscar liderar precio frente a Páramo (gap percibido -27.9pp es estructural — "
                  "competir frontalmente requeriría sacrificio de margen incompatible con el modelo de servicio "
                  "experiencial). El objetivo es REDUCIR la brecha percibida lo suficiente para que el precio deje "
                  "de ser barrera de entrada al 92% que no prefiere Gama hoy, manteniendo el premium justificable "
                  "por experiencia. Como dijo Cora: \"Gama no debe distanciarse mucho en precio de sus competidores, "
                  "pero esto no es excluyente de un buen servicio\".")

    add_h3(doc, "5.6.2 Vías creativas propuestas (input de Cora + análisis V7)")

    add_h4(doc, "VC-1 — Esquemas promocionales diferenciados por categoría")
    add_para(doc, "Conectar con análisis Cap 3.4 (categorías 'cuidar precio' vs 'ofrecer valor'). Promociones más "
                  "agresivas en categorías 'cuidar precio' (Galletas, Productos básicos) donde la brecha es pequeña "
                  "y donde retener hábito vale la inversión. Promociones cero o reducidas en categorías 'ofrecer "
                  "valor' (Congelados, Gaseosas, Salsas) donde la palanca es comunicar experiencia.")

    add_h4(doc, "VC-2 — Reactivar coleccionables / loyalty con producto físico canjeable")
    add_para(doc, "Cora compartió contexto histórico: \"En el pasado les funcionó muy bien estos esquemas "
                  "(coleccionables de tickets para canjear por productos para la cocina y el hogar). Sin embargo "
                  "recientemente su éxito no fue rotundo\". Vale abrir conversación con Gama para entender: ¿qué "
                  "cambió entre ola exitosa y reciente? ¿Tipo de producto canjeable, mecánica, comunicación, "
                  "contexto económico? Este insumo sería el punto de partida para diseñar una nueva generación "
                  "de esquemas que conecten con el shopper actual.")

    add_h4(doc, "VC-3 — Reorientar el mensaje del anzuelo promocional (recuperado de Cap 5.2)")
    add_para(doc, "Mantener PTL/DTLS como recursos comunicacionales pero cambiar el wording del anzuelo: en lugar "
                  "de \"precio bajo\" (que refuerza distancia de Gama vs Páramo y activa la barrera identitaria "
                  "del Seg 2 — ver Cap 4), usar \"razón específica para venir esta semana\". Las promociones son el "
                  "driver #3 (OR=3.64**); el mensaje correcto debe activarlas sin posicionarse en el territorio "
                  "donde Gama estructuralmente no compite.")

    add_h4(doc, "VC-4 — Comunicación más asertiva del DNA experiencial")
    add_para(doc, "La campaña actual llega al segmento equivocado (NSE E sobrerrepresentado, C+/C subrepresentado "
                  "— ver Cap 4.2). Reasignar peso de inversión hacia canales que alcancen C+/C amplificaría el DNA "
                  "donde el shopper ya lo valora. Esta es la palanca de defensa estructural — independiente del "
                  "precio.")

    add_h4(doc, "VC-5 — Programas de lealtad visible que premien recurrencia (no descuento puntual)")
    add_para(doc, "Esquemas que premian la frecuencia de compra (puntos visibles, beneficios escalonados) en lugar "
                  "de descuentos puntuales. Esto refuerza el vínculo experiencial del Núcleo Leal (Seg 3) sin "
                  "comunicar precio bajo al mercado amplio. Compatible con la observación de Cora sobre "
                  "\"abordar precio sin disminuir percepción de valor\".")

    add_h2(doc, "5.7 Prerrequisito metodológico — análisis de sensibilidad de precio antes de ejecutar")

    add_callout_box(doc,
        "Cora pidió explícitamente (email 18/05 17:04): cualquier recomendación de ajuste de precio debe ir "
        "acompañada de un análisis de sensibilidad respecto a la percepción de competidores. Esta es una nota "
        "metodológica obligatoria antes de cualquier intervención de pricing operativo.",
        color=ORANGE_W)

    add_h3(doc, "5.7.1 Por qué es necesario")
    add_bullet(doc, "El análisis actual (P30/P32) mide percepción de precio, NO precio real ni elasticidad. Un "
                    "shopper que percibe a Gama como cara puede estar equivocado o puede estar en lo correcto — "
                    "la encuesta no distingue.")
    add_bullet(doc, "Cualquier cambio de precio operativo en Gama produce un nuevo equilibrio competitivo. Si "
                    "Páramo responde con su propio ajuste, la ventaja se neutraliza. El análisis de sensibilidad "
                    "permite anticipar respuestas competitivas y dimensionar el impacto neto.")
    add_bullet(doc, "Sin sensibilidad medida, las recomendaciones de precio son intuiciones. Con sensibilidad, "
                    "son decisiones informadas.")

    add_h3(doc, "5.7.2 Recomendación específica para ola 2027")
    add_para(doc, "Incorporar al diseño de la próxima ola (ya documentado en agenda 2027) métricas de elasticidad "
                  "y sensibilidad cruzada:")
    add_bullet(doc, "Van Westendorp PSM en 3 categorías KVI: Lácteos/Charcutería, Carne/Pollo, Granos. Esto permite "
                    "calcular Rango de Precio Aceptable (RPA) y Punto de Precio Óptimo (PPO) por categoría.")
    add_bullet(doc, "Gabor-Granger en SKUs específicos donde Gama considera ajustes (top 10 del basket).")
    add_bullet(doc, "Conjoint o DCE para medir sensibilidad cruzada Gama vs Páramo (USD 10-20K — ya en agenda 2028).")
    add_bullet(doc, "Auditoría de pricing in-store (Gama vs Páramo vs Plan Suárez vs Rio) en categorías KVI — "
                    "comparar precio real vs precio percibido, identificar gaps donde la percepción está peor que "
                    "la realidad (oportunidad de comunicación) o mejor que la realidad (riesgo si se descubre).")


def build_anexos(doc):
    add_h1(doc, "Anexo A — Geografía y municipios (secundario, no central)")
    add_para(doc, "Análisis geográfico del estudio — incluido como contexto, NO como capítulo central según pedido "
                  "explícito de Cora en la reunión 18/05 (\"céntrate primero en total C+/C... después como anexo\").")
    add_para(doc, "Distribución muestral: Baruta (122) · Libertador (80) · Sucre (79) · Chacao (70) · El Hatillo "
                  "(31) · Altos Mirandinos (20).")
    add_para(doc, "Gama tiene su mayor liderazgo de urgencia en Libertador (15.0%, líder) y Chacao (17.1%, #2). "
                  "Esto es consistente con presencia física en esas zonas. CV-WOW-002: composición geográfica difiere "
                  "entre olas 2025-2026 — interpretar con precaución.")

    add_h1(doc, "Anexo B — Hipótesis Express formato")
    add_para(doc, "Hipótesis de Cora: las misiones de reposición son fuertes en Gama por formato Express (tiendas "
                  "más pequeñas, variedad reducida).")
    add_para(doc, "Status: COMPATIBLE con datos pero NO PROBADA. La variable de formato de tienda no está en la "
                  "BBDD 2026. Datos compatibles: Gama lidera urgencia en C+/C (16.3%) y D (15.0%), no en E (7.6%). "
                  "Patrón consistente con uso de tienda de conveniencia en segmentos de mayor ingreso/menor tiempo. "
                  "Para validación formal se requiere base operacional Gama (lista de sucursales × formato) cruzada "
                  "con georreferenciación de respondientes.")

    add_h1(doc, "Anexo C — P43-P45 El Recreo (apertura sucursal)")
    add_callout_box(doc,
        "Análisis táctico — fuera del estudio principal. Las preguntas P43-P45 aplican solo a residentes/transeúntes "
        "de El Recreo (n=21). Datos puramente indicativos.", color=ORANGE_W)

    add_h3(doc, "P44 — Gama vs cadenas de la zona (n=21)")
    headers = ["Percepción", "n", "%"]
    rows = [
        ["Igual que cadenas de la zona", "16", "76.2%"],
        ["Mucho mejor", "2", "9.5%"],
        ["Mejor", "1", "4.8%"],
        ["Mucho peor", "2", "9.5%"],
    ]
    add_table(doc, headers, rows, col_widths=[7, 2, 2])

    add_h3(doc, "P45 — Disposición a comprar si Gama abre sucursal (n=21)")
    headers = ["Nivel", "n", "%"]
    rows = [
        ["Muy Dispuesto", "3", "14.3%"],
        ["Dispuesto", "10", "47.6%"],
        ["Algo dispuesto", "5", "23.8%"],
        ["Poco dispuesto", "2", "9.5%"],
        ["Nada dispuesto", "1", "4.8%"],
        ["Alta disposición (Muy + Dispuesto)", "13", "61.9%"],
    ]
    add_table(doc, headers, rows, col_widths=[5, 2, 2])
    add_caveat_box(doc, "Bloque anexo n=21 — datos NO proyectables al mercado total. Solo verificación táctica.")


def build_decisiones_pendientes(doc):
    add_h1(doc, "Notas finales y próximos pasos")

    add_h3(doc, "Decisiones aplicadas en V8 (sobre la base aprobada de V7)")
    add_para(doc, "Mantenido de V7 (los 8 ajustes aprobados por Cora email 18/05 17:04):", bold=True)
    add_bullet(doc, "Rio precio-dominante con hipótesis de apertura de tiendas (Cap 1.3).")
    add_bullet(doc, "Plan Suárez 15.4% C+/C destacado en Cap 1.2 y Cap 1.3.")
    add_bullet(doc, "DNA \"outlier estructural\" matizado (no excluyente con abordar precio) — Cap 2.4.")
    add_bullet(doc, "Cap 3 contexto modelo mental + Cap 5.6 vías creativas precio + Cap 5.7 prerrequisito sensibilidad.")
    add_bullet(doc, "Caveats Pref-Gama (n=32) y categorías P30/P32 REFERENCIAL visibles.")

    add_para(doc, "Nuevo en V8 (CU-10, email Cora 18/05 19:25):", bold=True)
    add_bullet(doc, "Cap 1.4 nuevo — P22 Top Box puro reranking: Menor precio sube +6 posiciones, Rapidez +4, "
                    "Atención y Valer dinero bajan -3 cada uno. La corrección metodológica de Cora revela presión "
                    "funcional más severa de lo que el T2B sugería.")
    add_bullet(doc, "Cap 2.5 nuevo — Mundo de marca por cadena (P23 imagen total × experiencia P24). Identifica "
                    "Atención como atributo de SOMBRA de Gama (+62.5 pp entre preferentes vs total).")
    add_bullet(doc, "Cap 2.6 nuevo — Vecindad perceptual: Plazas es el competidor MÁS CERCANO a Gama (mismo cluster "
                    "experiencial, mismo atributo sombra Atención). Rio es la amenaza latente con ventaja en calidad.")
    add_bullet(doc, "Cap 4.4 nuevo — Drivers REFORMULADOS: regresión logística confirma que las importancias "
                    "DECLARADAS no predicen preferencia (pseudo-R² 0.021, no significativo). El modelo correcto es "
                    "PERCEPCIÓN P23 → PREFERENCIA. La campaña debe demostrar que Gama TIENE atención, no que la "
                    "atención es importante (todos ya lo saben).")
    add_bullet(doc, "Paleta de tablas/headers corregida — rojo Gama saturado (#E30613) en lugar del vino tinto V7.")

    add_h3(doc, "Entregables que acompañan a este Word V8")
    add_bullet(doc, "Documento metodológico (.docx) — actualizado con notas TB puro vs T2B + regresión logística + mundo de marca.")
    add_bullet(doc, "HTML interactivo (.html) — vigente desde V7, sin cambios (la data bruta es la misma BBDD 2026).")
    add_bullet(doc, "Plan V7 (.docx) — vigente como referencia del scope acordado.")
    add_bullet(doc, "CU-10 raw report (.md) — análisis completo de Cuanti que alimenta los nuevos bloques V8.")

    add_h3(doc, "Próximo paso — PPTX V8")
    add_para(doc, "Una vez aprobado este Word V8 por Owner y Cora, Vivienne (ya promovida a Opus 4.7 + brand kit Gama "
                  "V0.1 cargado pre-flight + protocolo 3 invocaciones obligatorio para evitar token explosion) renderea "
                  "PPTX V8 aplicando Pirámide Minto explícita por bloque (lámina descriptiva → lámina analítica) + "
                  "infografía vs bullet-point + paleta corregida (#E30613 saturado, NO vino tinto).")


def main():
    print("Iniciando construccion del Word V8...")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    print("  Portada..."); build_portada(doc)
    print("  Executive summary (con 7 hallazgos V8)..."); build_executive(doc)
    print("  Marco V8..."); build_marco(doc)
    print("  Ficha tecnica..."); build_ficha(doc)
    print("  Capitulo 1 - Embudo + drivers + TB puro NUEVO..."); build_cap1(doc)
    print("  Capitulo 2 - Posicionamiento + DNA + Mundo de marca + Vecindad NUEVO..."); build_cap2(doc)
    print("  Capitulo 3 - Precios + categorias..."); build_cap3(doc)
    print("  Capitulo 4 - Comunicacion + Drivers reformulados NUEVO..."); build_cap4(doc)
    print("  Capitulo 5 - Sintesis bifurcada..."); build_cap5(doc)
    print("  Anexos A/B/C..."); build_anexos(doc)
    print("  Notas finales V8..."); build_decisiones_pendientes(doc)

    doc.save(OUTPUT)
    print(f"\nGuardado: {OUTPUT}")


if __name__ == '__main__':
    main()
