"""Lee R-/B-/M-pantallas para consulta.docx — texto + conteo de imágenes."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from pathlib import Path

BASE = Path(r'C:\RAUL\03-projects\genteca\GME Estudios de mercado')
files = ['R-pantallas para consulta.docx',
         'B-pantallas para consulta.docx',
         'M-pantallas para consulta.docx']

for f in files:
    p = BASE / f
    print('=' * 70)
    print(f'FILE: {f}')
    print('=' * 70)
    doc = Document(str(p))
    # Imágenes embebidas
    n_images = 0
    for rel in doc.part.rels.values():
        if 'image' in rel.target_ref.lower():
            n_images += 1
    print(f'Imágenes embebidas: {n_images}')
    print(f'Párrafos: {len(doc.paragraphs)}')
    print()
    print('--- TEXTO ---')
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t:
            print(f'[{i}] {t}')
    # Tablas si hay
    if doc.tables:
        print('\n--- TABLAS ---')
        for ti, table in enumerate(doc.tables):
            print(f'Tabla {ti}: {len(table.rows)}x{len(table.columns)}')
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    txt = cell.text.strip()
                    if txt:
                        print(f'  [{ri},{ci}]: {txt[:200]}')
    print()
